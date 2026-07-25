import json
import io

import pytest
from docx import Document
from fastapi import HTTPException

from routers import all_in_one, jobs
from utils import copy_gen
from utils.owned_page import build_owned_page_registry
from utils.page_quality import (
    ADAPTIVE_POLICY_VERSION,
    PAGE_QUALITY_POLICY_VERSION,
    get_page_quality_policy,
    select_guidance_profile,
)


TEMPLATE_SECTIONS = [
    {
        "name": "hero",
        "label": "Hero",
        "purpose": "Lead the page.",
        "heading_level": "h1",
        "word_count": [80, 120],
        "keyword_slot": "primary",
        "prompt_rules": "Lead directly.",
    },
    {
        "name": "benefits",
        "label": "Benefits",
        "purpose": "Explain practical value.",
        "heading_level": "h2",
        "word_count": [100, 160],
        "keyword_slot": "supporting",
        "prompt_rules": "Explain the value.",
    },
    {
        "name": "support",
        "label": "Support",
        "purpose": "Close with support context.",
        "heading_level": "none",
        "word_count": [40, 80],
        "keyword_slot": "none",
        "prompt_rules": "Write body copy.",
    },
]


def _registry():
    return build_owned_page_registry(
        "## Existing benefits\n"
        "The existing page explains calibration support and routine maintenance.\n\n"
        "## Other context\n"
        "The team also publishes practical setup guidance."
    )


def test_plan_fields_are_bounded_and_only_h2_h3_headings_survive():
    registry = _registry()
    brief = copy_gen._normalise_strategy_brief(
        {
            "section_guidance": [
                {
                    "section": "hero",
                    "planned_heading": "Model-selected H1",
                    "coverage_points": ["Lead point"],
                },
                {
                    "section": "benefits",
                    "responsibility": "Explain practical value.",
                    "planned_heading": "How Calibration Support Reduces Rework",
                    "coverage_points": [
                        "Explain the operational problem",
                        "Connect maintenance to decision confidence",
                    ],
                    "owned_block_ids": ["O1", "O999", "O1"],
                    "owned_blocks": [{"id": "O999", "excerpt": "invented"}],
                    "retain_points": ["Keep the maintenance idea"],
                    "improve_points": ["Make the reader implication clearer"],
                    "depth_policy": "model_selected",
                },
                {
                    "section": "support",
                    "responsibility": "Close with support context.",
                    "planned_heading": "Should not survive",
                },
            ]
        },
        template_sections=TEMPLATE_SECTIONS,
        owned_page_registry=registry,
    )

    contracts = {
        item["section"]: item for item in brief["section_guidance"]
    }
    assert "planned_heading" not in contracts["hero"]
    assert "planned_heading" not in contracts["support"]
    assert (
        contracts["benefits"]["planned_heading"]
        == "How Calibration Support Reduces Rework"
    )
    assert contracts["benefits"]["owned_block_ids"] == ["O1"]
    assert contracts["benefits"]["owned_blocks"][0]["excerpt"] == registry["blocks"][0]["excerpt"]
    assert "depth_policy" not in contracts["benefits"]
    rejected = brief["owned_page_mapping_diagnostics"]["rejected_assignments"]
    assert {item["reason"] for item in rejected} == {"unknown_id", "duplicate_id"}


def test_legacy_normalisation_discards_unsolicited_planning_fields():
    brief = copy_gen._normalise_strategy_brief({
        "section_guidance": [{
            "section": "benefits",
            "responsibility": "Explain practical value.",
            "planned_heading": "Model-selected heading",
            "coverage_points": ["Model-selected coverage"],
            "owned_block_ids": ["O1"],
            "depth_policy": "proof_only",
        }]
    })

    assert brief["section_guidance"] == [{
        "section": "benefits",
        "responsibility": "Explain practical value.",
    }]
    assert "owned_page_mapping_diagnostics" not in brief


def test_markup_overlong_and_cross_section_owned_assignments_are_rejected():
    registry = _registry()
    brief = copy_gen._normalise_strategy_brief(
        {
            "section_guidance": [
                {
                    "section": "benefits",
                    "responsibility": "First contract.",
                    "planned_heading": "## Markup heading",
                    "owned_block_ids": ["O1"],
                },
                {
                    "section": "support",
                    "responsibility": "Second contract.",
                    "owned_block_ids": ["O1"],
                },
            ]
        },
        template_sections=TEMPLATE_SECTIONS,
        owned_page_registry=registry,
    )

    first, second = brief["section_guidance"]
    assert "planned_heading" not in first
    assert first["owned_block_ids"] == ["O1"]
    assert "owned_block_ids" not in second
    assert any(
        item["reason"] == "already_assigned"
        for item in brief["owned_page_mapping_diagnostics"]["rejected_assignments"]
    )


@pytest.mark.parametrize(
    "planned_heading",
    [
        "**Bold heading**",
        "_Italic heading_",
        "`Code heading`",
        "[Linked heading](https://example.com)",
        "![Image heading](https://example.com/image.png)",
        "> Quoted heading",
        "- List heading",
    ],
)
def test_planned_heading_rejects_markdown_not_just_heading_markers(
    planned_heading,
):
    brief = copy_gen._normalise_strategy_brief(
        {
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain practical value.",
                "planned_heading": planned_heading,
            }]
        },
        template_sections=[TEMPLATE_SECTIONS[1]],
    )

    assert "planned_heading" not in brief["section_guidance"][0]


def test_planning_ignores_unknown_and_duplicate_section_contracts():
    registry = _registry()
    brief = copy_gen._normalise_strategy_brief(
        {
            "section_guidance": [
                {
                    "section": "unknown",
                    "responsibility": "Must not consume mapped source.",
                    "owned_block_ids": ["O1"],
                },
                {
                    "section": "benefits",
                    "responsibility": "Use the canonical first contract.",
                    "planned_heading": "How Calibration Support Reduces Rework",
                    "owned_block_ids": ["O1"],
                },
                {
                    "section": "benefits",
                    "responsibility": "Conflicting duplicate.",
                    "planned_heading": "A Different Heading",
                    "owned_block_ids": ["O2"],
                },
            ]
        },
        template_sections=[TEMPLATE_SECTIONS[1]],
        owned_page_registry=registry,
    )

    assert len(brief["section_guidance"]) == 1
    contract = brief["section_guidance"][0]
    assert contract["responsibility"] == "Use the canonical first contract."
    assert contract["planned_heading"] == "How Calibration Support Reduces Rework"
    assert contract["owned_block_ids"] == ["O1"]


def test_page_refresh_keeps_related_attorney_profiles_in_one_section():
    registry = build_owned_page_registry(
        "## About Attorney Saman Dhukka\n"
        "Saman Dhukka focuses on estate planning and probate.\n\n"
        "## About Attorney Ali Dhukka\n"
        "Ali Dhukka handles personal injury and business litigation."
    )
    template_sections = [
        {
            **TEMPLATE_SECTIONS[1],
            "name": "solution",
            "label": "Legal Services",
        },
        {
            **TEMPLATE_SECTIONS[1],
            "name": "social_proof",
            "label": "Attorney Profiles",
        },
    ]
    brief = copy_gen._normalise_strategy_brief(
        {
            "verified_facts": [
                {
                    "id": "F1",
                    "fact": (
                        "Saman Dhukka focuses on estate planning and probate."
                    ),
                    "source": "current_page",
                    "source_excerpt": (
                        "Saman Dhukka focuses on estate planning and probate."
                    ),
                },
                {
                    "id": "F2",
                    "fact": (
                        "The second attorney handles personal injury and "
                        "business litigation."
                    ),
                    "source": "current_page",
                    "source_excerpt": (
                        "Ali Dhukka handles personal injury and business "
                        "litigation."
                    ),
                },
            ],
            "proof_fact_ids": ["F1", "F2"],
            "section_guidance": [
                {
                    "section": "solution",
                    "responsibility": "Explain the firm's services.",
                    "proof_fact_ids": ["F2"],
                    "owned_block_ids": ["O2"],
                },
                {
                    "section": "social_proof",
                    "responsibility": "Introduce the attorneys.",
                    "proof_fact_ids": ["F1"],
                    "owned_block_ids": ["O1"],
                    "planned_heading": (
                        "Attorney Saman Dhukka's Estate Planning Focus"
                    ),
                },
            ],
        },
        template_sections=template_sections,
        owned_page_registry=registry,
        page_copy_correction_enabled=True,
    )

    contracts = {
        contract["section"]: contract
        for contract in brief["section_guidance"]
    }
    profile_contract = contracts["social_proof"]
    assert profile_contract["owned_block_ids"] == ["O1", "O2"]
    assert [
        fact["id"] for fact in profile_contract["proof_facts"]
    ] == ["F1", "F2"]
    assert profile_contract["required_named_items"] == [
        "Saman Dhukka",
        "Ali Dhukka",
    ]
    assert (
        profile_contract["planned_heading"]
        == "Meet Attorneys Saman Dhukka and Ali Dhukka"
    )
    assert "owned_block_ids" not in contracts["solution"]
    assert "proof_facts" not in contracts["solution"]
    flags = []
    all_in_one._add_page_plan_qa_flags(
        flags,
        {
            "solution": "## Legal Services\nGeneral service context.",
            "social_proof": (
                "## Meet Attorneys Saman Dhukka and Ali Dhukka\n"
                "Saman Dhukka focuses on estate planning and probate."
            ),
        },
        {"sections": template_sections},
        brief,
        get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
        page_copy_correction_enabled=True,
    )
    missing_profile = next(
        flag
        for flag in flags
        if flag.get("code") == "page_related_profile_missing"
        and flag.get("section") == "social_proof"
    )
    assert missing_profile["missing_items"] == ["Ali Dhukka"]


def test_section_rerun_outcome_reports_media_limits_and_missing_people():
    evidence_context = {
        "requested_names": ["Saman Dhukka", "Ali Dhukka"],
        "matched_verified_fact_ids": ["F1", "F2"],
        "media_requested": True,
    }

    partial = copy_gen._reviewer_instruction_outcome(
        "Saman Dhukka and Ali Dhukka are both included.",
        evidence_context,
    )
    blocked = copy_gen._reviewer_instruction_outcome(
        "A generic replacement that omits both requested people.",
        {**evidence_context, "media_requested": False},
    )

    assert partial["status"] == "partially_applied"
    assert "image placement is not supported" in partial["message"]
    assert blocked["status"] == "blocked"
    assert blocked["missing_names"] == ["Saman Dhukka", "Ali Dhukka"]
    assert "instruction" not in partial


def test_plan_or_mapping_advisories_do_not_invalidate_strategy_readiness():
    brief = copy_gen._normalise_strategy_brief(
        {
            "search_intent": "Commercial",
            "page_goal": "Explain the service.",
            "primary_positioning": "Practical calibration support.",
            "headline_direction": "Lead with practical support.",
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain value.",
                "planned_heading": "Benefits",
                "owned_block_ids": ["O999"],
            }],
        },
        template_sections=[TEMPLATE_SECTIONS[1]],
        owned_page_registry=_registry(),
    )

    assert copy_gen.strategy_brief_issues(
        brief,
        [TEMPLATE_SECTIONS[1]],
        {"page_copy"},
    ) == []
    diagnostics = copy_gen.page_plan_diagnostics(
        brief,
        [TEMPLATE_SECTIONS[1]],
    )
    assert {finding["code"] for finding in diagnostics["findings"]} == {
        "planned_heading_generic",
    }
    assert diagnostics["owned_page_mapping"]["rejected_assignments"]


def test_plan_only_fields_do_not_make_a_missing_section_contract_ready():
    brief = copy_gen._normalise_strategy_brief(
        {
            "search_intent": "Commercial",
            "page_goal": "Explain the service.",
            "primary_positioning": "Practical calibration support.",
            "headline_direction": "Lead with practical support.",
            "section_guidance": [{
                "section": "benefits",
                "planned_heading": "How Calibration Support Reduces Rework",
                "coverage_points": ["Explain the operational problem"],
                "owned_block_ids": ["O1"],
            }],
        },
        template_sections=[TEMPLATE_SECTIONS[1]],
        owned_page_registry=_registry(),
    )

    assert copy_gen.strategy_brief_issues(
        brief,
        [TEMPLATE_SECTIONS[1]],
        {"page_copy"},
    ) == ["Section contracts are missing for: benefits."]


def test_section_prompt_uses_exact_heading_coverage_mapping_depth_and_guidance():
    registry = _registry()
    brief = copy_gen._normalise_strategy_brief(
        {
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain practical value.",
                "planned_heading": "How Calibration Support Reduces Rework",
                "coverage_points": [
                    "Explain the operational problem",
                    "Connect maintenance to decision confidence",
                ],
                "owned_block_ids": ["O1"],
                "retain_points": ["Keep the maintenance idea"],
                "improve_points": ["Make the reader implication clearer"],
            }]
        },
        template_sections=TEMPLATE_SECTIONS,
        owned_page_registry=registry,
    )
    brief["section_guidance"][0]["depth_policy"] = "explanatory"
    prompt = copy_gen._build_section_prompt(
        section={
            **TEMPLATE_SECTIONS[1],
            "adaptive_instruction": "Retain the planned explanatory depth.",
        },
        primary_keyword="calibration support",
        supporting_keyword="maintenance planning",
        lsi_keywords=[],
        business_type="b2b",
        brand_name="Example",
        h1="Calibration Support",
        page_type="service",
        paa_questions=[],
        competitor_excerpts=[],
        client_brief="",
        previous_section_text="",
        client_existing_content="",
        strategy_brief=brief,
        page_copy_guidance=select_guidance_profile("editorial_refresh"),
        page_quality_policy=get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
    )

    assert (
        "Start exactly with this H2 heading on the first line:\n"
        "## How Calibration Support Reduces Rework\n"
        "Do not rename, paraphrase, punctuate, or repeat it."
    ) in prompt
    assert "Coverage contract:" in prompt
    assert "Explain the operational problem" in prompt
    assert registry["blocks"][0]["excerpt"] in prompt
    assert "editorial source only; not an evidence allowlist" in prompt
    assert "Selected CopyPilot page-copy guidance" in prompt
    assert "Apply editorial-refresh page-copy guidance" in prompt


def test_guidance_profile_alone_does_not_activate_versioned_plan_behavior():
    brief = {
        "section_guidance": [{
            "section": "benefits",
            "responsibility": "Explain practical value.",
            "planned_heading": "How Calibration Support Reduces Rework",
            "coverage_points": ["Explain the operational problem"],
        }]
    }

    prompt = copy_gen._build_section_prompt(
        section=TEMPLATE_SECTIONS[1],
        primary_keyword="calibration support",
        supporting_keyword="maintenance planning",
        lsi_keywords=[],
        business_type="b2b",
        brand_name="Example",
        h1="Calibration Support",
        page_type="service",
        paa_questions=[],
        competitor_excerpts=[],
        client_brief="",
        previous_section_text="",
        client_existing_content="",
        strategy_brief=brief,
        page_copy_guidance=select_guidance_profile("editorial_refresh"),
    )

    assert "How Calibration Support Reduces Rework" not in prompt
    assert "Coverage contract:" not in prompt
    assert "Apply editorial-refresh page-copy guidance" in prompt


def test_page_plan_fields_do_not_enter_meta_prompt_formatting():
    brief = {
        "primary_positioning": "Practical calibration support.",
        "page_goal": "Explain the service.",
        "search_intent": "Commercial",
        "section_guidance": [{
            "section": "benefits",
            "planned_heading": "How Calibration Support Reduces Rework",
            "coverage_points": ["Explain the operational problem"],
            "owned_blocks": [{"id": "O1", "excerpt": "Owned source excerpt"}],
        }],
    }

    meta_prompt = copy_gen.format_strategy_brief_for_prompt(
        brief,
        output_type="meta",
    )

    assert "Practical calibration support." in meta_prompt
    assert "How Calibration Support Reduces Rework" not in meta_prompt
    assert "Owned source excerpt" not in meta_prompt


def test_h1_and_none_heading_contracts_remain_server_controlled():
    h1_prompt = copy_gen._build_section_prompt(
        section=TEMPLATE_SECTIONS[0],
        primary_keyword="calibration support",
        supporting_keyword="",
        lsi_keywords=[],
        business_type="b2b",
        brand_name="",
        h1="Canonical Calibration Support",
        page_type="service",
        paa_questions=[],
        competitor_excerpts=[],
        client_brief="",
        previous_section_text="",
        client_existing_content="",
        strategy_brief={
            "section_guidance": [{
                "section": "hero",
                "planned_heading": "Model H1",
            }]
        },
        page_copy_guidance=select_guidance_profile("balanced"),
        page_quality_policy=get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
    )
    none_prompt = copy_gen._build_section_prompt(
        section=TEMPLATE_SECTIONS[2],
        primary_keyword="",
        supporting_keyword="",
        lsi_keywords=[],
        business_type="b2b",
        brand_name="",
        h1="Canonical Calibration Support",
        page_type="service",
        paa_questions=[],
        competitor_excerpts=[],
        client_brief="",
        previous_section_text="",
        client_existing_content="",
        strategy_brief={},
        page_copy_guidance=select_guidance_profile("balanced"),
        page_quality_policy=get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
    )

    assert "Start exactly with this canonical H1: # Canonical Calibration Support" in h1_prompt
    assert "Model H1" not in h1_prompt
    assert "Do not add a heading. Write body copy only." in none_prompt


def test_legacy_h1_and_h3_prompt_contracts_remain_exactly_historical():
    common = {
        "primary_keyword": "calibration support",
        "supporting_keyword": "",
        "lsi_keywords": [],
        "business_type": "b2b",
        "brand_name": "",
        "h1": "Canonical Calibration Support",
        "page_type": "service",
        "paa_questions": [],
        "competitor_excerpts": [],
        "client_brief": "",
        "previous_section_text": "",
        "client_existing_content": "",
        "strategy_brief": {},
    }
    h1_prompt = copy_gen._build_section_prompt(
        section=TEMPLATE_SECTIONS[0],
        **common,
    )
    h3_prompt = copy_gen._build_section_prompt(
        section={**TEMPLATE_SECTIONS[1], "heading_level": "h3"},
        **common,
    )

    assert "Start with the H1 headline (# in markdown)." in h1_prompt
    assert "Start exactly with this canonical H1" not in h1_prompt
    assert "Use H3 subheadings (### in markdown) where appropriate." in h3_prompt


def test_v1_canonical_h1_is_enforced_for_rewritten_or_missing_headings():
    template = {"sections": [TEMPLATE_SECTIONS[0]]}

    rewritten, rewritten_changed = all_in_one._enforce_v1_canonical_page_h1(
        {"hero": "# Model Rewrite\nUseful body copy."},
        template,
        "Canonical Calibration Support",
    )
    missing, missing_changed = all_in_one._enforce_v1_canonical_page_h1(
        {"hero": "Useful body copy without a heading."},
        template,
        "Canonical Calibration Support",
    )

    assert rewritten_changed is True
    assert rewritten["hero"].startswith("# Canonical Calibration Support\n")
    assert missing_changed is True
    assert missing["hero"].startswith("# Canonical Calibration Support\n")


def test_v1_h1_qa_flags_missing_and_noncanonical_headings():
    template = {"sections": [TEMPLATE_SECTIONS[0]]}
    common = {
        "gen_meta": False,
        "gen_faqs": False,
        "gen_page_copy": True,
        "generated_title": "",
        "generated_description": "",
        "optimised_h1": "",
        "input_h1": "Canonical Calibration Support",
        "primary_keyword": "calibration support",
        "faq_items": [],
        "forbidden_phrases": [],
        "template": template,
        "strategy_brief": {},
        "page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION,
    }

    missing_codes = {
        flag["code"]
        for flag in all_in_one._collect_qa_flags(
            **common,
            section_results={"hero": "Body copy without a heading."},
        )
    }
    mismatch_codes = {
        flag["code"]
        for flag in all_in_one._collect_qa_flags(
            **common,
            section_results={"hero": "# Model Rewrite\nBody copy."},
        )
    }

    assert "page_h1_missing" in missing_codes
    assert "page_h1_canonical_mismatch" in mismatch_codes


def test_planning_schema_is_opt_in_and_keeps_profile_out_of_shared_strategy():
    prompts = []

    def provider(_api_key, prompt, **_kwargs):
        prompts.append(prompt)
        return json.dumps({
            "search_intent": "Commercial",
            "page_goal": "Explain the service.",
            "primary_positioning": "Practical support.",
            "headline_direction": "Lead with the service.",
            "meta_direction": "Describe the service.",
            "faq_direction": "Answer fit questions.",
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain value.",
                "planned_heading": "Practical Value for Operations Teams",
                "coverage_points": ["Explain the value"],
            }],
        })

    copy_gen.PROVIDER_FN["PlanTest"] = provider
    copy_gen.generate_strategy_brief(
        provider="PlanTest",
        api_key="key",
        url="https://example.com",
        keyword="calibration support",
        page_type="service",
        business_type="b2b",
        brand_name="Example",
        template_sections=[TEMPLATE_SECTIONS[1]],
        required_outputs=["meta", "page_copy"],
    )
    copy_gen.generate_strategy_brief(
        provider="PlanTest",
        api_key="key",
        url="https://example.com",
        keyword="calibration support",
        page_type="service",
        business_type="b2b",
        brand_name="Example",
        template_sections=[TEMPLATE_SECTIONS[1]],
        required_outputs=["meta", "page_copy"],
        enable_page_planning=True,
        owned_page_registry=_registry(),
    )

    assert '"planned_heading"' not in prompts[0]
    assert '"planned_heading"' in prompts[1]
    assert "O1" in prompts[1]
    assert "Apply balanced page-copy guidance" not in prompts[1]
    assert "required_named_items" not in prompts[1]
    assert "Facts are ceilings" not in prompts[1]


def test_initial_quality_strategy_requires_complete_source_evidence_and_specific_headings():
    prompts = []

    def provider(_api_key, prompt, **_kwargs):
        prompts.append(prompt)
        return json.dumps({
            "search_intent": "Commercial",
            "page_goal": "Help visitors choose a supported path.",
            "primary_positioning": "Practical production support.",
            "headline_direction": "Lead with the supported offer.",
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain the supported options.",
                "planned_heading": "Production Options for Different Project Needs",
                "coverage_points": ["Preserve every named product path"],
            }],
        })

    copy_gen.PROVIDER_FN["InitialQualityStrategyTest"] = provider
    registry = build_owned_page_registry(
        "## Product paths\n\n"
        "- Fabrics\n"
        "- Tape\n"
        "- Rentals\n\n"
        "## Resources\n\n"
        "- Know How Blog\n"
        "- Contact Us"
    )

    copy_gen.generate_strategy_brief(
        provider="InitialQualityStrategyTest",
        api_key="key",
        url="https://example.com",
        keyword="production supplies",
        page_type="homepage",
        business_type="b2b",
        brand_name="Example",
        template_sections=[TEMPLATE_SECTIONS[1]],
        required_outputs=["page_copy"],
        enable_page_planning=True,
        owned_page_registry=registry,
        page_quality_policy=get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
    )

    prompt = prompts[0]
    assert "Facts are ceilings" in prompt
    assert "every claim's subject, predicate, qualifier, cause, comparison, and outcome" in prompt
    assert "Put every relevant exact label" in prompt
    assert "no partial lists or labels outside the assigned block" in prompt
    assert "named resource, attributed proof, and next-step path" in prompt
    assert "Testimonials support only attributed sentiment" in prompt
    assert "Planning fields may organise verified material but add no client capability" in prompt
    assert "required_named_items" in prompt
    assert "generic readiness question" in prompt


def test_collection_planning_uses_preselected_section_keywords_without_reranking():
    prompts = []

    def provider(_api_key, prompt, **_kwargs):
        prompts.append(prompt)
        return json.dumps({
            "search_intent": "Commercial",
            "page_goal": "Help shoppers understand the collection.",
            "primary_positioning": "Lead with the category.",
            "headline_direction": "Use a direct category H1.",
            "section_guidance": [
                {
                    "section": "collection_story",
                    "responsibility": "Connect the category to the shopper.",
                    "planned_heading": "Personalised Knives for Meaningful Gifts",
                    "coverage_points": ["Explain the supported gifting angle"],
                },
                {
                    "section": "collection_value",
                    "responsibility": "Add non-promotional category context.",
                    "planned_heading": "Custom Engraved Pocket Knife Styles",
                    "coverage_points": ["Explain the supported category context"],
                },
            ],
        })

    copy_gen.PROVIDER_FN["CollectionHeadingKeywordTest"] = provider
    copy_gen.generate_strategy_brief(
        provider="CollectionHeadingKeywordTest",
        api_key="key",
        url="https://example.com/collections/knives",
        keyword="personalised knives",
        page_type="collection",
        business_type="ecommerce",
        brand_name="Example",
        template_sections=[
            {
                "name": "category_intro",
                "label": "Category Introduction",
                "purpose": "Introduce the category.",
                "heading_level": "h1",
            },
            {
                "name": "collection_story",
                "label": "Collection Story",
                "purpose": "Connect the category to customer motivation.",
                "heading_level": "h2",
            },
            {
                "name": "collection_value",
                "label": "Collection Context",
                "purpose": "Add useful non-promotional category depth.",
                "heading_level": "h2",
            },
            {
                "name": "collection_guidance",
                "label": "Helpful Buying Notes",
                "purpose": "Close with supported buying guidance.",
                "heading_level": "h2",
            },
        ],
        section_heading_keyword_assignments={
            "collection_story": "personalised knife gifts",
            "collection_value": "custom engraved pocket knives",
        },
        required_outputs=["page_copy"],
        enable_page_planning=True,
        page_quality_policy=get_page_quality_policy(
            PAGE_QUALITY_POLICY_VERSION
        ),
        page_copy_correction_enabled=True,
    )

    prompt = prompts[0]
    assert (
        "collection_story: Collection Story. Connect the category to customer "
        "motivation. Already-selected heading keyword: personalised knife gifts."
        in prompt
    )
    assert (
        "collection_value: Collection Context. Add useful non-promotional "
        "category depth. "
        "Already-selected heading keyword: custom engraved pocket knives."
        in prompt
    )
    assert "Do not select, replace, or rerank these keywords" in prompt
    assert "use its assigned keyword or a close grammatical variant" in prompt
    assert (
        "assign shipping, delivery, returns, discounts, sales, coupons, "
        "promotions, and other store incentives only to collection_guidance"
        in prompt
    )


def test_collection_promotional_language_is_reviewed_outside_guidance():
    template = all_in_one.get_template("collection_page")
    flags = all_in_one._collect_qa_flags(
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        generated_title="",
        generated_description="",
        optimised_h1="",
        input_h1="",
        primary_keyword="bear and son knives",
        faq_items=[],
        section_results={
            "category_intro": " ".join(["intro"] * 60),
            "collection_story": " ".join(["story"] * 80),
            "collection_value": (
                "## Practical Knife Options\n\n"
                "Orders over $50 receive free shipping alongside the supported "
                "category information."
            ),
            "collection_guidance": " ".join(["guidance"] * 45),
        },
        forbidden_phrases=[],
        template=template,
        page_type="collection",
        strategy_brief={},
    )

    matching = [
        flag
        for flag in flags
        if flag.get("code") == "collection_promotion_outside_guidance"
    ]

    assert len(matching) == 1
    assert matching[0]["section"] == "collection_value"
    assert matching[0]["severity"] == "review"


def test_collection_guidance_uses_strict_body_word_range():
    template = all_in_one.get_template("collection_page")
    flags = []
    all_in_one._add_section_word_count_flags(
        flags,
        {
            "collection_guidance": (
                "## Supported Buying Details\n\n"
                + " ".join(["guidance"] * 51)
            ),
        },
        template,
    )

    assert len(flags) == 1
    assert flags[0]["code"] == "section_word_count_above_target"
    assert flags[0]["section"] == "collection_guidance"
    assert flags[0]["actual_words"] == 51
    assert flags[0]["target_min"] == 40
    assert flags[0]["target_max"] == 50
    assert flags[0]["severity"] == "review"

    within_range_flags = []
    all_in_one._add_section_word_count_flags(
        within_range_flags,
        {
            "collection_guidance": (
                "## Supported Buying Details\n\n"
                + " ".join(["guidance"] * 50)
            ),
        },
        template,
    )
    assert within_range_flags == []


def test_initial_quality_planning_prompt_incremental_overhead_stays_bounded(
    monkeypatch,
):
    prompts = []

    def provider(_api_key, prompt, **_kwargs):
        prompts.append(prompt)
        return json.dumps({
            "search_intent": "Commercial",
            "page_goal": "Help visitors choose a supported path.",
            "primary_positioning": "Practical production support.",
            "headline_direction": "Lead with the supported offer.",
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain the supported options.",
                "planned_heading": "Production Options for Different Project Needs",
                "coverage_points": ["Preserve every named product path"],
            }],
        })

    monkeypatch.setitem(copy_gen.PROVIDER_FN, "PromptBudgetTest", provider)
    common = {
        "provider": "PromptBudgetTest",
        "api_key": "key",
        "url": "https://example.com",
        "keyword": "production supplies",
        "page_type": "homepage",
        "business_type": "b2b",
        "brand_name": "Example",
        "template_sections": [TEMPLATE_SECTIONS[1]],
        "required_outputs": ["page_copy"],
        "enable_page_planning": True,
        "owned_page_registry": _registry(),
    }

    copy_gen.generate_strategy_brief(**common)
    copy_gen.generate_strategy_brief(
        **common,
        page_quality_policy=get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
    )

    planning_prompt, quality_prompt = prompts
    quality_overhead = len(quality_prompt) - len(planning_prompt)

    assert 0 < quality_overhead <= 1_600


def test_assigned_proof_keeps_exact_source_boundary_and_validated_named_items():
    page_context = (
        "## Product paths\n\n"
        "- Fabrics\n"
        "- Tape\n"
        "- Rentals"
    )
    registry = build_owned_page_registry(page_context)
    source_excerpt = "Fabrics\n- Tape\n- Rentals"
    brief = copy_gen._normalise_strategy_brief(
        {
            "verified_facts": [{
                "id": "F1",
                "fact": "Product paths include Fabrics, Tape, and Rentals.",
                "source": "current_page",
                "source_excerpt": source_excerpt,
            }],
            "proof_fact_ids": ["F1"],
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain the available product paths.",
                "proof_fact_ids": ["F1"],
                "owned_block_ids": ["O1"],
                "required_named_items": [
                    "Fabrics",
                    "Tape",
                    "Rentals",
                    "Fabric",
                    "Invented Path",
                ],
            }],
        },
        evidence_sources={
            "current_page": page_context,
            "client_brief": "",
            "brand_profile": "",
        },
        template_sections=[TEMPLATE_SECTIONS[1]],
        owned_page_registry=registry,
    )

    contract = brief["section_guidance"][0]
    assert contract["required_named_items"] == ["Fabrics", "Tape", "Rentals"]
    assert contract["proof_facts"] == [{
        "id": "F1",
        "fact": "Product paths include Fabrics, Tape, and Rentals.",
        "source": "current_page",
        "source_excerpt": source_excerpt,
    }]

    prompt = copy_gen._build_section_prompt(
        section={
            **TEMPLATE_SECTIONS[1],
            "word_count": [250, 450],
            "depth_policy": "claim_sensitive",
            "adaptive_mode": "full",
        },
        primary_keyword="production supplies",
        supporting_keyword="",
        lsi_keywords=[],
        business_type="b2b",
        brand_name="Example",
        h1="Production Supplies",
        page_type="homepage",
        paa_questions=[],
        competitor_excerpts=[],
        client_brief="",
        previous_section_text="",
        client_existing_content="",
        strategy_brief=brief,
        page_quality_policy=get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
        initial_generation_quality_contract=True,
    )

    assert "Exact supporting excerpt:" in prompt
    assert source_excerpt in prompt
    assert "Required source names and paths:" in prompt
    assert "Invented Path" not in prompt
    assert "A proof point is a ceiling, not a seed" in prompt
    assert "Testimonials authorize only the attributed statement" in prompt


def test_initial_quality_prompt_targets_visible_body_depth_and_complete_ctas_only_once():
    page_policy = get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION)
    common = {
        "primary_keyword": "production supplies",
        "supporting_keyword": "",
        "lsi_keywords": [],
        "business_type": "b2b",
        "brand_name": "Example",
        "h1": "Production Supplies",
        "page_type": "homepage",
        "paa_questions": [],
        "competitor_excerpts": [],
        "client_brief": "",
        "previous_section_text": "",
        "client_existing_content": "",
        "page_quality_policy": page_policy,
    }
    body_section = {
        **TEMPLATE_SECTIONS[1],
        "word_count": [250, 450],
        "depth_policy": "claim_sensitive",
        "adaptive_mode": "full",
    }

    initial_prompt = copy_gen._build_section_prompt(
        section=body_section,
        strategy_brief={
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain the supported production options.",
            }],
        },
        initial_generation_quality_contract=True,
        **common,
    )
    rerun_prompt = copy_gen._build_section_prompt(
        section=body_section,
        strategy_brief={
            "section_guidance": [{
                "section": "benefits",
                "responsibility": "Explain the supported production options.",
            }],
        },
        **common,
    )

    assert "Deliver about 350 visible words" in initial_prompt
    assert "Only returned section copy counts toward this target" in initial_prompt
    assert "Deliver about 350 visible words" not in rerun_prompt
    assert (
        "Treat 250 words as the expected depth when the available evidence supports it"
        in rerun_prompt
    )

    paths = [
        "Contact Us",
        "Custom Curtain Quote Request",
        "Digital Printing Quote Request",
        "Fabric Finder",
        "Project Portfolio",
    ]
    cta_prompt = copy_gen._build_section_prompt(
        section={
            "name": "cta_close",
            "label": "Closing CTA",
            "purpose": "Direct the reader to a supported next step.",
            "heading_level": "h2",
            "word_count": [80, 140],
            "keyword_slot": "none",
            "prompt_rules": "Keep the close concise.",
            "depth_policy": "explanatory",
            "adaptive_mode": "full",
        },
        strategy_brief={
            "section_guidance": [{
                "section": "cta_close",
                "required_named_items": paths,
                "proof_points": ["The five named next-step paths are available."],
            }],
        },
        initial_generation_quality_contract=True,
        **common,
    )

    assert "Every CTA instruction must be a complete grammatical sentence" in cta_prompt
    assert "short introduction followed by bullets or separate complete sentences" in cta_prompt
    assert all(path in cta_prompt for path in paths)
    assert "Deliver about 110 visible words" not in cta_prompt


def test_initial_quality_page_generation_never_adds_a_repair_call():
    calls = []

    def provider(_api_key, prompt, **kwargs):
        calls.append((prompt, kwargs))
        return "## Supported Production Options\nStill deciding, use the supported path."

    copy_gen.PROVIDER_FN["InitialQualityCallCountTest"] = provider
    result = copy_gen.generate_page(
        template={
            "sections": [{
                "name": "benefits",
                "label": "Benefits",
                "purpose": "Explain supported options.",
                "heading_level": "h2",
                "word_count": [100, 160],
                "keyword_slot": "none",
                "prompt_rules": "Write directly.",
                "depth_policy": "claim_sensitive",
                "adaptive_mode": "full",
            }],
        },
        keyword_assignment={"benefits": {}},
        lsi_keywords={},
        business_type="b2b",
        brand_name="Example",
        h1="Production Options",
        page_type="homepage",
        paa_questions=[],
        ai_overview="",
        competitor_section_map={},
        client_brief="",
        client_existing_content="",
        provider="InitialQualityCallCountTest",
        api_key="key",
        model="quality-model",
        strategy_brief={
            "verified_facts": [{
                "id": "F1",
                "fact": "A supported production path exists.",
                "source": "current_page",
                "source_excerpt": "supported production path",
            }],
            "section_guidance": [{
                "section": "benefits",
                "planned_heading": "Supported Production Options",
                "proof_points": ["A supported production path exists."],
            }],
        },
        page_copy_guidance=select_guidance_profile("balanced"),
        page_quality_policy=get_page_quality_policy(PAGE_QUALITY_POLICY_VERSION),
    )

    assert len(calls) == 1
    assert calls[0][1]["model"] == "quality-model"
    assert result["benefits"].endswith("Still deciding, use the supported path.")


def test_owned_page_strategy_context_is_hard_bounded():
    registry = {
        "blocks": [
            {
                "id": f"O{index}",
                "heading": f"{'H' * 118}{index}",
                "excerpt": "x" * 300,
            }
            for index in range(1, 25)
        ]
    }

    bounded_registry = copy_gen._strategy_brief_prompt_registry(registry)
    formatted = copy_gen._strategy_brief_owned_page_block(bounded_registry)

    assert len(formatted) <= copy_gen.STRATEGY_BRIEF_PAGE_CONTEXT_CHAR_LIMIT
    assert bounded_registry["prompt_truncated"] is True
    assert len(bounded_registry["blocks"]) < len(registry["blocks"])
    assert len(formatted.split("\n\n")) == len(bounded_registry["blocks"])
    assert all(line.endswith("x" * 300) for line in formatted.split("\n\n"))


def test_deterministic_heading_and_depth_flags_are_version_gated():
    template = {"sections": [TEMPLATE_SECTIONS[1]]}
    strategy = {
        "section_guidance": [{
            "section": "benefits",
            "planned_heading": "How Calibration Support Reduces Rework",
        }]
    }
    legacy_flags = all_in_one._collect_qa_flags(
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        generated_title="",
        generated_description="",
        optimised_h1="",
        input_h1="",
        primary_keyword="calibration support",
        faq_items=[],
        section_results={"benefits": "## Generic Benefits\nShort copy."},
        forbidden_phrases=[],
        template=template,
        strategy_brief=strategy,
    )
    v1_flags = all_in_one._collect_qa_flags(
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        generated_title="",
        generated_description="",
        optimised_h1="",
        input_h1="",
        primary_keyword="calibration support",
        faq_items=[],
        section_results={"benefits": "## Generic Benefits\nShort copy."},
        forbidden_phrases=[],
        template=template,
        strategy_brief=strategy,
        page_quality_policy_version=PAGE_QUALITY_POLICY_VERSION,
    )

    legacy_codes = {flag["code"] for flag in legacy_flags}
    v1_codes = {flag["code"] for flag in v1_flags}
    assert "page_heading_plan_mismatch" not in legacy_codes
    assert "page_section_below_planned_depth" not in legacy_codes
    assert "page_heading_plan_mismatch" in v1_codes
    assert "page_section_below_planned_depth" in v1_codes
    assert "section_word_count_below_target" not in v1_codes


def test_legacy_collection_internal_heading_requires_editorial_review():
    flags = all_in_one._collect_qa_flags(
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        generated_title="",
        generated_description="",
        optimised_h1="",
        input_h1="",
        primary_keyword="rainbow foil",
        faq_items=[],
        section_results={
            "collection_guidance": (
                "## Helpful Buying Notes\n"
                "Choose pieces that support one coordinated table setting."
            ),
        },
        forbidden_phrases=[],
        template={
            "sections": [{
                "name": "collection_guidance",
                "label": "Helpful Buying Notes",
                "heading_level": "h2",
                "word_count": [60, 120],
            }],
        },
        strategy_brief={},
    )

    matching = [
        flag
        for flag in flags
        if flag.get("code") == "page_heading_generic"
    ]

    assert len(matching) == 1
    assert matching[0]["section"] == "collection_guidance"
    assert matching[0]["actual_heading"] == "Helpful Buying Notes"
    assert matching[0]["severity"] == "review"


@pytest.mark.parametrize(
    ("section_name", "internal_heading"),
    [
        ("collection_story", "Collection Story"),
        ("collection_value", "Collection Value"),
        ("collection_value", "Collection Context"),
    ],
)
def test_new_collection_internal_headings_require_editorial_review(
    section_name,
    internal_heading,
):
    flags = all_in_one._collect_qa_flags(
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        generated_title="",
        generated_description="",
        optimised_h1="",
        input_h1="",
        primary_keyword="personalised knives",
        faq_items=[],
        section_results={
            section_name: (
                f"## {internal_heading}\n"
                "Supported category copy for the intended customer."
            ),
        },
        forbidden_phrases=[],
        template={
            "sections": [{
                "name": section_name,
                "label": internal_heading,
                "heading_level": "h2",
                "word_count": [80, 110],
            }],
        },
        strategy_brief={},
    )

    matching = [
        flag
        for flag in flags
        if flag.get("code") == "page_heading_generic"
    ]

    assert len(matching) == 1
    assert matching[0]["section"] == section_name
    assert matching[0]["actual_heading"] == internal_heading
    assert matching[0]["severity"] == "review"


@pytest.mark.parametrize(
    ("strategy", "generated_copy", "expected_code"),
    [
        (
            {"section_guidance": [{"section": "benefits"}]},
            "## Practical Value for Operations Teams\nUseful copy.",
            "page_planned_heading_missing",
        ),
        (
            {
                "section_guidance": [{
                    "section": "benefits",
                    "planned_heading": "Benefits",
                }],
            },
            "## Benefits\nUseful copy.",
            "page_heading_generic",
        ),
    ],
)
def test_missing_plan_or_generic_visible_heading_requires_editorial_review(
    strategy,
    generated_copy,
    expected_code,
):
    flags = all_in_one._collect_qa_flags(
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        generated_title="",
        generated_description="",
        optimised_h1="",
        input_h1="",
        primary_keyword="calibration support",
        faq_items=[],
        section_results={"benefits": generated_copy},
        forbidden_phrases=[],
        template={"sections": [TEMPLATE_SECTIONS[1]]},
        strategy_brief=strategy,
        page_quality_policy_version=PAGE_QUALITY_POLICY_VERSION,
    )

    matching = [flag for flag in flags if flag["code"] == expected_code]

    assert len(matching) == 1
    assert matching[0]["severity"] == "review"
    legacy_codes = {
        flag["code"]
        for flag in all_in_one._collect_qa_flags(
            gen_meta=False,
            gen_faqs=False,
            gen_page_copy=True,
            generated_title="",
            generated_description="",
            optimised_h1="",
            input_h1="",
            primary_keyword="calibration support",
            faq_items=[],
            section_results={"benefits": generated_copy},
            forbidden_phrases=[],
            template={"sections": [TEMPLATE_SECTIONS[1]]},
            strategy_brief=strategy,
        )
    }
    assert "page_planned_heading_missing" not in legacy_codes
    assert ("page_heading_generic" in legacy_codes) == (
        expected_code == "page_heading_generic"
    )


def test_missing_required_source_name_is_flagged_without_rewriting_copy():
    common = {
        "gen_meta": False,
        "gen_faqs": False,
        "gen_page_copy": True,
        "generated_title": "",
        "generated_description": "",
        "optimised_h1": "",
        "input_h1": "",
        "primary_keyword": "production supplies",
        "faq_items": [],
        "forbidden_phrases": [],
        "template": {"sections": [TEMPLATE_SECTIONS[1]]},
        "strategy_brief": {
            "section_guidance": [{
                "section": "benefits",
                "planned_heading": "Production Paths",
                "required_named_items": ["Fabrics", "Tape", "Rentals"],
            }],
        },
        "page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION,
    }

    missing_flags = all_in_one._collect_qa_flags(
        **common,
        section_results={
            "benefits": "## Production Paths\nFabrics and Rentals are available."
        },
    )
    complete_flags = all_in_one._collect_qa_flags(
        **common,
        section_results={
            "benefits": "## Production Paths\nFabrics, Tape, and Rentals are available."
        },
    )

    matching = [
        flag
        for flag in missing_flags
        if flag["code"] == "page_required_source_item_missing"
    ]
    assert len(matching) == 1
    assert matching[0]["missing_items"] == ["Tape"]
    assert matching[0]["severity"] == "review"
    assert all(
        flag["code"] != "page_required_source_item_missing"
        for flag in complete_flags
    )


def test_new_job_resolves_balanced_and_persists_server_owned_versions(monkeypatch):
    monkeypatch.setenv("AIO_PAGE_COPY_QUALITY_V1_MODE", "on")

    settings, profile = all_in_one._new_job_page_quality_settings({}, "user-1")

    assert profile.id == "balanced"
    assert settings["page_copy_guidance"] == {"id": "balanced", "version": "1"}
    assert settings["page_quality_policy_version"] == PAGE_QUALITY_POLICY_VERSION
    assert settings["adaptive_policy_version"] == ADAPTIVE_POLICY_VERSION
    assert "prompt_instruction" not in repr(settings)


def test_meta_faq_only_job_does_not_stamp_page_copy_versions(monkeypatch):
    monkeypatch.setenv("AIO_PAGE_COPY_QUALITY_V1_MODE", "on")

    settings, profile = all_in_one._new_job_page_quality_settings(
        {
            "gen_page_copy": False,
            "gen_meta": True,
            "gen_faqs": True,
            "page_copy_guidance_profile_id": "conversion",
        },
        "user-1",
        page_copy_requested=False,
    )

    assert profile is None
    assert settings == {
        "gen_page_copy": False,
        "gen_meta": True,
        "gen_faqs": True,
        "owned_page_capture_version": all_in_one.AIO_OWNED_PAGE_CAPTURE_VERSION,
    }
    assert "page_quality_policy_version" not in settings
    assert "adaptive_policy_version" not in settings
    assert "owned_page_mapping_version" not in settings


def test_disabled_or_unknown_submitted_guidance_fails_visibly(monkeypatch):
    monkeypatch.setenv("AIO_PAGE_COPY_QUALITY_V1_MODE", "off")
    with pytest.raises(HTTPException, match="not enabled"):
        all_in_one._new_job_page_quality_settings(
            {"page_copy_guidance_profile_id": "balanced"},
            "user-1",
        )

    monkeypatch.setenv("AIO_PAGE_COPY_QUALITY_V1_MODE", "on")
    with pytest.raises(HTTPException, match="Unknown"):
        all_in_one._new_job_page_quality_settings(
            {"page_copy_guidance_profile_id": "arbitrary raw prompt"},
            "user-1",
        )


def test_emergency_rerun_switch_never_falls_back_to_legacy(monkeypatch):
    monkeypatch.setenv("AIO_PAGE_COPY_QUALITY_V1_RERUNS_ENABLED", "false")

    jobs._enforce_page_quality_rerun_available(
        {},
        page_copy_requested=True,
    )
    jobs._enforce_page_quality_rerun_available(
        {"page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION},
        page_copy_requested=False,
    )
    with pytest.raises(HTTPException, match="will not be rerun with legacy"):
        jobs._enforce_page_quality_rerun_available(
            {"page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION},
            page_copy_requested=True,
        )


def test_stored_quality_versions_never_fall_forward():
    legacy = all_in_one._stored_page_quality_context(
        {},
        page_copy_requested=True,
    )
    assert legacy["enabled"] is False

    versioned = all_in_one._stored_page_quality_context(
        {
            "page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION,
            "adaptive_policy_version": ADAPTIVE_POLICY_VERSION,
            "owned_page_mapping_version": "current-aio-owned-blocks-v1",
            "page_copy_guidance": {"id": "conversion", "version": "1"},
        },
        page_copy_requested=True,
    )
    assert versioned["enabled"] is True
    assert versioned["guidance"].id == "conversion"
    assert versioned["policy"] is get_page_quality_policy(
        PAGE_QUALITY_POLICY_VERSION
    )
    assert versioned["mapping_policy"].version == "current-aio-owned-blocks-v1"

    with pytest.raises(ValueError, match="unavailable"):
        all_in_one._stored_page_quality_context(
            {
                "page_quality_policy_version": "missing-version",
                "adaptive_policy_version": ADAPTIVE_POLICY_VERSION,
                "owned_page_mapping_version": "current-aio-owned-blocks-v1",
                "page_copy_guidance": {"id": "conversion", "version": "1"},
            },
            page_copy_requested=True,
        )


@pytest.mark.parametrize(
    "settings, expected_detail",
    [
        (
            {
                "page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION,
                "owned_page_mapping_version": "current-aio-owned-blocks-v1",
                "page_copy_guidance": {"id": "conversion", "version": "1"},
            },
            "adaptive policy version",
        ),
        (
            {
                "page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION,
                "adaptive_policy_version": ADAPTIVE_POLICY_VERSION,
                "page_copy_guidance": {"id": "conversion", "version": "1"},
            },
            "Owned-page mapping version",
        ),
        (
            {
                "page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION,
                "adaptive_policy_version": ADAPTIVE_POLICY_VERSION,
                "owned_page_mapping_version": "current-aio-owned-blocks-v1",
            },
            "stored page-copy guidance profile",
        ),
    ],
)
def test_reruns_fail_before_scheduling_when_stored_versions_are_unavailable(
    settings,
    expected_detail,
):
    with pytest.raises(HTTPException) as raised:
        jobs._validate_page_quality_rerun_settings(
            settings,
            page_copy_requested=True,
        )

    assert raised.value.status_code == 409
    assert "was not rerun" in raised.value.detail
    assert expected_detail in raised.value.detail

    jobs._validate_page_quality_rerun_settings(
        settings,
        page_copy_requested=False,
    )


def test_versioned_combined_docx_uses_generated_human_headings_once():
    template = {
        "name": "Service",
        "sections": [{
            **TEMPLATE_SECTIONS[1],
            "planned_heading": "How Calibration Support Reduces Rework",
        }],
    }
    payload = all_in_one._build_combined_docx(
        url="https://example.com",
        h1="Calibration Support",
        primary_keyword="calibration support",
        page_type="service",
        template=template,
        generated_title="Calibration Support Services",
        generated_description="Understand calibration support.",
        optimised_h1="Calibration Support",
        faq_items=[],
        faq_schema=None,
        section_results={
            "benefits": (
                "## How Calibration Support Reduces Rework\n"
                "Clear body copy."
            )
        },
        word_count=8,
        competitor_urls=[],
        gen_meta=True,
        gen_faqs=False,
        gen_page_copy=True,
        keyword_assignment={},
        page_quality_policy_version=PAGE_QUALITY_POLICY_VERSION,
    )
    document = Document(io.BytesIO(payload))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]

    assert paragraph_text.count("How Calibration Support Reduces Rework") == 1
    assert "Benefits" not in paragraph_text


def test_legacy_combined_docx_keeps_legacy_h1_and_section_label_format():
    template = {
        "name": "Service",
        "sections": [{
            **TEMPLATE_SECTIONS[1],
            "planned_heading": "How Calibration Support Reduces Rework",
        }],
    }
    payload = all_in_one._build_combined_docx(
        url="https://example.com",
        h1="Stored Calibration H1",
        primary_keyword="calibration support",
        page_type="service",
        template=template,
        generated_title="Calibration Support Services",
        generated_description="Understand calibration support.",
        optimised_h1="Optimised Calibration H1",
        faq_items=[],
        faq_schema=None,
        section_results={
            "benefits": (
                "## How Calibration Support Reduces Rework\n"
                "Clear body copy."
            )
        },
        word_count=8,
        competitor_urls=[],
        gen_meta=True,
        gen_faqs=False,
        gen_page_copy=True,
        keyword_assignment={},
    )
    document = Document(io.BytesIO(payload))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]

    assert paragraph_text[0] == "Stored Calibration H1"
    assert "Benefits" in paragraph_text
    assert "How Calibration Support Reduces Rework" not in paragraph_text


def test_page_copy_only_docx_keeps_generated_human_heading_once():
    template = {
        "name": "Service",
        "sections": [{
            **TEMPLATE_SECTIONS[1],
            "planned_heading": "How Calibration Support Reduces Rework",
            "adaptive_policy_version": ADAPTIVE_POLICY_VERSION,
        }],
    }
    payload = all_in_one._build_combined_docx(
        url="https://example.com",
        h1="Calibration Support",
        primary_keyword="calibration support",
        page_type="service",
        template=template,
        generated_title="",
        generated_description="",
        optimised_h1="",
        faq_items=[],
        faq_schema=None,
        section_results={
            "benefits": (
                "## How Calibration Support Reduces Rework\n"
                "Clear body copy."
            )
        },
        word_count=8,
        competitor_urls=[],
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        keyword_assignment={},
    )
    document = Document(io.BytesIO(payload))
    matching = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "How Calibration Support Reduces Rework"
    ]

    assert len(matching) == 1
    assert matching[0].style.name == "Heading 2"
    assert "Benefits" not in [paragraph.text for paragraph in document.paragraphs]


@pytest.mark.parametrize(
    (
        "heading_level",
        "planned_heading",
        "label",
        "expected_heading",
        "expected_style",
    ),
    [
        (
            "h1",
            "Calibration Support",
            "Hero",
            "Calibration Support",
            "Heading 1",
        ),
        (
            "h2",
            "How Calibration Support Reduces Rework",
            "Benefits",
            "How Calibration Support Reduces Rework",
            "Heading 2",
        ),
        (
            "h3",
            "",
            "Implementation Notes",
            "Implementation Notes",
            "Heading 3",
        ),
    ],
)
def test_page_copy_only_docx_adds_versioned_human_heading_fallback_once(
    heading_level,
    planned_heading,
    label,
    expected_heading,
    expected_style,
):
    template = {
        "name": "Service",
        "sections": [{
            **TEMPLATE_SECTIONS[1],
            "name": "details",
            "label": label,
            "heading_level": heading_level,
            "planned_heading": planned_heading,
            "adaptive_policy_version": ADAPTIVE_POLICY_VERSION,
        }],
    }
    payload = all_in_one._build_combined_docx(
        url="https://example.com",
        h1="Calibration Support",
        primary_keyword="calibration support",
        page_type="service",
        template=template,
        generated_title="",
        generated_description="",
        optimised_h1="",
        faq_items=[],
        faq_schema=None,
        section_results={"details": "Clear body copy without a Markdown heading."},
        word_count=8,
        competitor_urls=[],
        gen_meta=False,
        gen_faqs=False,
        gen_page_copy=True,
        keyword_assignment={},
    )
    document = Document(io.BytesIO(payload))
    matching = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == expected_heading
    ]

    assert len(matching) == 1
    assert matching[0].style.name == expected_style
    if label != expected_heading:
        assert label not in [paragraph.text for paragraph in document.paragraphs]

import logging
import time
import uuid
import base64
import re
from copy import deepcopy
from datetime import datetime, timezone
from typing import Literal
from urllib.parse import urlparse
from fastapi import APIRouter, Depends, BackgroundTasks, HTTPException, Response
from google.auth.exceptions import RefreshError
from pydantic import BaseModel

from auth import get_current_user, get_supabase
from abuse_protection import enforce_job_start, enforce_rate_limit, execute_active_job_write
from credentials import hydrate_job_settings, mark_gsc_reconnect_required, strip_secret_fields
from safe_logging import log_safe_exception, log_safe_external_failure
from utils.niches import get_niche_context
from utils.dfs import (
    get_search_volume, get_keyword_difficulty,
    get_ranked_keywords_for_url, get_keyword_ideas,
    get_serp_data,
)
from utils.keyword import rank_keywords, merge_keyword_pools, assign_keywords_to_sections
from utils.gsc import GscOAuthConfigError, get_gsc_client, get_top_queries_for_url
from utils.scraper import (
    scrape_url, map_competitor_sections,
    classify_competitor_relevance, is_editorial_competitor,
)
from utils.faq_scraper import (
    AIO_OWNED_PAGE_CAPTURE_VERSION,
    is_ecommerce_collection_page,
    scrape_page_context,
)
from utils.templates import TEMPLATES, get_template, get_templates_for_page_type, parse_custom_template
from utils.adaptive_templates import (
    adapt_template_for_generation,
    attach_depth_policies,
)
from utils.owned_page import (
    OWNED_PAGE_MAPPING_VERSION,
    SOURCE_BLOCK_PLAN_VERSION,
    SOURCE_ASSET_MANIFEST_VERSION,
    build_owned_page_registry,
    build_source_asset_manifest,
    get_owned_page_mapping_policy,
    resolve_source_block_plan_version,
)
from utils.page_quality import (
    ADAPTIVE_POLICY_VERSION,
    CLAIM_BOUND_RENDERER_VERSION,
    PAGE_QUALITY_POLICY_VERSION,
    PageQualityConfigurationError,
    get_adaptive_policy,
    get_page_quality_policy,
    guidance_capability_payload,
    page_quality_creation_enabled,
    resolve_claim_bound_renderer_version,
    resolve_stored_guidance_profile,
    select_guidance_profile,
)
from utils.page_types import default_template_key_for_page_type, normalize_page_type
from utils.language import find_non_us_english_spellings
from utils.copy_gen import (
    _claim_bound_canonical_h1,
    _has_non_negated_action_match,
    _is_collection_promotional_text,
    _supported_page_action_types,
    _source_asset_exact_phrases,
    _source_asset_forbidden_conflicts,
    _structured_source_asset_render_plan,
    _validated_source_asset_section_names,
    generate_page, generate_faq, generate_copy, generate_strategy_brief, sanitise,
    normalise_collection_references,
    page_plan_diagnostics, strategy_brief_issues,
    META_TITLE_PREFERRED_MIN, META_TITLE_PREFERRED_MAX,
    META_DESCRIPTION_PREFERRED_MIN, META_DESCRIPTION_PREFERRED_MAX,
)
from utils.docx_export import build_docx

router = APIRouter()
logger = logging.getLogger(__name__)

_GSC_RECONNECT_ERROR = "Google Search Console reconnect required."
_GSC_UNAVAILABLE_ERROR = "Selected Google Search Console connection unavailable."
_GSC_CONFIG_ERROR = "Google Search Console OAuth configuration missing."
_GSC_METHOD_LABELS = {"google_oauth", "service_account", "disabled", "unavailable"}
_OWNED_PAGE_CONTEXT_ERROR = "Owned-page context was unavailable."
_ROW_PROCESSING_ERROR = "This row could not be processed. Please try again."
_STRATEGY_BRIEF_ERROR = "Strategy brief generation was unavailable."
_SOURCE_ASSET_ADAPTIVE_MARKER = "__source_asset_material_present__"
_SOURCE_ASSET_ADAPTIVE_REASONS = frozenset({
    "no_owned_proof",
    "keyword_section_without_owned_proof",
})
_STANDARD_BUILT_IN_GENERATION_TEMPLATE_KEYS = frozenset(TEMPLATES)
_SOURCE_ASSET_ADAPTIVE_INSTRUCTION = (
    "The proof-only omission rule applies to newly authored factual claims, "
    "not to exact assigned source assets. Preserve those assigned assets as "
    "required editorial material, not factual proof, and do not infer any "
    "added claim from them."
)
_CAPTURE_QUALITY_INTEGER_FIELDS = frozenset({
    "empty_blocks_rejected",
    "duplicate_blocks_rejected",
    "filter_count",
    "heading_count",
    "low_signal_blocks_rejected",
    "mapped_block_count",
    "mapped_retained_chars",
    "navigation_links_rejected",
    "primary_retained_chars",
    "product_count",
    "raw_chars",
    "recovery_retained_chars",
    "retained_chars",
    "short_blocks_retained",
    "ui_noise_lines_rejected",
    "visible_link_labels_retained",
})
_CAPTURE_QUALITY_BOOLEAN_FIELDS = frozenset({
    "mapping_truncated",
    "recovery_attempted",
    "recovery_selected",
    "sparse",
})
_CAPTURE_SPARSE_REASONS = frozenset({
    "few_content_blocks",
    "few_headings",
    "no_products_detected",
    "retained_chars_below_target",
})
_SOURCE_ASSET_RERUN_ADAPTIVE_INSTRUCTION = (
    "This section remains available for a reviewer rerun, but its source "
    "assets are intentionally excluded from this rerun prompt. Do not "
    "reconstruct or replace them. They are not factual proof and authorize "
    "no invented replacement claims."
)
_PAGE_CLOSING_CTA_SECTION_NAMES = frozenset({
    "cta",
    "cta_close",
    "closing",
    "final_cta",
})
_PAGE_ACTION_PATTERNS = {
    "booking": re.compile(
        r"\b(?:book|booking|schedule|scheduling|appointments?)\b",
        re.IGNORECASE,
    ),
    "consultation": re.compile(r"\bconsultations?\b", re.IGNORECASE),
    "contact": re.compile(
        r"\b(?:contact|reach\s+out|get\s+in\s+touch|call|email|message|"
        r"speak\s+(?:to|with)|talk\s+(?:to|with))\b",
        re.IGNORECASE,
    ),
    "order": re.compile(r"\b(?:order|purchase|buy|shop)\b", re.IGNORECASE),
    "quote": re.compile(r"\b(?:quotes?|estimates?)\b", re.IGNORECASE),
    "visit": re.compile(r"\b(?:visit|directions?)\b", re.IGNORECASE),
}
_FREE_OFFER_RE = re.compile(
    r"\b(?:free|complimentary|no[-\s]?cost)\s+"
    r"(?:quotes?|estimates?|consultations?|assessments?|trials?)\b",
    re.IGNORECASE,
)
_BROAD_LOCATION_SCOPE_RE = re.compile(
    r"\b(?:(?:coverage\s+across|"
    r"serv(?:e|es|ing)\s+(?:clients?\s+)?throughout|"
    r"clients?\s+throughout|"
    r"available\s+(?:across|throughout))"
    r"[^.!?\n]{0,80}\b(?:area|region)|"
    r"(?:the\s+)?(?:entire|whole)\s+[^.!?\n]{0,50}"
    r"(?:area|region))\b",
    re.IGNORECASE,
)

_RATE_LIMITS = {
    "Claude": 0.5,
    "OpenAI": 0.5,
    "Gemini (free)": 5.0,
    "Mistral (free tier)": 2.0,
    "Groq (free tier)": 2.0,
}


def _new_job_page_quality_settings(
    submitted_settings: dict,
    user_id: object,
    *,
    page_copy_requested: bool = True,
) -> tuple[dict, object | None]:
    """Resolve server-owned versions once, before a new job is persisted."""
    settings = dict(submitted_settings)
    settings["owned_page_capture_version"] = AIO_OWNED_PAGE_CAPTURE_VERSION
    if not page_copy_requested:
        settings.pop("page_copy_guidance_profile_id", None)
        return settings, None

    submitted_profile_id = str(
        settings.get("page_copy_guidance_profile_id") or ""
    ).strip()
    if not page_quality_creation_enabled(user_id):
        if submitted_profile_id:
            raise HTTPException(
                status_code=400,
                detail="Current AIO page-copy quality guidance is not enabled for this account.",
            )
        settings.pop("page_copy_guidance_profile_id", None)
        return settings, None

    try:
        profile = select_guidance_profile(submitted_profile_id)
    except PageQualityConfigurationError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from None
    settings.update({
        "page_copy_guidance_profile_id": profile.id,
        "page_copy_guidance": profile.snapshot(),
        "page_quality_policy_version": PAGE_QUALITY_POLICY_VERSION,
        "adaptive_policy_version": ADAPTIVE_POLICY_VERSION,
        "owned_page_mapping_version": OWNED_PAGE_MAPPING_VERSION,
        "source_asset_manifest_version": SOURCE_ASSET_MANIFEST_VERSION,
        "claim_bound_renderer_version": CLAIM_BOUND_RENDERER_VERSION,
        "source_block_plan_version": SOURCE_BLOCK_PLAN_VERSION,
    })
    return settings, profile


def _stored_page_quality_context(
    settings: dict,
    *,
    page_copy_requested: bool,
) -> dict:
    """Resolve exact stored versions, or preserve the historical legacy path."""
    page_quality_version = str(
        settings.get("page_quality_policy_version") or ""
    ).strip()
    if not page_quality_version:
        dependent_version_fields = (
            "adaptive_policy_version",
            "owned_page_mapping_version",
            "source_asset_manifest_version",
            "claim_bound_renderer_version",
            "source_block_plan_version",
            "page_copy_guidance",
        )
        if page_copy_requested and any(
            settings.get(field)
            for field in dependent_version_fields
        ):
            raise PageQualityConfigurationError(
                "Stored page-copy quality configuration is missing its "
                "page-quality policy version."
            )
        return {
            "enabled": False,
            "guidance": None,
            "policy": None,
            "mapping_policy": None,
            "page_quality_policy_version": "",
            "adaptive_policy_version": "",
            "owned_page_mapping_version": "",
            "source_asset_manifest_version": "",
            "claim_bound_renderer_version": "",
            "source_block_plan_version": "",
        }
    if not page_copy_requested:
        return {
            "enabled": False,
            "guidance": None,
            "policy": None,
            "mapping_policy": None,
            "page_quality_policy_version": page_quality_version,
            "adaptive_policy_version": str(
                settings.get("adaptive_policy_version") or ""
            ).strip(),
            "owned_page_mapping_version": str(
                settings.get("owned_page_mapping_version") or ""
            ).strip(),
            "source_asset_manifest_version": str(
                settings.get("source_asset_manifest_version") or ""
            ).strip(),
            "claim_bound_renderer_version": str(
                settings.get("claim_bound_renderer_version") or ""
            ).strip(),
            "source_block_plan_version": str(
                settings.get("source_block_plan_version") or ""
            ).strip(),
        }

    page_policy = get_page_quality_policy(page_quality_version)
    adaptive_version = str(settings.get("adaptive_policy_version") or "").strip()
    adaptive_policy = get_adaptive_policy(adaptive_version)
    if adaptive_policy.version != page_policy.adaptive_policy_version:
        raise PageQualityConfigurationError(
            "Stored page-quality and adaptive policy versions do not match."
        )
    mapping_version = str(
        settings.get("owned_page_mapping_version") or ""
    ).strip()
    mapping_policy = get_owned_page_mapping_policy(mapping_version)
    source_asset_version = str(
        settings.get("source_asset_manifest_version") or ""
    ).strip()
    if (
        source_asset_version
        and source_asset_version != SOURCE_ASSET_MANIFEST_VERSION
    ):
        raise PageQualityConfigurationError(
            f'Source-asset manifest version "{source_asset_version}" is unavailable.'
        )
    guidance = resolve_stored_guidance_profile(
        settings.get("page_copy_guidance"),
        versioned_job=True,
    )
    claim_bound_renderer_version = resolve_claim_bound_renderer_version(
        settings.get("claim_bound_renderer_version")
    )
    source_block_plan_version = resolve_source_block_plan_version(
        settings.get("source_block_plan_version")
    )
    if bool(claim_bound_renderer_version) != bool(source_block_plan_version):
        raise PageQualityConfigurationError(
            "Stored claim-bound renderer and source-block plan versions do not match."
        )
    return {
        "enabled": True,
        "guidance": guidance,
        "policy": page_policy,
        "mapping_policy": mapping_policy,
        "page_quality_policy_version": page_policy.version,
        "adaptive_policy_version": adaptive_policy.version,
        "owned_page_mapping_version": mapping_version,
        "source_asset_manifest_version": source_asset_version,
        "claim_bound_renderer_version": claim_bound_renderer_version,
        "source_block_plan_version": source_block_plan_version,
    }


def _page_copy_correction_is_active(
    page_quality: dict,
    *,
    requested: bool,
) -> bool:
    """Limit the corrective path to newly stamped, compatible jobs."""
    return bool(
        requested
        and page_quality.get("enabled")
        and page_quality.get("source_asset_manifest_version")
        == SOURCE_ASSET_MANIFEST_VERSION
    )


def _uses_standard_built_in_page_generation(
    template_key: str,
    custom_template_text: str,
) -> bool:
    """Keep built-in page-copy templates AI-authored instead of source-rendered."""
    return bool(
        not str(custom_template_text or "").strip()
        and str(template_key or "").strip()
        in _STANDARD_BUILT_IN_GENERATION_TEMPLATE_KEYS
    )


def _claim_bound_rendering_is_active(
    page_quality: dict,
    *,
    requested: bool,
    template_key: str,
    custom_template_text: str,
) -> bool:
    return bool(
        requested
        and page_quality.get("enabled")
        and page_quality.get("claim_bound_renderer_version")
        == CLAIM_BOUND_RENDERER_VERSION
        and page_quality.get("source_block_plan_version")
        == SOURCE_BLOCK_PLAN_VERSION
        and not _uses_standard_built_in_page_generation(
            template_key,
            custom_template_text,
        )
    )


def _adapt_page_template_for_generation(
    template: dict,
    template_key: str,
    strategy_brief: dict | None,
    *,
    adaptive_policy_version: str = "",
    source_asset_manifest_version: str = "",
    include_source_asset_instruction: bool = True,
    correction_evidence_contract: bool = False,
) -> tuple[dict, list[dict]]:
    """Keep validated editorial assets without changing adaptive evidence rules."""
    baseline_template, baseline_plan = adapt_template_for_generation(
        template,
        template_key,
        strategy_brief,
        adaptive_policy_version=adaptive_policy_version,
        correction_evidence_contract=correction_evidence_contract,
    )
    if source_asset_manifest_version != SOURCE_ASSET_MANIFEST_VERSION:
        return baseline_template, baseline_plan
    source_section_names = _validated_source_asset_section_names(
        strategy_brief
    )
    if not source_section_names:
        return baseline_template, baseline_plan

    baseline_plan_by_section = {
        str(item.get("section") or "").strip().casefold(): item
        for item in baseline_plan
        if isinstance(item, dict)
    }
    bridged_section_names = {
        section_name
        for section_name in source_section_names
        if (
            baseline_plan_by_section.get(section_name, {}).get("mode")
            in {"omit", "compact"}
            and baseline_plan_by_section.get(section_name, {}).get("reason")
            in _SOURCE_ASSET_ADAPTIVE_REASONS
        )
    }
    if not bridged_section_names:
        return baseline_template, baseline_plan

    planning_brief = deepcopy(strategy_brief or {})
    for contract in planning_brief.get("section_guidance") or []:
        if (
            isinstance(contract, dict)
            and str(contract.get("section") or "").strip().casefold()
            in bridged_section_names
        ):
            contract["proof_points"] = [_SOURCE_ASSET_ADAPTIVE_MARKER]

    adapted_template, adapted_plan = adapt_template_for_generation(
        template,
        template_key,
        planning_brief,
        adaptive_policy_version=adaptive_policy_version,
        correction_evidence_contract=correction_evidence_contract,
    )
    for section in adapted_template.get("sections") or []:
        section_name = str(section.get("name") or "").strip().casefold()
        if section_name not in bridged_section_names:
            continue
        instruction = str(section.get("adaptive_instruction") or "").strip()
        source_asset_instruction = (
            _SOURCE_ASSET_ADAPTIVE_INSTRUCTION
            if include_source_asset_instruction
            else _SOURCE_ASSET_RERUN_ADAPTIVE_INSTRUCTION
        )
        section["adaptive_instruction"] = (
            f"{instruction} {source_asset_instruction}".strip()
        )
    for item in adapted_plan:
        section_name = str(item.get("section") or "").strip().casefold()
        if section_name not in bridged_section_names:
            continue
        item["proof_point_count"] = baseline_plan_by_section[
            section_name
        ].get("proof_point_count", 0)
        item["reason"] = "source_asset_material"
    return adapted_template, adapted_plan


def _scrape_owned_page_for_settings(
    settings: dict,
    url: str,
    scraper_override: str = "",
    business_type: str = "general",
    page_type: str = "general",
) -> dict:
    firecrawl_key = settings.get("firecrawl_api_key", "")
    mode = "ecommerce_collection" if is_ecommerce_collection_page(business_type, page_type) else "default"
    requested_provider = (
        "firecrawl"
        if scraper_override == "firecrawl" or settings.get("scrape_provider", "jina") == "firecrawl"
        else "jina"
    )
    capture_version = str(settings.get("owned_page_capture_version") or "").strip()
    capture_options = (
        {"capture_version": capture_version}
        if capture_version
        else {}
    )

    def annotate(result: dict, fallback_used: bool = False) -> dict:
        result["mode"] = result.get("mode") or mode
        result["requested_provider"] = requested_provider
        result["fallback_used"] = fallback_used or result.get("source") in {
            "cached_fallback",
            "live_selector_recovery",
        }
        result["raw_chars"] = int(result.get("raw_chars") or 0)
        result["cleaned_chars"] = int(result.get("cleaned_chars") or len(result.get("content") or ""))
        return result

    if requested_provider == "firecrawl":
        from utils.faq_scraper import scrape_page_context_firecrawl
        return annotate(
            scrape_page_context_firecrawl(
                firecrawl_key,
                url,
                mode=mode,
                **capture_options,
            )
        )

    jina_result = annotate(
        scrape_page_context(
            settings.get("jina_api_key", ""),
            url,
            mode=mode,
            **capture_options,
        )
    )
    if jina_result.get("success") or not settings.get("firecrawl_fallback"):
        return jina_result
    if not firecrawl_key:
        return jina_result

    from utils.faq_scraper import scrape_page_context_firecrawl
    firecrawl_result = annotate(
        scrape_page_context_firecrawl(
            firecrawl_key,
            url,
            mode=mode,
            **capture_options,
        ),
        fallback_used=True,
    )
    if not firecrawl_result.get("success"):
        firecrawl_result["error"] = f"Jina failed; {firecrawl_result.get('error') or 'Firecrawl could not scrape this page.'}"
    return firecrawl_result


def _safe_capture_quality_diagnostics(value: object) -> dict:
    if not isinstance(value, dict):
        return {}
    safe = {}
    for field in _CAPTURE_QUALITY_INTEGER_FIELDS:
        raw_value = value.get(field)
        if isinstance(raw_value, int) and not isinstance(raw_value, bool):
            safe[field] = max(0, min(raw_value, 100_000_000))
    for field in _CAPTURE_QUALITY_BOOLEAN_FIELDS:
        if isinstance(value.get(field), bool):
            safe[field] = value[field]
    retention_ratio = value.get("retention_ratio")
    if isinstance(retention_ratio, (int, float)) and not isinstance(retention_ratio, bool):
        safe["retention_ratio"] = max(0.0, min(float(retention_ratio), 1.0))
    reasons = value.get("sparse_reasons")
    if isinstance(reasons, list):
        safe["sparse_reasons"] = [
            reason
            for reason in reasons[:8]
            if reason in _CAPTURE_SPARSE_REASONS
        ]
    return safe


def _owned_page_scraper_available(settings: dict, scraper_override: str = "") -> bool:
    if scraper_override == "firecrawl" or settings.get("scrape_provider", "jina") == "firecrawl":
        return bool(settings.get("firecrawl_api_key"))
    return True

_GENERIC_OPENERS = (
    "Welcome to",
    "Are you looking for",
    "In today's world",
    "Whether you are",
    "Finding the right",
    "When it comes to",
    "Choosing the right",
    "Looking for",
    "There are many",
    "It can be difficult to",
    "If you are searching for",
    "Whether you need",
    "In the world of",
)

_KEYWORD_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "by", "for", "from", "in", "into",
    "near", "of", "on", "or", "the", "to", "with", "your",
}

_CONTENT_GAP_STOPWORDS = _KEYWORD_STOPWORDS | {
    "about", "also", "buyers", "client", "competitor", "competitors", "content",
    "cover", "covers", "daily", "does", "example", "explain", "explains",
    "from", "have", "into", "more", "page", "section", "their", "this",
    "what", "when", "where", "which", "while",
}

_REPEATED_PHRASE_STOPWORDS = _CONTENT_GAP_STOPWORDS | {
    "after", "all", "because", "before", "can", "each", "every", "has",
    "help", "helps", "how", "its", "like", "make", "makes", "need", "needs",
    "our", "over", "than", "that", "these", "they", "through", "use", "used",
    "using", "was", "way", "will", "you",
}

_QA_SEVERITIES = {"warning", "review", "error"}
_GENERIC_PAGE_HEADINGS = frozenset({
    "about",
    "benefits",
    "buying guide",
    "collection context",
    "collection guidance",
    "collection story",
    "collection value",
    "conclusion",
    "features",
    "helpful buying notes",
    "introduction",
    "our process",
    "our services",
    "overview",
    "services",
    "why choose us",
})
_COLLECTION_TEMPLATE_SECTION_NAMES = frozenset({
    "category_intro",
    "collection_story",
    "collection_value",
    "collection_guidance",
})
_COLLECTION_NON_PROMOTIONAL_SECTION_NAMES = (
    "category_intro",
    "collection_story",
    "collection_value",
)
_B2B_CONSUMER_CTAS = (
    "shop now",
    "add to cart",
    "grab yours",
    "buy today",
    "buy now",
    "order yours",
)
_META_NEXT_ACTION_PHRASES = (
    "explore", "browse", "compare", "discover", "find", "learn", "see", "view", "choose",
    "shop", "order", "request", "contact", "book", "schedule", "call", "visit", "talk",
    "get started", "get a quote", "get directions",
)
_FAQ_RISKY_TOPICS = (
    "shipping",
    "delivery",
    "return policy",
    "returns",
    "refund",
    "warranty",
    "guarantee",
    "in stock",
    "availability",
    "pricing",
)


def _safe_gsc_auth_method(settings: dict, gsc_credentials: dict | None, gsc_client=None) -> str:
    if not settings.get("use_gsc"):
        return "disabled"
    if not gsc_credentials or not gsc_client:
        return "unavailable"
    method = gsc_credentials.get("method")
    return method if method in _GSC_METHOD_LABELS else "unavailable"


def _is_cancelled(sb, job_id: str, user_id: str) -> bool:
    try:
        res = (
            sb.table("jobs")
            .select("status")
            .eq("id", job_id)
            .eq("user_id", user_id)
            .execute()
        )
        return bool(res.data and res.data[0].get("status") in {"cancelling", "cancelled"})
    except Exception:
        return False


def _is_missing_internal_link_suggestions_column(exc: Exception) -> bool:
    message = str(exc).lower()
    code = str(getattr(exc, "code", "") or "").upper()
    return "internal_link_suggestions" in message and (
        code in {"42703", "PGRST204", "PGRST205"}
        or "column" in message
        or "schema cache" in message
    )


def _execute_job_update(sb, job_id: str, user_id: str, update_data: dict):
    return (
        sb.table("jobs")
        .update(update_data)
        .eq("id", job_id)
        .eq("user_id", user_id)
        .execute()
    )


def _update_job(sb, job_id: str, user_id: str, data: dict):
    try:
        update_data = {**data, "updated_at": "now()"}
        if "current_step" in data and data["current_step"]:
            log_entry = {
                "ts": datetime.now(timezone.utc).isoformat(),
                "msg": data["current_step"],
            }
            try:
                res = (
                    sb.table("jobs")
                    .select("logs")
                    .eq("id", job_id)
                    .eq("user_id", user_id)
                    .execute()
                )
                current_logs = (res.data[0].get("logs") or []) if res.data else []
                current_logs.append(log_entry)
                update_data["logs"] = current_logs
            except Exception:
                pass
        try:
            _execute_job_update(sb, job_id, user_id, update_data)
        except Exception as exc:
            if (
                "internal_link_suggestions" in update_data
                and _is_missing_internal_link_suggestions_column(exc)
            ):
                fallback_data = {**update_data}
                fallback_data.pop("internal_link_suggestions", None)
                _execute_job_update(sb, job_id, user_id, fallback_data)
            else:
                raise
    except Exception:
        pass


def _build_brand_context(brand_profile: dict | None, niche: str = "") -> str:
    lines = []
    if brand_profile:
        if brand_profile.get("brand_voice"):
            lines.append("- Voice: " + brand_profile["brand_voice"])
        tone = brand_profile.get("tone") or brand_profile.get("tone_of_voice")
        if tone:
            lines.append("- Tone: " + tone)
        if brand_profile.get("target_audience"):
            lines.append("- Target audience: " + brand_profile["target_audience"])
        if brand_profile.get("usps"):
            lines.append("- Unique selling points: " + brand_profile["usps"])
        if brand_profile.get("key_messages"):
            lines.append("- Key messages to reinforce: " + brand_profile["key_messages"])
        if brand_profile.get("competitors"):
            lines.append("- Competitors to differentiate from: " + brand_profile["competitors"])
        if brand_profile.get("products_services"):
            lines.append("- Products/services: " + brand_profile["products_services"])
        if brand_profile.get("words_to_avoid"):
            lines.append("- Words to avoid: " + brand_profile["words_to_avoid"])
        if brand_profile.get("example_copy"):
            lines.append("- Example copy to emulate in style, not content:\n" + brand_profile["example_copy"])
        if brand_profile.get("guidelines"):
            lines.append("- Additional brand guidelines:\n" + brand_profile["guidelines"])

    parts = ["BRAND CONTEXT:\n" + "\n".join(lines)] if lines else []
    niche_context = get_niche_context(niche)
    if niche_context:
        parts.append("NICHE CONTEXT:\n" + niche_context)
    return "\n".join(parts)


def _build_brand_style_context(brand_profile: dict | None) -> str:
    lines = []
    profile = brand_profile or {}
    voice = profile.get("brand_voice") or profile.get("tone_of_voice")
    tone = profile.get("tone")
    if voice:
        lines.append("- Voice: " + str(voice))
    if tone:
        lines.append("- Tone: " + str(tone))
    return "BRAND STYLE:\n" + "\n".join(lines) if lines else ""


def _split_forbidden_phrases(*values) -> list[str]:
    phrases = []
    seen = set()
    for value in values:
        if not value:
            continue
        if isinstance(value, (list, tuple, set)):
            candidates = value
        else:
            candidates = re.split(r"[\n,;]+", str(value))
        for candidate in candidates:
            phrase = str(candidate).strip()
            key = phrase.lower()
            if phrase and key not in seen:
                phrases.append(phrase)
                seen.add(key)
    return phrases


def _contains_forbidden_phrase(text: str, phrase: str) -> bool:
    if not text or not phrase:
        return False
    phrase = phrase.strip()
    escaped = re.escape(phrase)
    if not escaped:
        return False
    left = r"(?<!\w)" if phrase[0].isalnum() else ""
    right = r"(?!\w)" if phrase[-1].isalnum() else ""
    return re.search(left + escaped + right, text, flags=re.IGNORECASE) is not None


def _normalise_phrase(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _add_qa_flag(
    flags: list[dict],
    code: str,
    message: str,
    output: str = "",
    phrase: str = "",
    severity: str = "review",
):
    flag = {
        "code": code,
        "message": message,
        "severity": severity if severity in _QA_SEVERITIES else "review",
    }
    if output:
        flag["output"] = output
    if phrase:
        flag["phrase"] = phrase
    flags.append(flag)


def _add_us_english_qa_flags(
    flags: list[dict],
    authored_outputs: list[tuple[str, str]],
    protected_phrases: list[str],
):
    for output, text in authored_outputs:
        matches = find_non_us_english_spellings(text, protected_phrases)
        if not matches:
            continue
        _add_qa_flag(
            flags,
            "non_us_english_spelling",
            f"Non-U.S. English spelling found in {output}.",
            output,
        )
        flags[-1]["details"] = matches[:5]


def _qa_status(flags: list[dict]) -> str:
    severities = {str(flag.get("severity") or "review") for flag in (flags or [])}
    if "error" in severities:
        return "error"
    if "review" in severities:
        return "review"
    if "warning" in severities:
        return "warning"
    return "ok"


def _result_failed(result: dict) -> bool:
    status = str((result or {}).get("status") or "")
    return bool((result or {}).get("error")) or status == "error" or status.startswith("skipped:")


def _add_strategy_qa_flag(
    flags: list[dict],
    strategy_status: str,
    strategy_issues: list[str] | None = None,
):
    if strategy_status == "ready" or strategy_status == "not_requested":
        return
    issues = [str(issue).strip() for issue in (strategy_issues or []) if str(issue).strip()]
    message = (
        "Strategy brief was unavailable; outputs were generated without the shared strategy layer."
        if strategy_status == "unavailable"
        else "Strategy brief is incomplete and needs review."
    )
    _add_qa_flag(flags, "strategy_brief_" + strategy_status, message, "strategy", severity="review")
    if issues:
        flags[-1]["details"] = issues[:6]


def _normalise_similarity_text(text: str) -> list[str]:
    return re.findall(r"[a-z0-9]+", (text or "").lower())


def _shingle_similarity(left: str, right: str, shingle_size: int = 3) -> float:
    left_tokens = _normalise_similarity_text(left)
    right_tokens = _normalise_similarity_text(right)
    if len(left_tokens) < 4 or len(right_tokens) < 4:
        return 0.0

    size = shingle_size if len(left_tokens) >= shingle_size and len(right_tokens) >= shingle_size else 1
    left_units = {" ".join(left_tokens[i:i + size]) for i in range(len(left_tokens) - size + 1)}
    right_units = {" ".join(right_tokens[i:i + size]) for i in range(len(right_tokens) - size + 1)}
    if not left_units or not right_units:
        return 0.0

    shingle_score = len(left_units & right_units) / len(left_units | right_units)
    token_left = set(left_tokens)
    token_right = set(right_tokens)
    token_score = len(token_left & token_right) / len(token_left | token_right) if token_left and token_right else 0.0
    return max(shingle_score, token_score)


def _first_page_copy_section(section_results: dict) -> str:
    for key, value in (section_results or {}).items():
        if not str(key).startswith("_") and str(value or "").strip():
            return str(value)
    return ""


def _full_page_copy_text(section_results: dict) -> str:
    return "\n\n".join(str(v) for k, v in (section_results or {}).items() if not str(k).startswith("_"))


def _first_sentence_for_qa(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return ""
    return re.split(r"[.!?]\s+", text, maxsplit=1)[0]


def _strip_leading_markdown_headings(text: str) -> str:
    lines = str(text or "").splitlines()
    while lines and (not lines[0].strip() or re.match(r"^#{1,6}\s+", lines[0].strip())):
        lines.pop(0)
    return "\n".join(lines).strip()


def _find_generic_opener(text: str) -> str:
    first_sentence = _normalise_phrase(_first_sentence_for_qa(text))
    if not first_sentence:
        return ""
    for opener in _GENERIC_OPENERS:
        if first_sentence.startswith(_normalise_phrase(opener)):
            return opener
    return ""


def _add_generic_opener_flags(
    flags: list[dict],
    generated_description: str,
    section_results: dict,
    strategy_brief: dict | None = None,
):
    opener = _find_generic_opener(generated_description)
    if opener:
        flags.append({
            "code": "generic_opener",
            "message": f'Generic opener found: "{opener}".',
            "output": "meta_description",
            "phrase": opener,
        })

    for section_name, text in (section_results or {}).items():
        if str(section_name).startswith("_"):
            continue
        authored_text = _without_exact_source_asset_phrases(
            str(text or ""),
            strategy_brief,
            section_name=str(section_name),
        )
        opener = _find_generic_opener(
            _strip_leading_markdown_headings(authored_text)
        )
        if opener:
            flags.append({
                "code": "generic_opener",
                "message": f'Generic opener found in section "{section_name}": "{opener}".',
                "output": "page_copy",
                "section": section_name,
                "phrase": opener,
            })


def _add_generic_page_reference_flags(
    flags: list[dict],
    section_results: dict,
    strategy_brief: dict | None = None,
):
    pattern = re.compile(
        r"\b(?:this page|this collection|this category|this range|on this page)\b",
        re.IGNORECASE,
    )
    for section_name, text in (section_results or {}).items():
        if str(section_name).startswith("_"):
            continue
        authored_text = _without_exact_source_asset_phrases(
            str(text or ""),
            strategy_brief,
            section_name=str(section_name),
        )
        match = pattern.search(authored_text)
        if match:
            _add_qa_flag(
                flags,
                "generic_page_reference",
                f'Generic page reference found in section "{section_name}": "{match.group(0)}".',
                "page_copy",
                match.group(0),
                severity="warning",
            )


def _extract_first_page_h1(section_results: dict) -> str:
    for text in (section_results or {}).values():
        for line in str(text or "").splitlines():
            match = re.match(r"^#\s+(.+?)\s*$", line.strip())
            if match:
                return match.group(1).strip()
    return ""


def _assemble_full_page_copy(section_results: dict, template: dict | None = None) -> str:
    if template:
        return "\n\n".join(
            str(section_results.get(section.get("name", ""), ""))
            for section in template.get("sections", [])
        )
    return _full_page_copy_text(section_results)


def _enforce_canonical_page_h1(section_results: dict, canonical_h1: str) -> tuple[dict, bool]:
    canonical_h1 = (canonical_h1 or "").strip()
    if not section_results or not canonical_h1:
        return section_results, False

    updated = dict(section_results)
    for section_name, text in (section_results or {}).items():
        if str(section_name).startswith("_"):
            continue

        lines = str(text or "").splitlines()
        for idx, line in enumerate(lines):
            match = re.match(r"^(#\s+)(.+?)\s*$", line.strip())
            if not match:
                continue
            page_h1 = match.group(2).strip()
            if _normalise_phrase(page_h1) == _normalise_phrase(canonical_h1):
                return section_results, False
            lines[idx] = f"# {canonical_h1}"
            updated[section_name] = "\n".join(lines)
            return updated, True

    return section_results, False


def _enforce_v1_canonical_page_h1(
    section_results: dict,
    template: dict | None,
    canonical_h1: str,
) -> tuple[dict, bool]:
    """Make the versioned H1 section match its server-selected heading exactly."""
    canonical_h1 = str(canonical_h1 or "").strip()
    if not section_results or not template or not canonical_h1:
        return section_results, False

    h1_section = next(
        (
            section
            for section in template.get("sections") or []
            if str(section.get("heading_level") or "").casefold() == "h1"
        ),
        None,
    )
    section_name = str((h1_section or {}).get("name") or "")
    text = str(section_results.get(section_name) or "")
    if not section_name or not text.strip():
        return section_results, False

    lines = text.splitlines()
    first_content_index = next(
        (index for index, line in enumerate(lines) if line.strip()),
        None,
    )
    if first_content_index is None:
        return section_results, False

    first_content_line = lines[first_content_index]
    heading_match = re.match(
        r"^\s*(#{1,3})\s+(.+?)\s*$",
        first_content_line,
    )
    if (
        heading_match
        and heading_match.group(1) == "#"
        and heading_match.group(2).strip() == canonical_h1
    ):
        return section_results, False

    if heading_match:
        lines[first_content_index] = f"# {canonical_h1}"
    else:
        lines.insert(first_content_index, f"# {canonical_h1}")

    updated = dict(section_results)
    updated[section_name] = "\n".join(lines)
    return updated, True


def _add_h1_alignment_flag(flags: list[dict], optimised_h1: str, section_results: dict):
    page_h1 = _extract_first_page_h1(section_results)
    if not page_h1 or not (optimised_h1 or "").strip():
        return
    if _normalise_phrase(page_h1) == _normalise_phrase(optimised_h1):
        return
    flags.append({
        "code": "page_h1_differs_from_meta_h1",
        "message": "Page-copy H1 differs from the optimized meta H1.",
        "output": "page_copy",
        "meta_h1": optimised_h1,
        "page_h1": page_h1,
    })


def _word_count_for_qa(text: str) -> int:
    return len(_normalise_similarity_text(text))


def _is_collection_template(template: dict | None) -> bool:
    section_names = {
        str(section.get("name") or "").strip().casefold()
        for section in (template or {}).get("sections") or []
        if isinstance(section, dict)
    }
    return _COLLECTION_TEMPLATE_SECTION_NAMES.issubset(section_names)


def _collection_section_body(text: str) -> str:
    return re.sub(
        r"^[ \t]{0,3}#{1,6}[ \t]+[^\n]*(?:\n|$)",
        "",
        str(text or ""),
        count=1,
    ).lstrip()


def _add_collection_section_role_flags(
    flags: list[dict],
    section_results: dict,
    template: dict | None,
) -> None:
    if not _is_collection_template(template):
        return

    sections_by_name = {
        str(section.get("name") or "").strip().casefold(): section
        for section in (template or {}).get("sections") or []
        if isinstance(section, dict)
    }
    for section_name in _COLLECTION_NON_PROMOTIONAL_SECTION_NAMES:
        text = str((section_results or {}).get(section_name) or "")
        if not text.strip() or not _is_collection_promotional_text(text):
            continue
        section = sections_by_name.get(section_name, {})
        flags.append({
            "code": "collection_promotion_outside_guidance",
            "message": (
                f'Section "{section.get("label", section_name)}" includes '
                "promotional or store-policy language reserved for the final "
                "collection guidance section."
            ),
            "output": "page_copy",
            "section": section_name,
            "severity": "review",
        })


def _limit_faq_items(faq_items: list, requested_count: int) -> tuple[list, bool]:
    if not isinstance(faq_items, list):
        return [], False
    try:
        requested = int(requested_count)
    except (TypeError, ValueError):
        return faq_items, False
    if requested < 0:
        return faq_items, False
    if len(faq_items) > requested:
        return faq_items[:requested], True
    return faq_items, False


def _contains_ngram(container: tuple[str, ...], candidate: tuple[str, ...]) -> bool:
    if len(candidate) > len(container):
        return False
    return any(
        container[idx:idx + len(candidate)] == candidate
        for idx in range(len(container) - len(candidate) + 1)
    )


def _repeated_phrase_candidates(text: str, min_count: int = 3, max_flags: int = 3) -> list[dict]:
    tokens = _normalise_similarity_text(text)
    if len(tokens) < 8:
        return []

    counts = {}
    for size in range(4, 1, -1):
        for idx in range(len(tokens) - size + 1):
            phrase = tuple(tokens[idx:idx + size])
            if phrase[0] in _REPEATED_PHRASE_STOPWORDS or phrase[-1] in _REPEATED_PHRASE_STOPWORDS:
                continue
            meaningful = [token for token in phrase if token not in _REPEATED_PHRASE_STOPWORDS]
            if len(set(meaningful)) < 2:
                continue
            counts[phrase] = counts.get(phrase, 0) + 1

    repeated = [
        (phrase, count)
        for phrase, count in counts.items()
        if count >= min_count
    ]
    repeated.sort(key=lambda item: (-len(item[0]), -item[1], " ".join(item[0])))

    selected = []
    selected_phrases = []
    for phrase, count in repeated:
        if any(_contains_ngram(existing, phrase) for existing in selected_phrases):
            continue
        selected_phrases.append(phrase)
        selected.append({"phrase": " ".join(phrase), "count": count})
        if len(selected) >= max_flags:
            break
    return selected


def _add_repeated_phrase_flags(flags: list[dict], page_copy_text: str):
    for repeated in _repeated_phrase_candidates(page_copy_text):
        flags.append({
            "code": "repeated_phrase",
            "message": "Phrase is repeated too often in page copy.",
            "output": "page_copy",
            "phrase": repeated["phrase"],
            "count": repeated["count"],
        })


def _meaningful_keyword_tokens(keyword: str) -> list[str]:
    return [
        token for token in _normalise_similarity_text(keyword)
        if len(token) > 2 and token not in _KEYWORD_STOPWORDS
    ]


def _keyword_present(keyword: str, text: str) -> bool:
    keyword_tokens = _normalise_similarity_text(keyword)
    text_tokens = _normalise_similarity_text(text)
    if not keyword_tokens or not text_tokens:
        return False

    keyword_phrase = " ".join(keyword_tokens)
    text_phrase = " ".join(text_tokens)
    if f" {keyword_phrase} " in f" {text_phrase} ":
        return True

    meaningful_tokens = _meaningful_keyword_tokens(keyword)
    if not meaningful_tokens:
        return False
    if len(meaningful_tokens) == 1:
        return meaningful_tokens[0] in set(text_tokens)

    matched = sum(1 for token in meaningful_tokens if token in set(text_tokens))
    required = len(meaningful_tokens) if len(meaningful_tokens) <= 3 else int(len(meaningful_tokens) * 0.75 + 0.999)
    return matched >= required


def _keyword_token_root(token: str) -> str:
    token = str(token or "").lower()
    if len(token) > 4 and token.endswith("ies"):
        return token[:-3] + "y"
    if len(token) > 4 and token.endswith("s") and not token.endswith("ss"):
        return token[:-1]
    return token


def _qa_keyword_present(keyword: str, text: str) -> bool:
    if _keyword_present(keyword, text):
        return True
    meaningful_tokens = _meaningful_keyword_tokens(keyword)
    text_tokens = _normalise_similarity_text(text)
    if not meaningful_tokens or not text_tokens:
        return False
    text_roots = {_keyword_token_root(token) for token in text_tokens}
    matched = sum(1 for token in meaningful_tokens if _keyword_token_root(token) in text_roots)
    required = len(meaningful_tokens) if len(meaningful_tokens) <= 3 else int(len(meaningful_tokens) * 0.75 + 0.999)
    return matched >= required


def _generated_output_texts(
    generated_title: str,
    generated_description: str,
    optimised_h1: str,
    faq_items: list,
    section_results: dict,
) -> list[tuple[str, str]]:
    outputs = [
        ("meta_title", generated_title or ""),
        ("meta_description", generated_description or ""),
        ("meta_h1", optimised_h1 or ""),
        ("page_copy", _full_page_copy_text(section_results)),
    ]
    for index, item in enumerate(faq_items or []):
        if isinstance(item, dict):
            outputs.append((f"faq_{index + 1}", f"{item.get('question', '')}\n{item.get('answer', '')}"))
    return outputs


def _add_exclamation_flags(flags: list[dict], outputs: list[tuple[str, str]]):
    for output, text in outputs:
        if "!" in text:
            _add_qa_flag(
                flags,
                "exclamation_mark_present",
                f"Exclamation mark found in {output.replace('_', ' ')}.",
                output,
                "!",
                severity="review",
            )


def _without_exact_source_asset_phrases(
    text: str,
    strategy_brief: dict | None,
    *,
    section_name: str = "",
) -> str:
    value = str(text or "")
    validated_sections = _validated_source_asset_section_names(
        strategy_brief
    )
    requested_section = str(section_name or "").strip().casefold()
    section_names = (
        [requested_section]
        if requested_section in validated_sections
        else ([] if requested_section else sorted(validated_sections))
    )
    for source_section_name in section_names:
        for phrase in sorted(
            _source_asset_exact_phrases(
                strategy_brief,
                source_section_name,
            ),
            key=len,
            reverse=True,
        ):
            value = value.replace(phrase, "", 1)
    return value


def _page_copy_without_materialized_source_units(
    section_results: dict,
    strategy_brief: dict | None,
    forbidden_phrases: list[str] | None = None,
) -> str:
    """Return authored page text for repetition review, excluding exact inserts."""
    sections = []
    for section_name, text in (section_results or {}).items():
        if str(section_name).startswith("_"):
            continue
        authored_text = str(text or "")
        for item in _structured_source_asset_render_plan(
            strategy_brief,
            str(section_name),
            forbidden_phrases or [],
        ):
            rendered = str(item.get("rendered") or "")
            if rendered:
                authored_text = authored_text.replace(rendered, "", 1)
        sections.append(authored_text)
    return "\n\n".join(sections)


def _structured_source_duplicate_findings(
    section_results: dict,
    strategy_brief: dict | None,
    forbidden_phrases: list[str] | None = None,
) -> list[dict]:
    """Find exact structured-source phrases repeated outside canonical units."""
    findings = []
    for section_name, text in (section_results or {}).items():
        if str(section_name).startswith("_"):
            continue
        remainder = str(text or "")
        materialized_items = []
        for item in _structured_source_asset_render_plan(
            strategy_brief,
            str(section_name),
            forbidden_phrases or [],
        ):
            rendered = str(item.get("rendered") or "")
            if not rendered or rendered not in remainder:
                continue
            remainder = remainder.replace(rendered, "", 1)
            materialized_items.append(item)

        duplicate_phrases = []
        asset_ids = []
        for item in materialized_items:
            phrases = (
                item.get("items") or []
                if item.get("kind") == "named_list"
                else [item.get("quote"), item.get("attribution")]
            )
            item_has_duplicate = False
            for phrase in phrases:
                exact_phrase = str(phrase or "").strip()
                is_named_list_item = item.get("kind") == "named_list"
                is_single_word_list_item = bool(
                    is_named_list_item
                    and len(re.findall(r"[^\W_]+", exact_phrase, re.UNICODE)) == 1
                )
                duplicate_present = bool(
                    exact_phrase
                    and (
                        re.search(
                            (
                                r"(?m)^[ \t]*(?:[-+*]|\d+[.)])[ \t]+"
                                + re.escape(exact_phrase)
                                + r"[ \t]*$"
                            ),
                            remainder,
                        )
                        if is_single_word_list_item
                        else exact_phrase in remainder
                    )
                )
                if duplicate_present:
                    duplicate_phrases.append(exact_phrase)
                    item_has_duplicate = True
            if item_has_duplicate:
                asset_ids.append(str(item.get("asset_id") or ""))

        if duplicate_phrases:
            findings.append({
                "section": str(section_name),
                "asset_ids": list(dict.fromkeys(
                    asset_id
                    for asset_id in asset_ids
                    if asset_id
                )),
                "duplicate_phrases": list(dict.fromkeys(
                    duplicate_phrases
                )),
            })
    return findings


def _section_authored_evidence_texts(contract: dict | None) -> list[str]:
    values = []
    section_contract = contract if isinstance(contract, dict) else {}
    for item in section_contract.get("proof_facts") or []:
        if not isinstance(item, dict):
            continue
        text = str(
            item.get("source_excerpt") or item.get("fact") or ""
        ).strip()
        if text:
            values.append(text)
    for asset in section_contract.get("source_assets") or []:
        if not isinstance(asset, dict) or asset.get("kind") != "direct_statement":
            continue
        text = str(asset.get("statement") or "").strip()
        if text:
            values.append(text)
    return values


def _page_action_types(
    text: str,
    *,
    support: bool = False,
    brand_name: str = "",
) -> set[str]:
    value = str(text or "")
    if support:
        return _supported_page_action_types(
            value,
            brand_name=brand_name,
        )
    action_types = _supported_page_action_types(
        value,
        brand_name=brand_name,
    )
    action_types.update({
        action_type
        for action_type, pattern in _PAGE_ACTION_PATTERNS.items()
        if _has_non_negated_action_match(pattern, value)
    })
    return action_types


def _add_page_copy_evidence_backstop_flags(
    flags: list[dict],
    section_results: dict,
    template: dict | None,
    strategy_brief: dict | None,
    forbidden_phrases: list[str] | None,
    *,
    page_type: str,
    brand_name: str = "",
):
    """Flag narrow correction-path leaks that deterministic evidence can prove."""
    contracts = _page_plan_contracts(strategy_brief)
    sections_by_name = {
        str(section.get("name") or "").strip().casefold(): section
        for section in (template or {}).get("sections") or []
        if isinstance(section, dict)
    }
    diagnostics = (
        (strategy_brief or {}).get("source_asset_mapping_diagnostics")
        if isinstance(strategy_brief, dict)
        else {}
    )
    known_asset_ids = {
        str(value)
        for value in (
            diagnostics.get("assigned_asset_ids") or []
            if isinstance(diagnostics, dict)
            and diagnostics.get("active") is True
            else []
        )
        if re.fullmatch(r"A[1-9]\d*", str(value))
    }

    for raw_section_name, raw_text in (section_results or {}).items():
        section_name = str(raw_section_name or "").strip().casefold()
        if not section_name or section_name.startswith("_"):
            continue
        text = str(raw_text or "")
        contract = contracts.get(section_name, {})
        section = sections_by_name.get(section_name, {})
        evidence_sparse = section.get("evidence_sparse") is True
        authored_text = _page_copy_without_materialized_source_units(
            {section_name: text},
            strategy_brief,
            forbidden_phrases,
        )
        section_asset_ids = sorted(
            {
                str(value)
                for value in contract.get("source_asset_ids") or []
                if str(value) in known_asset_ids
            }
        )
        leaked_asset_ids = [
            asset_id
            for asset_id in section_asset_ids
            if re.search(
                rf"(?im)^[ \t]*(?:#{1,6}[ \t]+)?"
                rf"(?:source[ \t]+)?{re.escape(asset_id)}[ \t]*:",
                authored_text,
            )
        ]
        if leaked_asset_ids:
            flags.append({
                "code": "page_internal_source_asset_label",
                "message": (
                    f'Section "{section.get("label", section_name)}" exposes '
                    "an internal source-asset label."
                ),
                "output": "page_copy",
                "section": section_name,
                "asset_ids": leaked_asset_ids,
                "severity": "review",
            })

        evidence_texts = _section_authored_evidence_texts(contract)
        evidence_text = "\n".join(evidence_texts)
        if section_name in _PAGE_CLOSING_CTA_SECTION_NAMES:
            generated_actions = _page_action_types(
                authored_text,
                brand_name=brand_name,
            )
            supported_actions = _page_action_types(
                evidence_text,
                support=True,
                brand_name=brand_name,
            )
            unsupported_actions = sorted(
                generated_actions - supported_actions
            )
            if unsupported_actions:
                flags.append({
                    "code": "page_unsupported_action_type",
                    "message": (
                        f'Section "{section.get("label", section_name)}" adds '
                        "a next-step action type not supported by its exact "
                        "same-section evidence."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "action_types": unsupported_actions,
                    "severity": "review",
                })

            generated_offers = {
                _normalise_phrase(match.group(0))
                for match in _FREE_OFFER_RE.finditer(authored_text)
            }
            supported_evidence = _normalise_phrase(evidence_text)
            unsupported_offers = sorted(
                offer
                for offer in generated_offers
                if not _contains_forbidden_phrase(
                    supported_evidence,
                    offer,
                )
            )
            if unsupported_offers:
                flags.append({
                    "code": "page_unsupported_offer_qualifier",
                    "message": (
                        f'Section "{section.get("label", section_name)}" adds '
                        "an unsupported free or no-cost offer."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "offers": unsupported_offers,
                    "severity": "review",
                })

        if (
            evidence_sparse
            and str(page_type or "").strip().casefold() == "local"
            and _BROAD_LOCATION_SCOPE_RE.search(authored_text)
            and not _BROAD_LOCATION_SCOPE_RE.search(evidence_text)
        ):
            flags.append({
                "code": "page_unsupported_location_scope",
                "message": (
                    f'Section "{section.get("label", section_name)}" broadens '
                    "named location evidence into unsupported area-wide "
                    "coverage."
                ),
                "output": "page_copy",
                "section": section_name,
                "severity": "review",
            })


def _add_b2b_consumer_cta_flags(
    flags: list[dict],
    business_type: str,
    outputs: list[tuple[str, str]],
):
    if str(business_type or "").casefold() != "b2b":
        return
    for output, text in outputs:
        for phrase in _B2B_CONSUMER_CTAS:
            if _contains_forbidden_phrase(text, phrase):
                _add_qa_flag(
                    flags,
                    "b2b_consumer_cta",
                    f'Consumer CTA "{phrase}" found in B2B {output.replace("_", " ")}.',
                    output,
                    phrase,
                    severity="review",
                )


def _add_brand_in_h1_flags(
    flags: list[dict],
    brand_name: str,
    gen_meta: bool,
    gen_page_copy: bool,
    optimised_h1: str,
    section_results: dict,
):
    brand = str(brand_name or "").strip()
    if not brand:
        return
    candidates = []
    if gen_meta and str(optimised_h1 or "").strip():
        candidates.append(("meta_h1", str(optimised_h1)))
    page_h1 = _extract_first_page_h1(section_results) if gen_page_copy else ""
    if page_h1:
        candidates.append(("page_h1", page_h1))
    for output, h1 in candidates:
        if _contains_forbidden_phrase(h1, brand):
            _add_qa_flag(
                flags,
                "brand_name_in_h1",
                "H1 contains the brand name, which conflicts with the AIO headline rule.",
                output,
                brand,
                severity="review",
            )


def _add_meta_length_flags(
    flags: list[dict],
    gen_meta: bool,
    generated_title: str,
    generated_description: str,
):
    if not gen_meta:
        return
    title = str(generated_title or "").strip()
    description = str(generated_description or "").strip()
    if title and not META_TITLE_PREFERRED_MIN <= len(title) <= META_TITLE_PREFERRED_MAX:
        severity = "review" if len(title) > 90 else "warning"
        _add_qa_flag(
            flags,
            "meta_title_outside_preferred_range",
            f"Meta title is {len(title)} characters; the preferred range is "
            f"{META_TITLE_PREFERRED_MIN} to {META_TITLE_PREFERRED_MAX}.",
            "meta_title",
            severity=severity,
        )
        flags[-1].update({
            "actual_length": len(title),
            "preferred_min": META_TITLE_PREFERRED_MIN,
            "preferred_max": META_TITLE_PREFERRED_MAX,
        })
    if description and not META_DESCRIPTION_PREFERRED_MIN <= len(description) <= META_DESCRIPTION_PREFERRED_MAX:
        severity = "review" if len(description) > 200 else "warning"
        _add_qa_flag(
            flags,
            "meta_description_outside_preferred_range",
            f"Meta description is {len(description)} characters; the preferred range is "
            f"{META_DESCRIPTION_PREFERRED_MIN} to {META_DESCRIPTION_PREFERRED_MAX}.",
            "meta_description",
            severity=severity,
        )
        flags[-1].update({
            "actual_length": len(description),
            "preferred_min": META_DESCRIPTION_PREFERRED_MIN,
            "preferred_max": META_DESCRIPTION_PREFERRED_MAX,
        })


def _add_keyword_placement_flags(
    flags: list[dict],
    primary_keyword: str,
    gen_meta: bool,
    gen_page_copy: bool,
    optimised_h1: str,
    section_results: dict,
):
    keyword = str(primary_keyword or "").strip()
    if not keyword:
        return
    if gen_meta and str(optimised_h1 or "").strip() and not _qa_keyword_present(keyword, optimised_h1):
        _add_qa_flag(
            flags,
            "target_keyword_missing_from_h1",
            "Target keyword or a close grammatical variant was not found in the optimized H1.",
            "meta_h1",
            severity="warning",
        )
        flags[-1]["keyword"] = keyword

    if not gen_page_copy:
        return
    page_copy = _full_page_copy_text(section_results)
    body_lines = [
        line
        for line in page_copy.splitlines()
        if not re.match(r"^\s*#{1,6}\s+", line)
    ]
    first_words = _normalise_similarity_text("\n".join(body_lines))[:100]
    if first_words and not _qa_keyword_present(keyword, " ".join(first_words)):
        _add_qa_flag(
            flags,
            "target_keyword_missing_from_first_100_words",
            "Target keyword or a close grammatical variant was not found in the first 100 body words.",
            "page_copy",
            severity="warning",
        )
        flags[-1]["keyword"] = keyword

    h2_labels = []
    for line in page_copy.splitlines():
        match = re.match(r"^\s*##\s+(.+?)\s*$", line)
        if match:
            h2_labels.append(match.group(1))
    if h2_labels and not any(_qa_keyword_present(keyword, label) for label in h2_labels):
        _add_qa_flag(
            flags,
            "target_keyword_missing_from_h2",
            "Target keyword or a close grammatical variant was not found in any H2 heading.",
            "page_copy",
            severity="warning",
        )
        flags[-1]["keyword"] = keyword


def _meta_next_action_expected(business_type: str, page_type: str) -> bool:
    business = str(business_type or "").casefold()
    page = str(page_type or "").casefold()
    return business in {"ecommerce", "service", "local"} or any(
        term in page for term in ("service", "product", "collection", "category", "location", "landing")
    )


def _add_meta_field_quality_flags(
    flags: list[dict],
    *,
    gen_meta: bool,
    primary_keyword: str,
    generated_title: str,
    generated_description: str,
    business_type: str,
    page_type: str,
):
    if not gen_meta:
        return
    keyword = str(primary_keyword or "").strip()
    title = str(generated_title or "").strip()
    description = str(generated_description or "").strip()

    if keyword and title and not _qa_keyword_present(keyword, title):
        _add_qa_flag(
            flags,
            "target_keyword_missing_from_meta_title",
            "Target keyword or a close grammatical variant was not found in the meta title.",
            "meta_title",
            severity="warning",
        )
        flags[-1]["keyword"] = keyword
    if keyword and description and not _qa_keyword_present(keyword, description):
        _add_qa_flag(
            flags,
            "target_keyword_missing_from_meta_description",
            "Target keyword or a close grammatical variant was not found in the meta description.",
            "meta_description",
            severity="warning",
        )
        flags[-1]["keyword"] = keyword
    if (
        description
        and _meta_next_action_expected(business_type, page_type)
        and not any(_contains_forbidden_phrase(description, phrase) for phrase in _META_NEXT_ACTION_PHRASES)
    ):
        _add_qa_flag(
            flags,
            "meta_description_missing_action",
            "Meta description does not include a clear next action for this business or page type.",
            "meta_description",
            severity="warning",
        )


def _add_faq_quality_flags(flags: list[dict], faq_items: list):
    seen_questions = set()
    for index, item in enumerate(faq_items or []):
        if not isinstance(item, dict):
            continue
        question = str(item.get("question") or "").strip()
        answer = str(item.get("answer") or "").strip()
        output = f"faq_{index + 1}"
        normalised_question = _normalise_phrase(question.rstrip("?"))
        if normalised_question and normalised_question in seen_questions:
            _add_qa_flag(flags, "duplicate_faq_question", "Duplicate FAQ question found.", output, severity="review")
        seen_questions.add(normalised_question)
        if question and not question.endswith("?"):
            _add_qa_flag(flags, "faq_question_missing_question_mark", "FAQ question does not end with a question mark.", output, severity="warning")
        if question and not answer:
            _add_qa_flag(flags, "faq_answer_missing", "FAQ question has no answer.", output, severity="review")
        elif len(answer.split()) < 15:
            _add_qa_flag(
                flags,
                "faq_answer_very_short",
                "FAQ answer is very short and may not fully answer the question.",
                output,
                severity="warning",
            )
        combined = f"{question}\n{answer}"
        for topic in _FAQ_RISKY_TOPICS:
            if _contains_forbidden_phrase(combined, topic):
                _add_qa_flag(
                    flags,
                    "faq_risky_mutable_topic",
                    f'FAQ discusses mutable or policy-sensitive topic "{topic}" and needs evidence review.',
                    output,
                    topic,
                    severity="warning",
                )
                break


def _add_keyword_presence_flags(
    flags: list[dict],
    primary_keyword: str,
    gen_meta: bool,
    gen_page_copy: bool,
    meta_text: str,
    page_copy_text: str,
):
    keyword = (primary_keyword or "").strip()
    if not keyword:
        return

    if gen_meta and meta_text.strip() and not _keyword_present(keyword, meta_text):
        flags.append({
            "code": "target_keyword_missing_from_meta",
            "message": "Target keyword was not found in the generated meta output.",
            "output": "meta",
            "keyword": keyword,
        })

    if gen_page_copy and page_copy_text.strip() and not _keyword_present(keyword, page_copy_text):
        flags.append({
            "code": "target_keyword_missing_from_page_copy",
            "message": "Target keyword was not found in the generated page copy.",
            "output": "page_copy",
            "keyword": keyword,
        })


def _add_section_word_count_flags(flags: list[dict], section_results: dict, template: dict | None):
    if not section_results or not template:
        return

    for section in template.get("sections", []):
        section_name = section.get("name", "")
        text = section_results.get(section_name, "")
        if not str(text or "").strip():
            continue

        target = section.get("word_count") or []
        if len(target) != 2:
            continue
        target_min, target_max = target
        strict_collection_guidance = bool(
            section_name == "collection_guidance"
            and _is_collection_template(template)
        )
        countable_text = (
            _collection_section_body(str(text))
            if strict_collection_guidance
            else str(text)
        )
        actual_words = _word_count_for_qa(countable_text)
        severe_min = (
            target_min
            if strict_collection_guidance
            else int(target_min * 0.6)
        )
        tolerated_min = (
            target_min
            if strict_collection_guidance
            else int(target_min * 0.8)
        )
        tolerated_max = (
            target_max
            if strict_collection_guidance
            else int(target_max * 1.2)
        )

        if actual_words < severe_min:
            flags.append({
                "code": "section_word_count_below_target",
                "message": f"Section '{section.get('label', section_name)}' is substantially shorter than the target range.",
                "output": "page_copy",
                "section": section_name,
                "section_label": section.get("label", section_name),
                "actual_words": actual_words,
                "target_min": target_min,
                "target_max": target_max,
                "severity": "review",
            })
        elif actual_words < tolerated_min:
            flags.append({
                "code": "section_word_count_below_target",
                "message": f"Section '{section.get('label', section_name)}' is shorter than the target range.",
                "output": "page_copy",
                "section": section_name,
                "section_label": section.get("label", section_name),
                "actual_words": actual_words,
                "target_min": target_min,
                "target_max": target_max,
                "severity": "warning",
            })
        elif actual_words > tolerated_max:
            flags.append({
                "code": "section_word_count_above_target",
                "message": f"Section '{section.get('label', section_name)}' is longer than the target range.",
                "output": "page_copy",
                "section": section_name,
                "section_label": section.get("label", section_name),
                "actual_words": actual_words,
                "target_min": target_min,
                "target_max": target_max,
                "severity": "review",
            })


def _first_markdown_heading(text: str) -> tuple[str, str] | None:
    for line in str(text or "").splitlines():
        if not line.strip():
            continue
        match = re.match(r"^\s*(#{1,3})\s+(.+?)\s*$", line)
        if not match:
            return None
        return f"h{len(match.group(1))}", match.group(2).strip()
    return None


def _page_plan_contracts(strategy_brief: dict | None) -> dict[str, dict]:
    brief = strategy_brief if isinstance(strategy_brief, dict) else {}
    return {
        str(item.get("section") or "").strip().casefold(): item
        for item in brief.get("section_guidance") or []
        if isinstance(item, dict) and str(item.get("section") or "").strip()
    }


_SOURCE_LIST_ITEM_RE = re.compile(r"^\s*[-+*]\s+(.+?)\s*$")
_SOURCE_TESTIMONIAL_ATOMIC_GAP_CHARS = 40
_SOURCE_TESTIMONIAL_FORMATTING_GAP_RE = re.compile(
    r"""^[\s>*_~`|:;,'"\-–—]*$"""
)


def _normalise_exact_source_phrase(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip())


def _exact_source_phrase_pattern(text: str) -> str:
    parts = re.split(r"(\s+)", str(text or "").strip())
    return "".join(
        r"\s+" if part.isspace() else re.escape(part)
        for part in parts
        if part
    )


def _source_named_list_is_one_unit(text: str, items: list) -> bool:
    required_items = [
        _normalise_exact_source_phrase(item)
        for item in items
        if str(item or "").strip()
    ]
    if not required_items:
        return True

    runs = []
    current_run = []
    contiguous_runs = []
    current_contiguous_run = []
    for line in str(text or "").splitlines():
        match = _SOURCE_LIST_ITEM_RE.match(line)
        if match:
            item = _normalise_exact_source_phrase(match.group(1))
            current_run.append(item)
            current_contiguous_run.append(item)
            continue
        if line.strip() and current_run:
            runs.append(current_run)
            current_run = []
        if current_contiguous_run:
            contiguous_runs.append(current_contiguous_run)
            current_contiguous_run = []
    if current_run:
        runs.append(current_run)
    if current_contiguous_run:
        contiguous_runs.append(current_contiguous_run)

    strict_matching_runs = [
        run
        for run in contiguous_runs
        if run == required_items
    ]
    if strict_matching_runs:
        return len(strict_matching_runs) == 1

    loose_matching_runs = [
        run
        for run in runs
        if run == required_items
    ]
    return len(loose_matching_runs) == 1


def _source_testimonial_is_atomic(
    text: str,
    quote: str,
    attribution: str,
) -> bool:
    source_text = str(text or "").replace("\r\n", "\n")
    quote_pattern = _exact_source_phrase_pattern(quote)
    attribution_pattern = _exact_source_phrase_pattern(attribution)
    if not quote_pattern or not attribution_pattern:
        return False

    for quote_match in re.finditer(quote_pattern, source_text):
        line_start = source_text.rfind("\n", 0, quote_match.start()) + 1
        prefix = source_text[line_start:quote_match.start()]
        if (
            len(prefix) > _SOURCE_TESTIMONIAL_ATOMIC_GAP_CHARS
            or not _SOURCE_TESTIMONIAL_FORMATTING_GAP_RE.fullmatch(prefix)
        ):
            continue
        tail = source_text[quote_match.end():]
        for attribution_match in re.finditer(attribution_pattern, tail):
            gap = tail[:attribution_match.start()]
            if len(gap) > _SOURCE_TESTIMONIAL_ATOMIC_GAP_CHARS:
                break
            attribution_end = quote_match.end() + attribution_match.end()
            line_end = source_text.find("\n", attribution_end)
            if line_end < 0:
                line_end = len(source_text)
            suffix = source_text[attribution_end:line_end]
            if (
                _SOURCE_TESTIMONIAL_FORMATTING_GAP_RE.fullmatch(gap)
                and len(suffix) <= _SOURCE_TESTIMONIAL_ATOMIC_GAP_CHARS
                and _SOURCE_TESTIMONIAL_FORMATTING_GAP_RE.fullmatch(
                    suffix
                )
            ):
                return True
    return False


def _add_legacy_generic_heading_qa_flags(
    flags: list[dict],
    section_results: dict,
    template: dict | None,
):
    for section in (template or {}).get("sections") or []:
        expected_level = str(section.get("heading_level") or "").casefold()
        if expected_level not in {"h2", "h3"}:
            continue
        section_name = str(section.get("name") or "")
        actual = _first_markdown_heading(section_results.get(section_name, ""))
        if not actual:
            continue
        actual_heading = actual[1]
        if (
            actual_heading.casefold() not in _GENERIC_PAGE_HEADINGS
            and actual_heading.casefold()
            != str(section.get("label") or "").strip().casefold()
        ):
            continue
        flags.append({
            "code": "page_heading_generic",
            "message": (
                f'Section "{section.get("label", section_name)}" uses '
                "a generic reader-facing heading."
            ),
            "output": "page_copy",
            "section": section_name,
            "actual_heading": actual_heading,
            "severity": "review",
        })


def _add_page_plan_qa_flags(
    flags: list[dict],
    section_results: dict,
    template: dict | None,
    strategy_brief: dict | None,
    page_quality_policy,
    canonical_h1: str = "",
    forbidden_phrases: list[str] | None = None,
    page_copy_correction_enabled: bool = False,
):
    if not section_results or not template:
        return
    claim_bound_rendering = bool(
        isinstance(strategy_brief, dict)
        and strategy_brief.get("claim_bound_renderer_version")
        == CLAIM_BOUND_RENDERER_VERSION
        and strategy_brief.get("source_block_plan_version")
        == SOURCE_BLOCK_PLAN_VERSION
    )
    contracts = _page_plan_contracts(strategy_brief)
    actual_headings: dict[str, dict] = {}
    exact_headings_enabled = bool(
        page_quality_policy
        and page_quality_policy.exact_planned_headings
        and not claim_bound_rendering
    )
    coverage_enabled = bool(
        page_quality_policy
        and page_quality_policy.coverage_points
        and not claim_bound_rendering
    )
    source_asset_quality_enabled = bool(
        (
            strategy_brief.get("source_asset_manifest_version")
            if isinstance(strategy_brief, dict)
            else ""
        )
        == SOURCE_ASSET_MANIFEST_VERSION
    )
    raw_source_asset_diagnostics = (
        (strategy_brief or {}).get("source_asset_mapping_diagnostics")
        if isinstance(strategy_brief, dict)
        else None
    )
    source_asset_diagnostics = (
        raw_source_asset_diagnostics
        if isinstance(raw_source_asset_diagnostics, dict)
        else {}
    )
    validated_source_asset_sections = (
        _validated_source_asset_section_names(strategy_brief)
        if source_asset_quality_enabled
        else set()
    )
    unassigned_source_asset_ids = [
        str(value)
        for value in source_asset_diagnostics.get(
            "unassigned_asset_ids"
        ) or []
        if str(value)
    ]
    if (
        source_asset_quality_enabled
        and not claim_bound_rendering
        and source_asset_diagnostics.get("active")
        and unassigned_source_asset_ids
    ):
        flags.append({
            "code": "page_source_assets_unassigned",
            "message": (
                "One or more owned-page source assets were not assigned to a "
                "page section and need relevance review."
            ),
            "output": "page_copy",
            "asset_ids": unassigned_source_asset_ids,
            "severity": "review",
        })

    for section in template.get("sections") or []:
        section_name = str(section.get("name") or "")
        expected_level = str(section.get("heading_level") or "").casefold()
        contract = contracts.get(section_name.casefold(), {})
        evidence_sparse = bool(
            page_copy_correction_enabled
            and section.get("evidence_sparse") is True
        )
        planned_heading = str(
            (
                section.get("planned_heading")
                if page_copy_correction_enabled
                else contract.get("planned_heading")
            )
            or ""
        ).strip()
        actual = _first_markdown_heading(section_results.get(section_name, ""))

        if exact_headings_enabled and expected_level == "h1":
            if not actual:
                flags.append({
                    "code": "page_h1_missing",
                    "message": (
                        f'Section "{section.get("label", section_name)}" does not '
                        "start with its required H1."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "severity": "review",
                })
            elif (
                actual[0] == "h1"
                and str(canonical_h1 or "").strip()
                and actual[1] != str(canonical_h1).strip()
            ):
                flags.append({
                    "code": "page_h1_canonical_mismatch",
                    "message": (
                        f'Section "{section.get("label", section_name)}" does not '
                        "use its exact canonical H1."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "canonical_h1": str(canonical_h1).strip(),
                    "actual_heading": actual[1],
                    "severity": "review",
                })

        if exact_headings_enabled and expected_level in {"h2", "h3"}:
            if evidence_sparse and not actual:
                flags.append({
                    "code": "page_heading_missing",
                    "message": (
                        f'Section "{section.get("label", section_name)}" does '
                        f"not start with its required {expected_level.upper()}."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "expected_level": expected_level,
                    "severity": "review",
                })
            elif not planned_heading and not evidence_sparse:
                flags.append({
                    "code": "page_planned_heading_missing",
                    "message": (
                        f'Section "{section.get("label", section_name)}" has no '
                        "accepted reader-facing heading plan."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "severity": "review",
                })
            elif planned_heading and (
                not actual or actual[1] != planned_heading
            ):
                flags.append({
                    "code": "page_heading_plan_mismatch",
                    "message": (
                        f'Section "{section.get("label", section_name)}" does not use '
                        "its exact accepted planned heading."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "planned_heading": planned_heading,
                    "actual_heading": actual[1] if actual else "",
                    "severity": "review",
                })

        if exact_headings_enabled and actual:
            actual_level, actual_heading = actual
            if (
                expected_level in {"h2", "h3"}
                and not evidence_sparse
                and (
                    actual_heading.casefold() in _GENERIC_PAGE_HEADINGS
                    or actual_heading.casefold()
                    == str(section.get("label") or "").strip().casefold()
                )
            ):
                flags.append({
                    "code": "page_heading_generic",
                    "message": (
                        f'Section "{section.get("label", section_name)}" uses '
                        "a generic reader-facing heading."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "actual_heading": actual_heading,
                    "severity": "review",
                })
            heading_entry = actual_headings.setdefault(
                actual_heading.casefold(),
                {"heading": actual_heading, "sections": []},
            )
            heading_entry["sections"].append(section_name)
            if expected_level in {"h1", "h2", "h3"} and actual_level != expected_level:
                flags.append({
                    "code": "page_heading_hierarchy",
                    "message": (
                        f'Section "{section.get("label", section_name)}" starts with '
                        f"{actual_level.upper()} instead of {expected_level.upper()}."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "expected_level": expected_level,
                    "actual_level": actual_level,
                    "severity": "review",
                })
            elif expected_level == "none":
                flags.append({
                    "code": "page_heading_hierarchy",
                    "message": (
                        f'Section "{section.get("label", section_name)}" adds '
                        "a heading even though the template requires body copy only."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "expected_level": "none",
                    "actual_level": actual_level,
                    "severity": "review",
                })

        target = section.get("word_count") or []
        text = str(section_results.get(section_name) or "")
        source_asset_section_enabled = (
            section_name.casefold() in validated_source_asset_sections
        )
        source_asset_conflicts = (
            _source_asset_forbidden_conflicts(
                strategy_brief,
                section_name,
                forbidden_phrases or [],
            )
            if source_asset_section_enabled
            else []
        )
        conflicting_source_asset_ids = {
            conflict["asset_id"]
            for conflict in source_asset_conflicts
            if conflict.get("asset_id")
        }
        conflicting_source_phrase_keys = {
            str(conflict.get("source_phrase") or "").strip().casefold()
            for conflict in source_asset_conflicts
            if str(conflict.get("source_phrase") or "").strip()
        }
        if source_asset_conflicts:
            flags.append({
                "code": "page_source_asset_forbidden_conflict",
                "message": (
                    f'Section "{section.get("label", section_name)}" has '
                    "assigned exact source material that was deferred because "
                    "it conflicts with a configured forbidden phrase."
                ),
                "output": "page_copy",
                "section": section_name,
                "asset_ids": sorted(conflicting_source_asset_ids),
                "forbidden_phrases": list(dict.fromkeys(
                    str(conflict.get("forbidden_phrase") or "").strip()
                    for conflict in source_asset_conflicts
                    if str(conflict.get("forbidden_phrase") or "").strip()
                )),
                "severity": "review",
            })
        required_named_items = []
        seen_required_items = set()
        raw_required_items = contract.get("required_named_items") or []
        if source_asset_quality_enabled and not source_asset_section_enabled:
            raw_required_items = []
        elif not source_asset_section_enabled:
            raw_required_items = raw_required_items[:12]
        for value in raw_required_items:
            required_item = str(value or "").strip()
            if required_item.casefold() in conflicting_source_phrase_keys:
                continue
            if not source_asset_section_enabled:
                required_item = required_item[:160]
            required_key = required_item.casefold()
            if required_item and required_key not in seen_required_items:
                seen_required_items.add(required_key)
                required_named_items.append(required_item)
        missing_named_items = [
            required_item
            for required_item in required_named_items
            if not _contains_forbidden_phrase(
                _normalise_phrase(text),
                _normalise_phrase(required_item),
            )
        ]
        if text.strip() and missing_named_items:
            flags.append({
                "code": "page_required_source_item_missing",
                "message": (
                    f'Section "{section.get("label", section_name)}" omits '
                    "one or more required names or source paths."
                ),
                "output": "page_copy",
                "section": section_name,
                "missing_items": missing_named_items,
                "severity": "review",
            })
        related_profile_subjects = [
            str(value or "").strip()
            for value in contract.get("related_profile_subjects") or []
            if str(value or "").strip()
        ]
        if related_profile_subjects and text.strip():
            profile_body = re.sub(
                r"^\s{0,3}#{1,6}\s+[^\n]*(?:\n|$)",
                "",
                text,
                count=1,
            )
            missing_profile_subjects = [
                name
                for name in related_profile_subjects
                if not _contains_forbidden_phrase(
                    _normalise_phrase(profile_body),
                    _normalise_phrase(name),
                )
            ]
            if missing_profile_subjects:
                flags.append({
                    "code": "page_related_profile_missing",
                    "message": (
                        f'Section "{section.get("label", section_name)}" '
                        "names a related profile in its plan but omits that "
                        "person from the section body."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "missing_items": missing_profile_subjects,
                    "severity": "review",
                })
        if source_asset_section_enabled and text.strip():
            missing_testimonials = []
            unpreserved_source_lists = []
            source_statements_needing_review = []
            normalized_text = _normalise_phrase(text)
            for asset in contract.get("source_assets") or []:
                if not isinstance(asset, dict):
                    continue
                if (
                    str(asset.get("id") or "")
                    in conflicting_source_asset_ids
                ):
                    continue
                if asset.get("kind") == "direct_statement":
                    statement = str(asset.get("statement") or "").strip()
                    if statement and not _contains_forbidden_phrase(
                        normalized_text,
                        _normalise_phrase(statement),
                    ):
                        source_statements_needing_review.append(
                            str(asset.get("id") or "").strip()
                        )
                    continue
                if asset.get("kind") == "named_list":
                    if not _source_named_list_is_one_unit(
                        text,
                        asset.get("items") or [],
                    ):
                        unpreserved_source_lists.append(
                            str(asset.get("id") or "").strip()
                        )
                    continue
                if asset.get("kind") != "testimonial":
                    continue
                asset_id = str(asset.get("id") or "").strip()
                missing_components = []
                for component, value in (
                    ("quote", asset.get("quote")),
                    ("attribution", asset.get("attribution")),
                ):
                    exact_value = str(value or "").strip()
                    if exact_value and not _contains_forbidden_phrase(
                        normalized_text,
                        _normalise_phrase(exact_value),
                    ):
                        missing_components.append(component)
                if (
                    not missing_components
                    and not _source_testimonial_is_atomic(
                        text,
                        str(asset.get("quote") or ""),
                        str(asset.get("attribution") or ""),
                    )
                ):
                    missing_components.append("atomic_pair")
                if missing_components:
                    missing_testimonials.append({
                        "asset_id": asset_id,
                        "missing_components": missing_components,
                    })
            if missing_testimonials:
                flags.append({
                    "code": "page_required_testimonial_missing",
                    "message": (
                        f'Section "{section.get("label", section_name)}" omits '
                        "or changes an assigned testimonial quote or attribution."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "missing_testimonials": missing_testimonials,
                    "severity": "review",
                })
            if unpreserved_source_lists:
                flags.append({
                    "code": "page_required_source_list_not_preserved",
                    "message": (
                        f'Section "{section.get("label", section_name)}" does '
                        "not preserve an assigned named source list as one "
                        "complete list."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "asset_ids": unpreserved_source_lists,
                    "severity": "review",
                })
            if source_statements_needing_review:
                flags.append({
                    "code": (
                        "page_source_statement_preservation_needs_review"
                    ),
                    "message": (
                        f'Section "{section.get("label", section_name)}" does '
                        "not preserve one or more assigned direct source "
                        "statements verbatim. Review whether each supported "
                        "proposition remains intact without added claims."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "asset_ids": source_statements_needing_review,
                    "severity": "review",
                })
        if coverage_enabled and text.strip() and len(target) == 2:
            target_min = int(target[0])
            actual_words = _word_count_for_qa(text)
            if source_asset_quality_enabled:
                target_midpoint = (
                    int(target[0]) + int(target[1])
                ) // 2
                planned_depth_threshold = (
                    target_min
                    if page_copy_correction_enabled
                    else int(target_midpoint * 0.85)
                )
            else:
                target_midpoint = None
                planned_depth_threshold = int(target_min * 0.8)
            if actual_words < planned_depth_threshold:
                depth_target_label = (
                    "planned minimum depth"
                    if page_copy_correction_enabled
                    else "planned depth review threshold"
                )
                flags.append({
                    "code": "page_section_below_planned_depth",
                    "message": (
                        f'Section "{section.get("label", section_name)}" is '
                        f"substantially below its {depth_target_label}."
                    ),
                    "output": "page_copy",
                    "section": section_name,
                    "actual_words": actual_words,
                    "target_min": target_min,
                    **(
                        {
                            "target_midpoint": target_midpoint,
                            "review_threshold": planned_depth_threshold,
                        }
                        if target_midpoint is not None
                        else {}
                    ),
                    "severity": "review",
                })

    for heading_entry in (
        actual_headings.values()
        if exact_headings_enabled
        else []
    ):
        heading = heading_entry["heading"]
        section_names = heading_entry["sections"]
        if heading and len(section_names) > 1:
            flags.append({
                "code": "page_heading_duplicate",
                "message": f'The heading "{heading}" is repeated across page sections.',
                "output": "page_copy",
                "sections": section_names,
                "severity": "review",
            })


def _build_page_quality_diagnostics(
    *,
    strategy_brief: dict | None,
    template: dict | None,
    section_results: dict,
    page_quality: dict,
) -> dict:
    plan = page_plan_diagnostics(
        strategy_brief,
        (template or {}).get("sections", []),
    )
    contracts = _page_plan_contracts(strategy_brief)
    sections = []
    findings = list(plan.get("findings") or [])

    for section in (template or {}).get("sections") or []:
        section_name = str(section.get("name") or "")
        contract = contracts.get(section_name.casefold(), {})
        actual = _first_markdown_heading(section_results.get(section_name, ""))
        expected_level = str(section.get("heading_level") or "").casefold()
        actual_heading = actual[1] if actual else ""
        if expected_level in {"h2", "h3"} and not actual_heading:
            findings.append({
                "code": "actual_heading_not_reliably_parsed",
                "section": section_name,
                "message": "No leading Markdown heading could be parsed reliably.",
            })
        elif actual_heading and (
            actual_heading.casefold()
            == str(section.get("label") or "").strip().casefold()
            or actual_heading.casefold() in _GENERIC_PAGE_HEADINGS
        ):
            findings.append({
                "code": "actual_heading_generic",
                "section": section_name,
                "message": "The generated heading is generic or repeats the internal template label.",
            })
        sections.append({
            "section": section_name,
            "section_label": section.get("label") or section_name,
            "expected_heading_level": expected_level,
            "planned_heading": contract.get("planned_heading") or "",
            "actual_heading": actual_heading,
            "actual_heading_level": actual[0] if actual else "",
            "coverage_points": list(contract.get("coverage_points") or []),
            "depth_policy": contract.get("depth_policy") or section.get("depth_policy") or "",
            "owned_block_ids": list(contract.get("owned_block_ids") or []),
            "source_asset_ids": list(contract.get("source_asset_ids") or []),
            "retain_points": list(contract.get("retain_points") or []),
            "improve_points": list(contract.get("improve_points") or []),
            "actual_words": _word_count_for_qa(
                str(section_results.get(section_name) or "")
            ),
            "planned_word_range": list(section.get("word_count") or []),
        })

    guidance = page_quality.get("guidance")
    return {
        "page_quality_policy_version": page_quality["page_quality_policy_version"],
        "adaptive_policy_version": page_quality["adaptive_policy_version"],
        "owned_page_mapping_version": page_quality["owned_page_mapping_version"],
        "source_asset_manifest_version": page_quality.get(
            "source_asset_manifest_version",
            "",
        ),
        "claim_bound_renderer_version": page_quality.get(
            "claim_bound_renderer_version",
            "",
        ),
        "source_block_plan_version": page_quality.get(
            "source_block_plan_version",
            "",
        ),
        "source_block_plan": deepcopy(
            (strategy_brief or {}).get("source_block_plan")
        ) if isinstance(strategy_brief, dict) else None,
        "guidance_profile": (
            {
                "id": guidance.id,
                "label": guidance.label,
                "version": guidance.version,
            }
            if guidance
            else None
        ),
        "summary": {
            key: value
            for key, value in plan.items()
            if key not in {"findings"}
        },
        "sections": sections,
        "findings": findings,
    }


def _template_for_page_copy(
    template: dict,
    separate_faq_output_enabled: bool,
    *,
    versioned_blog_h1: bool = False,
) -> dict:
    page_template = deepcopy(template)
    sections = page_template.get("sections") or []
    if (
        versioned_blog_h1
        and str(page_template.get("page_type") or "").casefold() == "blog"
        and not any(
            str(section.get("heading_level") or "").casefold() == "h1"
            for section in sections
        )
    ):
        intro = next(
            (
                section
                for section in sections
                if str(section.get("name") or "").casefold() == "intro"
                and str(section.get("heading_level") or "none").casefold()
                == "none"
            ),
            None,
        )
        if intro is not None:
            intro["heading_level"] = "h1"

    if not separate_faq_output_enabled:
        return page_template

    adjusted_sections = []
    for section in sections:
        name = str(section.get("name", "")).lower()
        label = str(section.get("label", "")).lower()
        is_faq_section = "faq" in name or label == "frequently asked questions"
        if not is_faq_section:
            adjusted_sections.append(section)

    page_template["sections"] = adjusted_sections
    return page_template


def _add_similarity_flag(result: dict, code: str, message: str, output: str, previous_row: int, similarity: float):
    flags = result.setdefault("qa_flags", [])
    if any(flag.get("code") == code and flag.get("output") == output for flag in flags):
        return
    flags.append({
        "code": code,
        "message": message,
        "output": output,
        "severity": "review",
        "similar_to_row": previous_row,
        "similarity": round(similarity, 3),
    })
    if result.get("status") in {"ok", "warning"}:
        result["status"] = "review"


def _apply_cross_row_uniqueness_flags(result: dict, previous_results: list[dict]) -> dict:
    if not result or result.get("status") in {"error", "skipped: no keywords found"}:
        return result

    checks = [
        {
            "field": "generated_title",
            "output": "meta_title",
            "code": "meta_title_similar_to_row",
            "message": "Meta title is very similar to a previous row.",
            "threshold": 0.90,
            "min_tokens": 4,
        },
        {
            "field": "generated_description",
            "output": "meta_description",
            "code": "meta_description_similar_to_row",
            "message": "Meta description is very similar to a previous row.",
            "threshold": 0.82,
            "min_tokens": 8,
        },
        {
            "field": "page_intro",
            "output": "page_intro",
            "code": "page_intro_similar_to_row",
            "message": "The first page-copy section is very similar to a previous row.",
            "threshold": 0.78,
            "min_tokens": 8,
        },
        {
            "field": "page_copy",
            "output": "page_copy",
            "code": "page_copy_similar_to_row",
            "message": "Full page copy is very similar to a previous row.",
            "threshold": 0.85,
            "min_tokens": 80,
        },
    ]

    current_texts = {
        "generated_title": result.get("generated_title") or "",
        "generated_description": result.get("generated_description") or "",
        "page_intro": _first_page_copy_section(result.get("section_results") or {}),
        "page_copy": _full_page_copy_text(result.get("section_results") or {}),
    }

    for check in checks:
        current_text = current_texts[check["field"]]
        if len(_normalise_similarity_text(current_text)) < check["min_tokens"]:
            continue

        best_match = None
        for previous_index, previous in enumerate(previous_results, start=1):
            previous_texts = {
                "generated_title": previous.get("generated_title") or "",
                "generated_description": previous.get("generated_description") or "",
                "page_intro": _first_page_copy_section(previous.get("section_results") or {}),
                "page_copy": _full_page_copy_text(previous.get("section_results") or {}),
            }
            previous_text = previous_texts[check["field"]]
            if len(_normalise_similarity_text(previous_text)) < check["min_tokens"]:
                continue

            similarity = _shingle_similarity(current_text, previous_text)
            if similarity >= check["threshold"] and (best_match is None or similarity > best_match[1]):
                best_match = (previous_index, similarity)

        if best_match:
            _add_similarity_flag(
                result,
                check["code"],
                check["message"],
                check["output"],
                best_match[0],
                best_match[1],
            )

    return result


def _extract_gap_topic_candidates(text: str) -> list[str]:
    tokens = [
        token for token in _normalise_similarity_text(text)
        if len(token) > 3 and token not in _CONTENT_GAP_STOPWORDS
    ]
    candidates = []
    seen = set()
    for index in range(len(tokens) - 1):
        phrase = f"{tokens[index]} {tokens[index + 1]}"
        if phrase not in seen:
            candidates.append(phrase)
            seen.add(phrase)
    if candidates:
        return candidates
    return [token for token in tokens if token not in seen]


def _topic_present_in_text(topic: str, text: str) -> bool:
    topic_tokens = _normalise_similarity_text(topic)
    text_tokens = set(_normalise_similarity_text(text))
    if not topic_tokens or not text_tokens:
        return False
    return any(token in text_tokens for token in topic_tokens)


def _build_content_gap_summary(competitor_section_map: dict, section_results: dict) -> list[dict]:
    diagnostics = []
    if not competitor_section_map or not section_results:
        return diagnostics

    for section, excerpts in competitor_section_map.items():
        generated_text = section_results.get(section, "")
        if not str(generated_text or "").strip():
            continue

        missing_topics = []
        seen = set()
        for excerpt in excerpts or []:
            for topic in _extract_gap_topic_candidates(str(excerpt)):
                if topic in seen or _topic_present_in_text(topic, generated_text):
                    continue
                missing_topics.append(topic)
                seen.add(topic)
                if len(missing_topics) >= 3:
                    break
            if len(missing_topics) >= 3:
                break

        if missing_topics:
            diagnostics.append({
                "section": section,
                "missing_topics": missing_topics,
                "summary": (
                    "Competitors mention "
                    + ", ".join(missing_topics)
                    + " but this section does not clearly cover them."
                ),
            })

    return diagnostics


def _row_topic_text(row: dict) -> str:
    parts = [
        row.get("primary_keyword") or "",
        row.get("h1") or "",
        row.get("generated_title") or "",
        row.get("generated_description") or "",
        row.get("optimised_h1") or "",
        _full_page_copy_text(row.get("section_results") or {}),
    ]
    for item in row.get("faq_items") or []:
        if isinstance(item, dict):
            parts.append(f"{item.get('question', '')} {item.get('answer', '')}")
    return "\n".join(str(part) for part in parts if str(part or "").strip())


def _build_internal_link_suggestions(results: list[dict], max_per_source: int = 3) -> list[dict]:
    eligible = [
        row for row in results
        if (row.get("url") or "").startswith("http")
        and row.get("status") not in {"error", "skipped: invalid URL", "skipped: no keywords found"}
    ]
    suggestions = []

    for source in eligible:
        source_url = source.get("url", "")
        source_text = _row_topic_text(source)
        if not source_text.strip():
            continue

        source_suggestions = []
        source_primary = " ".join(_normalise_similarity_text(source.get("primary_keyword") or ""))
        for target in eligible:
            target_url = target.get("url", "")
            if not target_url or target_url == source_url:
                continue

            anchor = (target.get("primary_keyword") or target.get("h1") or "").strip()
            if not anchor:
                continue
            if source_primary and source_primary == " ".join(_normalise_similarity_text(anchor)):
                continue

            if _keyword_present(anchor, source_text):
                confidence = 0.9
                reason = "Target keyword appears naturally in the source page copy."
            else:
                target_tokens = _meaningful_keyword_tokens(anchor)
                source_tokens = set(_normalise_similarity_text(source_text))
                if not target_tokens:
                    continue
                matched = [token for token in target_tokens if token in source_tokens]
                if len(matched) < max(2, len(target_tokens)):
                    continue
                confidence = min(0.85, 0.6 + (0.08 * len(matched)))
                reason = "Source page topic overlaps with the target page keyword."

            source_suggestions.append({
                "source_url": source_url,
                "target_url": target_url,
                "anchor_text": anchor,
                "confidence": round(confidence, 2),
                "reason": reason,
            })

        source_suggestions.sort(key=lambda item: item["confidence"], reverse=True)
        suggestions.extend(source_suggestions[:max_per_source])

    return suggestions[:200]


def _collect_qa_flags(
    *,
    gen_meta: bool,
    gen_faqs: bool,
    gen_page_copy: bool,
    generated_title: str,
    generated_description: str,
    optimised_h1: str,
    input_h1: str,
    primary_keyword: str,
    faq_items: list,
    requested_faq_count: int = 0,
    section_results: dict,
    forbidden_phrases: list[str],
    template: dict | None = None,
    brand_name: str = "",
    business_type: str = "general",
    page_type: str = "general",
    strategy_brief: dict | None = None,
    page_quality_policy_version: str = "",
    page_copy_correction_enabled: bool = False,
    page_copy_quality_block_reasons: list[str] | None = None,
) -> list[dict]:
    flags = []
    claim_bound_rendering = bool(
        isinstance(strategy_brief, dict)
        and strategy_brief.get("claim_bound_renderer_version")
        == CLAIM_BOUND_RENDERER_VERSION
        and strategy_brief.get("source_block_plan_version")
        == SOURCE_BLOCK_PLAN_VERSION
    )

    if gen_page_copy and page_copy_quality_block_reasons:
        _add_qa_flag(
            flags,
            "page_copy_quality_blocked",
            (
                "Page copy was withheld because the evidence-locked source "
                "contract could not be completed safely."
            ),
            "page_copy",
            severity="error",
        )
        flags[-1]["details"] = [
            str(reason)
            for reason in page_copy_quality_block_reasons[:10]
            if str(reason or "")
        ]

    if gen_meta:
        if not (generated_title or "").strip():
            _add_qa_flag(flags, "meta_missing_title", "Meta title was requested but no title was generated.", "meta")
        if not (generated_description or "").strip():
            _add_qa_flag(flags, "meta_missing_description", "Meta description was requested but no description was generated.", "meta")
        if not (optimised_h1 or "").strip():
            _add_qa_flag(flags, "meta_missing_h1", "Optimized H1 was requested but no H1 was generated.", "meta")
        if (generated_title or "").strip().lower() == (input_h1 or "").strip().lower() and (input_h1 or "").strip():
            _add_qa_flag(flags, "meta_title_matches_h1", "Generated title matches the input H1.", "meta")

    if gen_faqs and not faq_items:
        _add_qa_flag(flags, "faq_missing", "FAQs were requested but no FAQ items were generated.", "faq")
    elif gen_faqs and requested_faq_count > 0 and len(faq_items) < requested_faq_count:
        _add_qa_flag(
            flags,
            "faq_count_incomplete",
            f"Only {len(faq_items)} of {requested_faq_count} requested FAQs were generated.",
            "faq",
            severity="review",
        )

    page_copy_text = "\n\n".join(str(v) for k, v in (section_results or {}).items() if not str(k).startswith("_"))
    if gen_page_copy and not page_copy_text.strip():
        _add_qa_flag(flags, "page_copy_missing", "Page copy was requested but no page sections were generated.", "page_copy")
    elif gen_page_copy:
        if (
            page_copy_correction_enabled
            and re.search(
                r"\[\[[ \t]*COPYPILOT_SOURCE_",
                page_copy_text,
                re.IGNORECASE,
            )
        ):
            _add_qa_flag(
                flags,
                "internal_source_marker",
                "An internal source-placement marker remains in page copy.",
                "page_copy",
                severity="review",
            )
        if page_copy_correction_enabled and not claim_bound_rendering:
            for finding in _structured_source_duplicate_findings(
                section_results,
                strategy_brief,
                forbidden_phrases,
            ):
                flags.append({
                    "code": "page_structured_source_duplicate",
                    "message": (
                        f'Section "{finding["section"]}" repeats exact '
                        "source wording outside its canonical source unit."
                    ),
                    "output": "page_copy",
                    "section": finding["section"],
                    "asset_ids": finding["asset_ids"],
                    "duplicate_phrases": finding["duplicate_phrases"],
                    "severity": "review",
                })
            _add_page_copy_evidence_backstop_flags(
                flags,
                section_results,
                template,
                strategy_brief,
                forbidden_phrases,
                page_type=page_type,
                brand_name=brand_name,
            )
        if not claim_bound_rendering:
            _add_collection_section_role_flags(
                flags,
                section_results,
                template,
            )
            _add_section_word_count_flags(flags, section_results, template)
        page_quality_policy = (
            get_page_quality_policy(page_quality_policy_version)
            if page_quality_policy_version
            else None
        )
        if page_quality_policy:
            _add_page_plan_qa_flags(
                flags,
                section_results,
                template,
                strategy_brief,
                page_quality_policy,
                canonical_h1=optimised_h1 or input_h1,
                forbidden_phrases=forbidden_phrases,
                page_copy_correction_enabled=page_copy_correction_enabled,
            )
            planned_depth_sections = {
                str(flag.get("section") or "")
                for flag in flags
                if flag.get("code") == "page_section_below_planned_depth"
            }
            flags[:] = [
                flag
                for flag in flags
                if not (
                    flag.get("code") == "section_word_count_below_target"
                    and str(flag.get("section") or "") in planned_depth_sections
                )
            ]
        else:
            _add_legacy_generic_heading_qa_flags(
                flags,
                section_results,
                template,
            )
        authored_page_copy_text = "" if claim_bound_rendering else (
            _page_copy_without_materialized_source_units(
                section_results,
                strategy_brief,
                forbidden_phrases,
            )
            if page_quality_policy and page_copy_correction_enabled
            else page_copy_text
        )
        _add_repeated_phrase_flags(flags, authored_page_copy_text)

    if gen_meta:
        _add_generic_opener_flags(
            flags,
            generated_description or "",
            section_results if gen_page_copy else {},
            strategy_brief,
        )
    elif gen_page_copy:
        _add_generic_opener_flags(
            flags,
            "",
            section_results,
            strategy_brief,
        )
    if gen_page_copy:
        _add_generic_page_reference_flags(
            flags,
            section_results,
            strategy_brief,
        )

    if gen_meta and gen_page_copy:
        _add_h1_alignment_flag(flags, optimised_h1 or "", section_results)

    meta_text = "\n".join([generated_title or "", generated_description or "", optimised_h1 or ""])
    _add_keyword_presence_flags(
        flags,
        primary_keyword=primary_keyword,
        gen_meta=gen_meta,
        gen_page_copy=gen_page_copy,
        meta_text=meta_text,
        page_copy_text=page_copy_text,
    )
    _add_meta_field_quality_flags(
        flags,
        gen_meta=gen_meta,
        primary_keyword=primary_keyword,
        generated_title=generated_title,
        generated_description=generated_description,
        business_type=business_type,
        page_type=page_type,
    )
    _add_keyword_placement_flags(
        flags,
        primary_keyword=primary_keyword,
        gen_meta=gen_meta,
        gen_page_copy=gen_page_copy,
        optimised_h1=optimised_h1,
        section_results=section_results,
    )
    _add_meta_length_flags(
        flags,
        gen_meta=gen_meta,
        generated_title=generated_title,
        generated_description=generated_description,
    )
    _add_brand_in_h1_flags(
        flags,
        brand_name=brand_name,
        gen_meta=gen_meta,
        gen_page_copy=gen_page_copy,
        optimised_h1=optimised_h1,
        section_results=section_results,
    )
    _add_faq_quality_flags(flags, faq_items)

    outputs = _generated_output_texts(
        generated_title,
        generated_description,
        optimised_h1,
        faq_items,
        section_results,
    )
    authored_outputs = [
        (
            output,
            (
                _without_exact_source_asset_phrases(
                    text,
                    strategy_brief,
                )
                if output == "page_copy"
                else text
            ),
        )
        for output, text in outputs
    ]
    _add_exclamation_flags(flags, authored_outputs)
    _add_b2b_consumer_cta_flags(
        flags,
        business_type,
        authored_outputs,
    )
    _add_us_english_qa_flags(
        flags,
        authored_outputs,
        [brand_name, input_h1],
    )

    seen_matches = set()
    for output, text in outputs:
        for phrase in forbidden_phrases:
            key = (output, phrase.lower())
            if key in seen_matches:
                continue
            if _contains_forbidden_phrase(text, phrase):
                _add_qa_flag(
                    flags,
                    "forbidden_phrase",
                    f"Forbidden phrase found in {output}.",
                    output,
                    phrase,
                )
                seen_matches.add(key)

    for flag in flags:
        flag.setdefault("severity", "review")
    return flags


def _process_single_row(
    row: dict,
    settings: dict,
    gsc_client,
    branded_terms: list,
    used_keywords: set,
    sb,
    job_id: str,
    row_num: int,
    total_rows: int,
    user_id: str = "",
    brand_profile: dict = None,
    gsc_auth_method: str = "disabled",
    scraper_override: str = "",
    page_copy_correction_enabled: bool = True,
) -> dict:
    row_started_at = time.monotonic()

    def step(msg: str):
        _update_job(sb, job_id, user_id, {"current_step": f"Row {row_num}/{total_rows}: {msg}"})

    url          = (row.get("url") or "").strip()
    manual_kws   = [k.strip() for k in (row.get("keyword") or "").split(",") if k.strip()]
    h1_raw       = (row.get("h1") or "").strip()
    h1           = "" if h1_raw.lower() == "none" else h1_raw
    page_type    = normalize_page_type(row.get("page_type") or settings.get("page_type", "service"), default="service")
    template_key = row.get("template_key") or settings.get("template_key") or default_template_key_for_page_type(page_type)
    if page_type == "landing_page" and template_key == "service_page":
        template_key = default_template_key_for_page_type(page_type)
    custom_template_text = str(
        settings.get("custom_template_text") or ""
    ).strip()

    # What to generate — from row overrides or job-level settings
    gen_page_copy = row.get("gen_page_copy", settings.get("gen_page_copy", True))
    gen_meta      = row.get("gen_meta",      settings.get("gen_meta",      True))
    gen_faqs      = row.get("gen_faqs",      settings.get("gen_faqs",      True))
    row_num_faqs  = row.get("num_faqs")
    num_faqs      = int(settings.get("num_faqs", 5) if row_num_faqs is None else row_num_faqs)
    page_quality = _stored_page_quality_context(
        settings,
        page_copy_requested=bool(gen_page_copy),
    )
    page_quality_enabled = bool(page_quality["enabled"])
    page_quality_policy = page_quality["policy"]
    page_copy_guidance = page_quality["guidance"]
    page_copy_correction_active = _page_copy_correction_is_active(
        page_quality,
        requested=bool(
            page_copy_correction_enabled
            and gen_page_copy
        ),
    )
    claim_bound_rendering_active = _claim_bound_rendering_is_active(
        page_quality,
        requested=bool(gen_page_copy),
        template_key=template_key,
        custom_template_text=custom_template_text,
    )
    active_claim_bound_renderer_version = (
        page_quality.get("claim_bound_renderer_version", "")
        if claim_bound_rendering_active
        else ""
    )
    active_source_block_plan_version = (
        page_quality.get("source_block_plan_version", "")
        if claim_bound_rendering_active
        else ""
    )
    page_planning_enabled = bool(
        page_quality_policy
        and (
            page_quality_policy.exact_planned_headings
            or page_quality_policy.coverage_points
            or page_quality_policy.bounded_owned_page_reuse
        )
    )

    dfs_login    = settings["dfs_login"]
    dfs_password = settings["dfs_password"]
    provider     = settings.get("provider", "Claude")
    model        = settings.get("model") or None
    api_key      = settings.get("api_key", "")
    brand_name   = str(settings.get("brand_name") or (brand_profile or {}).get("brand_name") or "").strip()
    business_type = settings.get("business_type", "general")
    min_volume   = int(settings.get("min_volume", 10))
    location_code = int(settings.get("location_code", 2840))
    include_brand = settings.get("include_brand", True)
    forbidden_phrases = settings.get("forbidden_phrases", "")
    forbidden_phrase_list = _split_forbidden_phrases(
        forbidden_phrases,
        (brand_profile or {}).get("words_to_avoid", ""),
    )
    forbidden_phrase_text = ", ".join(forbidden_phrase_list)

    # Build client brief — merge niche context + brand profile + custom brief
    explicit_client_brief = settings.get("client_brief", "") or ""
    client_brief = explicit_client_brief
    _niche_ctx = get_niche_context(settings.get("niche", ""))
    if _niche_ctx:
        client_brief = (client_brief + "\n\n" + _niche_ctx).strip()
    if brand_profile:
        parts = []
        if brand_profile.get("tone_of_voice"):
            parts.append("Tone of voice: " + brand_profile["tone_of_voice"])
        if brand_profile.get("key_messages"):
            parts.append("Key messages: " + brand_profile["key_messages"])
        if brand_profile.get("words_to_avoid"):
            parts.append("Words to avoid: " + brand_profile["words_to_avoid"])
        if brand_profile.get("guidelines"):
            parts.append(brand_profile["guidelines"])
        if parts:
            client_brief = (client_brief + "\n\n" + "\n".join(parts)).strip()
    brand_context = _build_brand_context(brand_profile, settings.get("niche", ""))
    brand_style_context = _build_brand_style_context(brand_profile)

    use_gsc  = settings.get("use_gsc", False)
    site_url = settings.get("site_url", "")
    jina_key = settings.get("jina_api_key", "")

    run_diagnostics = {
        "provider": provider,
        "model": model or "",
        "gsc_auth_method": gsc_auth_method,
        "page_type": page_type,
        "template_key": template_key,
        "page_copy_quality": {
            "enabled": page_quality_enabled,
            "correction_enabled": page_copy_correction_active,
            "page_quality_policy_version": page_quality["page_quality_policy_version"],
            "adaptive_policy_version": page_quality["adaptive_policy_version"],
            "owned_page_mapping_version": page_quality["owned_page_mapping_version"],
            "source_asset_manifest_version": page_quality.get(
                "source_asset_manifest_version",
                "",
            ),
            "claim_bound_renderer_version": active_claim_bound_renderer_version,
            "source_block_plan_version": active_source_block_plan_version,
            "claim_bound_rendering": claim_bound_rendering_active,
            "guidance_profile_id": (
                page_copy_guidance.id if page_copy_guidance else ""
            ),
            "guidance_profile_version": (
                page_copy_guidance.version if page_copy_guidance else ""
            ),
        },
        "generation_requested": {
            "meta": bool(gen_meta),
            "faqs": bool(gen_faqs),
            "page_copy": bool(gen_page_copy),
        },
        "input_signal_counts": {
            "manual_keywords": len(manual_kws),
            "dfs_ranked": 0,
            "gsc_queries": 0,
            "keyword_pool": 0,
            "ranked_keywords": 0,
            "serp_organic": 0,
            "paa_questions": 0,
            "ai_overview_sections": 0,
            "competitor_candidates": 0,
            "competitor_scrape_successes": 0,
            "competitor_rejected": 0,
            "competitors_scraped": 0,
            "page_context_chars": 0,
            "scraped_page_chars": 0,
        },
        "scrape": {
            "page_context_success": False,
            "page_context_source": "",
            "page_context_error": "",
            "requested_provider": "",
            "content_mode": "default",
            "fallback_used": False,
            "raw_response_chars": 0,
            "retained_context_chars": 0,
            "capture_version": "",
            "quality_diagnostics": {},
            "client_existing_content_success": False,
        },
        "output_counts": {
            "faq_items": 0,
            "sections": 0,
            "word_count": 0,
        },
        "duration_ms": 0,
    }

    def _finish_diagnostics() -> dict:
        run_diagnostics["duration_ms"] = int((time.monotonic() - row_started_at) * 1000)
        return run_diagnostics

    def _empty(status: str) -> dict:
        return {
            "url": url, "primary_keyword": None, "keyword_source": status,
            "gsc_auth_method": gsc_auth_method,
            "generated_title": None, "generated_description": None, "optimised_h1": None,
            "faq_items": [], "faq_schema": None,
            "word_count": 0, "template_name": None,
            "competitor_urls": [], "docx_b64": None, "status": status,
            "run_diagnostics": _finish_diagnostics(),
        }

    if not url or not url.startswith("http"):
        return _empty("skipped: invalid URL")

    # ─────────────────────────────────────────────────────────────────────
    # STEP 1 — Keyword pipeline (shared across all outputs)
    # ─────────────────────────────────────────────────────────────────────
    step("fetching DFS ranked keywords...")
    dfs_ranked = []
    try:
        dfs_ranked = get_ranked_keywords_for_url(url, dfs_login, dfs_password, location_code)
        run_diagnostics["input_signal_counts"]["dfs_ranked"] = len(dfs_ranked)
        step("DFS ranked: " + str(len(dfs_ranked)) + " keywords found")
    except Exception as exc:
        log_safe_exception(logger, "aio.keywords.ranked_failed", exc)
        step("Ranked keyword data was unavailable; continuing.")

    # Optional GSC layer
    gsc_queries = []
    if use_gsc and gsc_client and site_url:
        step("fetching GSC queries...")
        try:
            gsc_queries = get_top_queries_for_url(gsc_client, site_url, url, top_n=10)
            run_diagnostics["input_signal_counts"]["gsc_queries"] = len(gsc_queries)
            step("GSC: " + str(len(gsc_queries)) + " queries")
        except Exception:
            pass

    all_kws = list({r["keyword"] for r in dfs_ranked} | set(manual_kws) | {q["query"] for q in gsc_queries})
    all_kws = [k for k in all_kws if k]

    vol_map = {}
    diff_map = {}
    if all_kws:
        step("fetching keyword volumes...")
        try:
            vol_map = get_search_volume(all_kws, dfs_login, dfs_password, location_code)
        except Exception as exc:
            log_safe_exception(logger, "aio.keywords.volume_failed", exc)
            step("Keyword volume data was unavailable; continuing.")
        try:
            diff_map = get_keyword_difficulty(all_kws, dfs_login, dfs_password, location_code)
        except Exception as exc:
            log_safe_exception(logger, "aio.keywords.difficulty_failed", exc)
            step("Keyword difficulty data was unavailable; continuing.")

    pool   = merge_keyword_pools(gsc_queries, dfs_ranked, manual_kws, vol_map, diff_map)
    run_diagnostics["input_signal_counts"]["keyword_pool"] = len(pool)
    pool   = [k for k in pool if k.get("volume", 0) >= min_volume]
    ranked = rank_keywords(pool, branded_terms, h1=h1, exclude_position_one=True)
    ranked = [k for k in ranked if not k.get("branded")]
    manual_primary = manual_kws[0] if manual_kws else ""

    if manual_primary:
        manual_key = manual_primary.lower()
        existing_manual = next(
            (k for k in pool if str(k.get("keyword", "")).strip().lower() == manual_key),
            {},
        )
        manual_entry = {
            **existing_manual,
            "keyword": manual_primary,
            "volume": existing_manual.get("volume", vol_map.get(manual_primary, 0)),
            "difficulty": existing_manual.get("difficulty", diff_map.get(manual_primary, 1)),
            "score": existing_manual.get("score", 1.0),
            "branded": False,
        }
        ranked = [manual_entry] + [
            k for k in ranked
            if str(k.get("keyword", "")).strip().lower() != manual_key
        ]
        keyword_source = "manual"
    else:
        keyword_source = "dfs+gsc" if gsc_queries else "dfs"

    run_diagnostics["input_signal_counts"]["ranked_keywords"] = len(ranked)

    if not ranked and manual_kws:
        ranked = [{"keyword": k, "volume": 10, "difficulty": 1, "score": 1.0} for k in manual_kws]
        run_diagnostics["input_signal_counts"]["ranked_keywords"] = len(ranked)
        keyword_source = "manual"

    if not ranked and not use_gsc and h1:
        ranked = [{"keyword": h1, "volume": 0, "difficulty": 50, "score": 0.0}]
        run_diagnostics["input_signal_counts"]["ranked_keywords"] = len(ranked)
        keyword_source = "h1 fallback"

    if not ranked:
        step("✗ no keywords found — skipping")
        return _empty("skipped: no keywords found")

    primary_keyword = ranked[0]["keyword"]
    if not manual_primary and primary_keyword.lower() in used_keywords:
        for r in ranked[1:]:
            if r["keyword"].lower() not in used_keywords:
                primary_keyword = r["keyword"]
                break
    used_keywords.add(primary_keyword.lower())

    if not h1 and primary_keyword:
        h1 = primary_keyword.title()

    step("keyword: " + primary_keyword)

    # ─────────────────────────────────────────────────────────────────────
    # STEP 2 — SERP (shared: organic URLs + PAA + AI Overview)
    # ─────────────────────────────────────────────────────────────────────
    step("fetching SERP...")
    serp_data       = {"organic": [], "paa_items": [], "ai_overview": ""}
    paa_questions   = []
    ai_overview     = ""
    ai_overview_sections = []
    organic_results = []
    try:
        serp_data       = get_serp_data(dfs_login, dfs_password, primary_keyword, location_code)
        if serp_data.get("error"):
            log_safe_external_failure(
                logger,
                "aio.serp.failed",
                serp_data.get("error"),
            )
            step("Search-result data was unavailable; continuing.")
        paa_questions   = serp_data.get("paa_items") or serp_data.get("paa") or []
        ai_overview_sections = serp_data.get("ai_overview_sections") or []
        ai_overview     = serp_data.get("ai_overview_raw") or serp_data.get("ai_overview") or ""
        organic_results = serp_data.get("organic") or []
        run_diagnostics["input_signal_counts"]["paa_questions"] = len(paa_questions)
        run_diagnostics["input_signal_counts"]["ai_overview_sections"] = len(ai_overview_sections)
        run_diagnostics["input_signal_counts"]["serp_organic"] = len(organic_results)
        step("SERP: " + ("AIO ✓" if ai_overview else "AIO ✗") + ", PAA: " + str(len(paa_questions)) + ", organic: " + str(len(organic_results)))
    except Exception as exc:
        log_safe_exception(logger, "aio.serp.failed", exc)
        step("Search-result data was unavailable; continuing.")

    # ─────────────────────────────────────────────────────────────────────
    # STEP 3 — Owned-page context (shared by strategy and every output)
    # ─────────────────────────────────────────────────────────────────────
    page_context = ""
    scraped_page_content = ""
    owned_page_scrape = {}
    owned_page_registry = None
    source_asset_manifest = None
    if gen_faqs or gen_meta or gen_page_copy:
        if settings.get("scrape_pages", True) and _owned_page_scraper_available(settings, scraper_override):
            step("scraping page context...")
            try:
                scrape_result = _scrape_owned_page_for_settings(
                    settings,
                    url,
                    scraper_override=scraper_override,
                    business_type=business_type,
                    page_type=page_type,
                )
                owned_page_scrape = scrape_result
                run_diagnostics["scrape"]["page_context_source"] = scrape_result.get("source") or ""
                run_diagnostics["scrape"]["requested_provider"] = scrape_result.get("requested_provider") or ""
                run_diagnostics["scrape"]["content_mode"] = scrape_result.get("mode") or "default"
                run_diagnostics["scrape"]["fallback_used"] = bool(scrape_result.get("fallback_used"))
                run_diagnostics["scrape"]["raw_response_chars"] = int(scrape_result.get("raw_chars") or 0)
                run_diagnostics["scrape"]["retained_context_chars"] = int(
                    scrape_result.get("cleaned_chars") or len(scrape_result.get("content") or "")
                )
                scrape_capture_version = str(
                    scrape_result.get("capture_version") or ""
                )
                run_diagnostics["scrape"]["capture_version"] = (
                    scrape_capture_version
                    if scrape_capture_version == AIO_OWNED_PAGE_CAPTURE_VERSION
                    else ""
                )
                run_diagnostics["scrape"]["quality_diagnostics"] = (
                    _safe_capture_quality_diagnostics(
                        scrape_result.get("quality_diagnostics")
                    )
                )
                if scrape_result.get("success"):
                    scraped_page_content = scrape_result["content"]
                    page_context = scraped_page_content
                    run_diagnostics["scrape"]["page_context_success"] = True
                    run_diagnostics["input_signal_counts"]["scraped_page_chars"] = len(scraped_page_content)
                    step("page context: " + str(len(page_context)) + " chars")
                else:
                    log_safe_external_failure(
                        logger,
                        "aio.scrape.owned_page_failed",
                        scrape_result.get("error"),
                    )
                    run_diagnostics["scrape"]["page_context_error"] = (
                        _OWNED_PAGE_CONTEXT_ERROR
                    )
                    scrape_result["error"] = _OWNED_PAGE_CONTEXT_ERROR
                    step("Owned-page context was unavailable; continuing without it.")
            except Exception as exc:
                log_safe_exception(
                    logger,
                    "aio.scrape.owned_page_failed",
                    exc,
                )
                owned_page_scrape = {
                    "success": False,
                    "error": _OWNED_PAGE_CONTEXT_ERROR,
                    "source": "",
                    "requested_provider": settings.get("scrape_provider", "jina"),
                    "mode": "ecommerce_collection" if is_ecommerce_collection_page(business_type, page_type) else "default",
                }
                run_diagnostics["scrape"]["page_context_error"] = (
                    _OWNED_PAGE_CONTEXT_ERROR
                )
                step("Owned-page context was unavailable; continuing without it.")

        if client_brief:
            page_context = (page_context + "\n\n" + client_brief).strip()
        run_diagnostics["input_signal_counts"]["page_context_chars"] = len(page_context)
        if (
            page_quality_policy
            and page_quality_policy.bounded_owned_page_reuse
        ):
            owned_page_registry = build_owned_page_registry(
                scraped_page_content,
                mapping_version=page_quality["owned_page_mapping_version"],
            )
            run_diagnostics["owned_page_mapping"] = {
                "version": owned_page_registry["version"],
                "registry_block_count": len(owned_page_registry["blocks"]),
                "source_char_count": owned_page_registry["source_char_count"],
                "retained_char_count": owned_page_registry["retained_char_count"],
                "truncated": owned_page_registry["truncated"],
            }
            source_asset_version = page_quality.get(
                "source_asset_manifest_version",
                "",
            )
            if source_asset_version:
                source_asset_manifest = build_source_asset_manifest(
                    owned_page_registry,
                    manifest_version=source_asset_version,
                )
                manifest_diagnostics = source_asset_manifest["diagnostics"]
                run_diagnostics["source_asset_manifest"] = {
                    "version": source_asset_manifest["version"],
                    "manifest_hash": source_asset_manifest["manifest_hash"],
                    "asset_count": manifest_diagnostics["asset_count"],
                    "source_truncated": manifest_diagnostics[
                        "source_truncated"
                    ],
                    "registry_truncated": manifest_diagnostics[
                        "registry_truncated"
                    ],
                    "structured_assets_suppressed": manifest_diagnostics[
                        "structured_assets_suppressed"
                    ],
                }

    # ─────────────────────────────────────────────────────────────────────
    # STEP 4 — Competitor scraping (for page copy)
    # ─────────────────────────────────────────────────────────────────────
    competitor_urls_used = []
    competitor_section_map = {}
    kw_assignment = {}
    lsi_map = {}
    template = None
    resolved_template_key = template_key

    if gen_page_copy:
        step("resolving template...")
        if custom_template_text:
            template = parse_custom_template(custom_template_text, page_type)
        else:
            try:
                template = get_template(template_key)
            except ValueError:
                resolved_template_key = default_template_key_for_page_type(page_type)
                template = get_template(resolved_template_key)
        template = _template_for_page_copy(
            template,
            bool(gen_faqs),
            versioned_blog_h1=bool(
                page_quality_policy
                and page_quality_policy.exact_planned_headings
                and not custom_template_text
            ),
        )

        kw_assignment = assign_keywords_to_sections(ranked, template["sections"])
        competitor_section_map = {s["name"]: [] for s in template["sections"]}

        step("scraping competitors...")
        client_domain = urlparse(url).netloc
        if organic_results:
            try:
                scored = []
                for sr in organic_results[:8]:
                    comp_url = sr.get("url") or sr.get("link") or sr.get("relative_url") or ""
                    if not comp_url.startswith("http"):
                        continue
                    if client_domain and client_domain in urlparse(comp_url).netloc:
                        continue
                    run_diagnostics["input_signal_counts"]["competitor_candidates"] += 1
                    sc = scrape_url(comp_url, api_key=jina_key)
                    if not sc["success"]:
                        continue
                    run_diagnostics["input_signal_counts"]["competitor_scrape_successes"] += 1
                    if not is_editorial_competitor(sc, page_type):
                        run_diagnostics["input_signal_counts"]["competitor_rejected"] += 1
                        continue
                    sc["relevance"] = classify_competitor_relevance(sc, business_type, page_type)
                    sc["comp_url"]  = comp_url
                    scored.append(sc)
                scored.sort(key=lambda x: x["relevance"], reverse=True)
                top = scored[:3]
                competitor_urls_used = [c["comp_url"] for c in top]
                run_diagnostics["input_signal_counts"]["competitors_scraped"] = len(competitor_urls_used)
                if top:
                    competitor_section_map = map_competitor_sections(top, template["sections"])
                step("competitors: " + str(len(competitor_urls_used)) + " scraped")
            except Exception as exc:
                log_safe_exception(logger, "aio.scrape.competitor_failed", exc)
                step("Competitor evidence was unavailable; continuing.")
        else:
            step("competitors unavailable: no organic results")

    standard_built_in_generation_active = bool(
        gen_page_copy
        and _uses_standard_built_in_page_generation(
            resolved_template_key,
            custom_template_text,
        )
    )
    claim_bound_rendering_active = _claim_bound_rendering_is_active(
        page_quality,
        requested=bool(gen_page_copy),
        template_key=resolved_template_key,
        custom_template_text=custom_template_text,
    )
    active_claim_bound_renderer_version = (
        page_quality.get("claim_bound_renderer_version", "")
        if claim_bound_rendering_active
        else ""
    )
    active_source_block_plan_version = (
        page_quality.get("source_block_plan_version", "")
        if claim_bound_rendering_active
        else ""
    )
    generation_source_asset_manifest = (
        None
        if standard_built_in_generation_active
        else source_asset_manifest
    )
    run_diagnostics["page_copy_quality"].update({
        "claim_bound_renderer_version": active_claim_bound_renderer_version,
        "source_block_plan_version": active_source_block_plan_version,
        "claim_bound_rendering": claim_bound_rendering_active,
        "standard_built_in_generation": standard_built_in_generation_active,
        "exact_source_asset_preservation": bool(
            generation_source_asset_manifest
        ),
    })

    # ─────────────────────────────────────────────────────────────────────
    # STEP 5 — Generate strategy brief, then meta copy
    # ─────────────────────────────────────────────────────────────────────
    strategy_brief = {}
    strategy_status = "not_requested"
    strategy_issues = []
    adaptive_section_plan = []
    adaptive_template_family = ""
    required_strategy_outputs = [
        output
        for output, enabled in (
            ("meta", gen_meta),
            ("faq", gen_faqs),
            ("page_copy", gen_page_copy),
        )
        if enabled
    ]
    if gen_meta or gen_faqs or gen_page_copy:
        step("building strategy brief...")
        try:
            section_heading_keyword_assignments = {}
            if resolved_template_key == "collection_page":
                for section_name in (
                    "collection_story",
                    "collection_value",
                ):
                    assigned_keyword = str(
                        kw_assignment.get(section_name, {}).get("supporting")
                        or ""
                    ).strip()
                    if assigned_keyword:
                        section_heading_keyword_assignments[
                            section_name
                        ] = assigned_keyword
            strategy_brief = generate_strategy_brief(
                provider=provider,
                api_key=api_key,
                model=model,
                url=url,
                keyword=primary_keyword,
                page_type=page_type,
                business_type=business_type,
                brand_name=brand_name,
                h1=h1,
                brand_context=brand_context,
                client_brief=client_brief,
                evidence_client_brief=explicit_client_brief,
                page_context=scraped_page_content,
                ai_overview=ai_overview,
                paa_questions=paa_questions,
                competitor_section_map=competitor_section_map,
                template_sections=(template or {}).get("sections", []),
                section_heading_keyword_assignments=(
                    section_heading_keyword_assignments
                ),
                required_outputs=required_strategy_outputs,
                enable_page_planning=page_planning_enabled,
                owned_page_registry=owned_page_registry,
                source_asset_manifest=generation_source_asset_manifest,
                page_quality_policy=page_quality_policy,
                page_copy_correction_enabled=page_copy_correction_active,
                claim_bound_renderer_version=active_claim_bound_renderer_version,
                source_block_plan_version=active_source_block_plan_version,
            )
            if page_quality_enabled:
                strategy_brief = attach_depth_policies(
                    strategy_brief,
                    (
                        ""
                        if settings.get("custom_template_text", "").strip()
                        else resolved_template_key
                    ),
                    page_quality["adaptive_policy_version"],
                )
            strategy_issues = strategy_brief_issues(
                strategy_brief,
                (template or {}).get("sections", []),
                required_strategy_outputs,
            )
            strategy_status = "ready" if not strategy_issues else "needs_review"
            step("strategy brief ready" if not strategy_issues else "strategy brief needs review")
        except Exception as exc:
            log_safe_exception(logger, "aio.strategy.failed", exc)
            strategy_brief = {}
            strategy_status = "unavailable"
            strategy_issues = [_STRATEGY_BRIEF_ERROR]
            step("Strategy brief was unavailable; continuing.")

    if gen_page_copy and template:
        adaptive_key = "" if settings.get("custom_template_text", "").strip() else resolved_template_key
        template, adaptive_section_plan = _adapt_page_template_for_generation(
            template,
            adaptive_key,
            strategy_brief,
            adaptive_policy_version=(
                page_quality["adaptive_policy_version"]
                if page_quality_enabled
                else ""
            ),
            source_asset_manifest_version=page_quality.get(
                "source_asset_manifest_version",
                "",
            ),
            correction_evidence_contract=page_copy_correction_active,
        )
        adaptive_template_family = str(template.get("_adaptive_family") or "")
        adaptive_mode_counts = {
            mode: sum(1 for item in adaptive_section_plan if item.get("mode") == mode)
            for mode in ("full", "compact", "omit")
        }
        run_diagnostics["adaptive_template"] = {
            "family": adaptive_template_family,
            "sections": adaptive_section_plan,
        }
        step(
            "template plan: "
            + str(adaptive_mode_counts["full"])
            + " full, "
            + str(adaptive_mode_counts["compact"])
            + " compact, "
            + str(adaptive_mode_counts["omit"])
            + " omitted"
        )

    generated_title = None
    generated_description = None
    optimised_h1 = None
    input_h1_for_qa = h1
    evidence_contract_ready = bool(
        strategy_status == "ready" and strategy_brief.get("verified_facts")
    )

    if gen_meta:
        step("generating meta copy...")
        try:
            meta_context_parts = []
            if scraped_page_content and not evidence_contract_ready:
                meta_context_parts.append("SCRAPED PAGE CONTENT:\n" + scraped_page_content[:10000])
            if client_brief and not evidence_contract_ready:
                meta_context_parts.append("CLIENT BRIEF:\n" + client_brief)
            meta_result = generate_copy(
                provider=provider,
                api_key=api_key,
                model=model,
                url=url,
                keyword=primary_keyword,
                page_type=page_type,
                brand_name=brand_name if include_brand else "",
                forbidden_phrases=forbidden_phrase_text,
                context="\n\n".join(meta_context_parts),
                brand_context=brand_style_context if evidence_contract_ready else brand_context,
                business_type=business_type,
                h1=h1,
                strategy_brief=strategy_brief,
            )
            generated_title       = meta_result.get("title", "")
            generated_description = meta_result.get("description", "")
            optimised_h1          = meta_result.get("h1_optimised", "")
            # Use optimized H1 as page H1 if we didn't have one
            if not h1 and optimised_h1:
                h1 = optimised_h1
            step("✓ meta — title: " + str(len(generated_title or "")) + " chars, desc: " + str(len(generated_description or "")) + " chars")
        except Exception as exc:
            log_safe_exception(logger, "aio.meta.failed", exc)
            step("Meta copy could not be generated; continuing with other requested outputs.")

    # ─────────────────────────────────────────────────────────────────────
    # STEP 6 — Generate FAQs
    # ─────────────────────────────────────────────────────────────────────
    faq_items  = []
    faq_schema = None
    faq_script = None

    if gen_faqs:
        step("generating FAQs...")
        try:
            ai_ov_for_faq = ai_overview
            paa_for_faq   = paa_questions

            faq_items = generate_faq(
                provider=provider,
                api_key=api_key,
                model=model,
                keyword=primary_keyword,
                page_type=page_type,
                brand_name=brand_name if include_brand else "",
                business_type=business_type,
                h1=h1,
                num_faqs=num_faqs,
                paa_items=paa_for_faq,
                ai_overview_sections=ai_overview_sections,
                ai_overview_raw=ai_ov_for_faq,
                forbidden_phrases=forbidden_phrases,
                page_context=page_context,
                brand_profile=brand_profile,
            )
            faq_items, faqs_trimmed = _limit_faq_items(faq_items, num_faqs)
            if faqs_trimmed:
                step("FAQs trimmed to requested count: " + str(len(faq_items)))
            step("✓ FAQs: " + str(len(faq_items)) + " generated")

            # Build schema
            from utils.dfs import _extract_ai_overview_text
            try:
                from faq_saas_schema import build_faq_schema
            except ImportError:
                pass

            # Inline FAQ schema builder
            def _build_faq_schema(items):
                entities = [{"@type": "Question", "name": i["question"], "acceptedAnswer": {"@type": "Answer", "text": i["answer"]}} for i in items]
                schema = {"@context": "https://schema.org", "@type": "FAQPage", "mainEntity": entities}
                import json
                raw_json = json.dumps(schema, ensure_ascii=False, indent=2)
                script = f"<script type=\"application/ld+json\">\n{raw_json}\n</script>"
                return raw_json, script

            faq_schema, faq_script = _build_faq_schema(faq_items)
            run_diagnostics["output_counts"]["faq_items"] = len(faq_items)

        except Exception as exc:
            log_safe_exception(logger, "aio.faq.failed", exc)
            step("FAQs could not be generated; continuing with other requested outputs.")

    # ─────────────────────────────────────────────────────────────────────
    # STEP 7 — Generate full page copy
    # ─────────────────────────────────────────────────────────────────────
    section_results = {}
    full_page       = ""
    word_count      = 0
    page_copy_quality_blocked = False
    page_copy_quality_block_reasons = []
    source_block_plan = None
    if claim_bound_rendering_active:
        page_copy_canonical_h1 = _claim_bound_canonical_h1(
            source_asset_manifest,
            input_h1_for_qa,
            primary_keyword,
            forbidden_phrase_list,
        )
    else:
        page_copy_canonical_h1 = (
            (optimised_h1 or h1)
            if (
                page_quality_policy
                and page_quality_policy.exact_planned_headings
            )
            else h1
        )

    if gen_page_copy and template:
        step("generating page copy (" + str(len(template["sections"])) + " sections)...")

        # LSI keywords
        supporting_kws = list({v.get("supporting", "") for v in kw_assignment.values() if v.get("supporting")})
        for sk in supporting_kws[:3]:
            try:
                ideas = get_keyword_ideas(sk, dfs_login, dfs_password, location_code, limit=10)
                lsi_map[sk] = [i["keyword"] for i in ideas[:3]]
            except Exception:
                lsi_map[sk] = []
                step("Related keyword ideas unavailable; continuing without them.")

        # Client existing content
        client_existing_content = ""
        if scraped_page_content:
            client_existing_content = scraped_page_content[:800]
            run_diagnostics["scrape"]["client_existing_content_success"] = True
        elif (
            settings.get("scrape_pages", True)
            and settings.get("scrape_provider", "jina") == "jina"
            and scraper_override != "firecrawl"
        ):
            try:
                existing = scrape_url(url, api_key=jina_key)
                if existing["success"]:
                    client_existing_content = existing.get("body_text", "")[:800]
                    run_diagnostics["scrape"]["client_existing_content_success"] = bool(client_existing_content)
            except Exception:
                pass

        try:
            def on_section(i, total, label):
                step("section " + str(i+1) + "/" + str(total) + ": " + label)
                if _is_cancelled(sb, job_id, user_id):
                    raise InterruptedError("cancelled")

            page_result = generate_page(
                template=template,
                keyword_assignment=kw_assignment,
                lsi_keywords=lsi_map,
                business_type=business_type,
                brand_name=brand_name,
                h1=page_copy_canonical_h1,
                page_type=page_type,
                paa_questions=paa_questions,
                ai_overview=ai_overview,
                competitor_section_map=competitor_section_map,
                client_brief=client_brief,
                client_existing_content=client_existing_content,
                provider=provider,
                api_key=api_key,
                model=model,
                forbidden_phrases=forbidden_phrase_text,
                progress_callback=on_section,
                strategy_brief=strategy_brief,
                brand_style_context=brand_style_context if evidence_contract_ready else "",
                page_copy_guidance=page_copy_guidance,
                page_quality_policy=page_quality_policy,
                page_copy_correction_enabled=page_copy_correction_active,
                claim_bound_renderer_version=active_claim_bound_renderer_version,
                source_block_plan_version=active_source_block_plan_version,
                source_asset_manifest=generation_source_asset_manifest,
            )
            page_copy_quality_blocked = bool(
                page_result.get("_quality_blocked")
            )
            page_copy_quality_block_reasons = [
                str(reason)
                for reason in page_result.get("_quality_block_reasons") or []
                if str(reason or "")
            ]
            source_block_plan = page_result.get("_source_block_plan")
            section_results = {k: v for k, v in page_result.items() if not k.startswith("_")}
            if (
                resolved_template_key == "collection_page"
                and not claim_bound_rendering_active
            ):
                section_results = {
                    name: normalise_collection_references(
                        text,
                        primary_keyword or h1,
                        protected_exact_phrases=(
                            _source_asset_exact_phrases(
                                strategy_brief,
                                name,
                            )
                        ),
                    )
                    for name, text in section_results.items()
                }
            full_page       = page_result.get("_full_page", "")
            word_count      = page_result.get("_word_count", 0)
            if claim_bound_rendering_active:
                h1_replaced = False
            elif (
                page_quality_policy
                and page_quality_policy.exact_planned_headings
            ):
                section_results, h1_replaced = _enforce_v1_canonical_page_h1(
                    section_results,
                    template,
                    page_copy_canonical_h1,
                )
            else:
                section_results, h1_replaced = _enforce_canonical_page_h1(
                    section_results,
                    optimised_h1 or "",
                )
            if h1_replaced:
                full_page = _assemble_full_page_copy(section_results, template)
                word_count = len(full_page.split())
                step(
                    "page copy H1 aligned to canonical H1"
                    if (
                        page_quality_policy
                        and page_quality_policy.exact_planned_headings
                    )
                    else "page copy H1 aligned to meta H1"
                )
            run_diagnostics["output_counts"]["sections"] = len(section_results)
            run_diagnostics["output_counts"]["word_count"] = word_count
            step("✓ page copy: " + str(word_count) + " words")
        except InterruptedError:
            raise
        except Exception as exc:
            log_safe_exception(logger, "aio.page_copy.failed", exc)
            if claim_bound_rendering_active:
                page_copy_quality_blocked = True
                page_copy_quality_block_reasons = [
                    "claim_bound_renderer_failed"
                ]
                section_results = {}
                full_page = ""
                word_count = 0
                step(
                    "Evidence-locked page copy was withheld because its "
                    "source contract could not be completed safely."
                )
            else:
                step("Page copy could not be generated; continuing with other requested outputs.")

    # ─────────────────────────────────────────────────────────────────────
    # STEP 8 — Finalise generated output
    # ─────────────────────────────────────────────────────────────────────
    if gen_page_copy and section_results:
        if claim_bound_rendering_active:
            pass
        elif (
            page_quality_policy
            and page_quality_policy.exact_planned_headings
        ):
            section_results, _ = _enforce_v1_canonical_page_h1(
                section_results,
                template,
                page_copy_canonical_h1,
            )
        else:
            section_results, _ = _enforce_canonical_page_h1(
                section_results,
                optimised_h1 or "",
            )
        full_page = _assemble_full_page_copy(section_results, template)
        word_count = len(full_page.split())

    if not input_h1_for_qa and optimised_h1:
        h1 = optimised_h1
    run_diagnostics["output_counts"]["faq_items"] = len(faq_items)
    run_diagnostics["output_counts"]["sections"] = len(section_results)
    run_diagnostics["output_counts"]["word_count"] = word_count

    # STEP 9 — Build combined docx
    docx_b64 = None
    step("building combined docx...")
    try:
        docx_bytes = _build_combined_docx(
            url=url,
            h1=h1,
            primary_keyword=primary_keyword,
            page_type=page_type,
            template=template,
            generated_title=generated_title,
            generated_description=generated_description,
            optimised_h1=optimised_h1,
            faq_items=faq_items,
            faq_schema=faq_schema,
            section_results=section_results,
            word_count=word_count,
            competitor_urls=competitor_urls_used,
            gen_meta=gen_meta,
            gen_faqs=gen_faqs,
            gen_page_copy=gen_page_copy,
            keyword_assignment=kw_assignment,
            page_quality_policy_version=page_quality["page_quality_policy_version"],
            claim_bound_renderer_version=active_claim_bound_renderer_version,
            page_copy_canonical_h1=page_copy_canonical_h1,
        )
        docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")
        step("✓ done")
    except Exception as exc:
        log_safe_exception(logger, "aio.docx.failed", exc)
        step("DOCX export could not be built; generated results remain available.")

    stored_competitor_section_map = {
        section: [str(excerpt)[:500] for excerpt in excerpts[:3] if str(excerpt).strip()]
        for section, excerpts in (competitor_section_map or {}).items()
        if excerpts
    }
    content_gap_summary = _build_content_gap_summary(stored_competitor_section_map, section_results)
    qa_flags = _collect_qa_flags(
        gen_meta=gen_meta,
        gen_faqs=gen_faqs,
        gen_page_copy=gen_page_copy,
        generated_title=generated_title or "",
        generated_description=generated_description or "",
        optimised_h1=optimised_h1 or "",
        input_h1=input_h1_for_qa,
        primary_keyword=primary_keyword,
        faq_items=faq_items,
        requested_faq_count=num_faqs,
        section_results=section_results,
        forbidden_phrases=forbidden_phrase_list,
        template=template,
        brand_name=brand_name,
        business_type=business_type,
        page_type=page_type,
        strategy_brief=strategy_brief,
        page_quality_policy_version=page_quality["page_quality_policy_version"],
        page_copy_correction_enabled=page_copy_correction_active,
        page_copy_quality_block_reasons=page_copy_quality_block_reasons,
    )
    _add_strategy_qa_flag(qa_flags, strategy_status, strategy_issues)
    quality_diagnostics = (
        _build_page_quality_diagnostics(
            strategy_brief=strategy_brief,
            template=template,
            section_results=section_results,
            page_quality=page_quality,
        )
        if page_quality_enabled
        else None
    )

    if owned_page_scrape.get("success"):
        source_labels = {
            "live": "Jina live",
            "live_selector_recovery": "Jina live, selector recovery",
            "cached_fallback": "Jina cached fallback",
            "firecrawl": "Firecrawl",
        }
        scrape_status = "Success: " + source_labels.get(
            str(owned_page_scrape.get("source") or ""),
            str(owned_page_scrape.get("source") or "owned page"),
        )
    elif owned_page_scrape:
        scrape_status = "Failed: " + str(
            owned_page_scrape.get("error") or "No page context was retained."
        )
    elif settings.get("scrape_pages", True):
        scrape_status = "Not available"
    else:
        scrape_status = "Disabled"

    return {
        "url":                  url,
        "h1":                   h1,
        "input_h1":             input_h1_for_qa,
        "primary_keyword":      primary_keyword,
        "keyword_source":       keyword_source,
        "model":                model or "",
        "gsc_auth_method":      gsc_auth_method,
        "kw_volume":            (ranked[0].get("volume") if ranked else None),
        "generated_title":      generated_title,
        "generated_description": generated_description,
        "optimised_h1":         optimised_h1,
        "title_length":         len(generated_title or ""),
        "description_length":   len(generated_description or ""),
        "faq_items":            faq_items,
        "faq_schema":           faq_schema,
        "faq_count":            len(faq_items),
        "word_count":           word_count,
        "template_name":        template["name"] if template else None,
        "section_results":      section_results,
        "keyword_assignment":   kw_assignment if gen_page_copy else {},
        "lsi_keywords":         lsi_map if gen_page_copy else {},
        "competitor_section_map": stored_competitor_section_map if gen_page_copy else {},
        "content_gap_summary":  content_gap_summary if gen_page_copy else [],
        "strategy_brief":       strategy_brief,
        "strategy_status":      strategy_status,
        "strategy_issues":      strategy_issues,
        "page_quality_policy_version": page_quality["page_quality_policy_version"] or None,
        "adaptive_policy_version": page_quality["adaptive_policy_version"] or None,
        "owned_page_mapping_version": page_quality["owned_page_mapping_version"] or None,
        "source_asset_manifest_version": page_quality.get(
            "source_asset_manifest_version"
        ) or None,
        "claim_bound_renderer_version": (
            active_claim_bound_renderer_version or None
        ),
        "source_block_plan_version": (
            active_source_block_plan_version or None
        ),
        "page_copy_quality_blocked": page_copy_quality_blocked,
        "page_copy_quality_block_reasons": page_copy_quality_block_reasons,
        "source_block_plan": source_block_plan,
        "page_copy_guidance": (
            {
                "id": page_copy_guidance.id,
                "label": page_copy_guidance.label,
                "version": page_copy_guidance.version,
            }
            if page_copy_guidance
            else None
        ),
        "quality_diagnostics": quality_diagnostics,
        "scrape_status":        scrape_status,
        "page_context_preview": scraped_page_content,
        "adaptive_template_family": adaptive_template_family,
        "adaptive_section_plan": adaptive_section_plan,
        "competitor_urls":      competitor_urls_used,
        "docx_b64":             docx_b64,
        "qa_flags":             qa_flags,
        "run_diagnostics":      _finish_diagnostics(),
        "status":               _qa_status(qa_flags),
    }


def _append_page_copy_markdown(
    doc,
    section: dict,
    text: str,
    *,
    skip_heading_text: str = "",
):
    actual_heading = _first_markdown_heading(text)
    heading_level = str(section.get("heading_level") or "").casefold()
    if not actual_heading and heading_level in {"h1", "h2", "h3"}:
        fallback_heading = (
            str(section.get("planned_heading") or "").strip()
            or str(section.get("label") or section.get("name") or "").strip()
        )
        if fallback_heading:
            doc.add_heading(fallback_heading, level=int(heading_level[1]))

    skipped_heading = False
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+?)\s*$", line)
        if heading_match:
            heading_text = heading_match.group(2).strip()
            if (
                not skipped_heading
                and skip_heading_text
                and heading_text.casefold() == skip_heading_text.casefold()
            ):
                skipped_heading = True
                continue
            doc.add_heading(heading_text, level=len(heading_match.group(1)))
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(
                re.sub(r"^\d+\.\s+", "", line),
                style="List Number",
            )
        else:
            doc.add_paragraph(line)


def _build_combined_docx(
    url, h1, primary_keyword, page_type, template,
    generated_title, generated_description, optimised_h1,
    faq_items, faq_schema, section_results, word_count, competitor_urls,
    gen_meta, gen_faqs, gen_page_copy,
    keyword_assignment=None,
    page_quality_policy_version="",
    claim_bound_renderer_version="",
    page_copy_canonical_h1="",
):
    """Build a single docx with meta, FAQs, and page copy in one document."""
    keyword_assignment = keyword_assignment or {}
    is_versioned_page_copy = bool(str(page_quality_policy_version or "").strip())
    claim_bound_page_copy = bool(
        gen_page_copy
        and claim_bound_renderer_version == CLAIM_BOUND_RENDERER_VERSION
    )
    canonical_document_h1 = (
        page_copy_canonical_h1
        if claim_bound_page_copy
        else ((optimised_h1 or h1) if is_versioned_page_copy else h1)
    )
    if gen_page_copy and not gen_meta and not gen_faqs and template:
        return build_docx(
            url=url,
            page_type=page_type,
            template_name=template["name"],
            primary_keyword=primary_keyword,
            section_results=section_results,
            template_sections=template.get("sections", []),
            keyword_assignment=keyword_assignment,
            word_count=word_count,
            competitor_urls=competitor_urls,
            h1=canonical_document_h1,
        )

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # Title
    title_para = doc.add_heading(canonical_document_h1 or url, level=1)

    # Metadata table
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = "Table Grid"
    meta_table.rows[0].cells[0].text = "URL"
    meta_table.rows[0].cells[1].text = url

    for label, value in [
        ("Primary Keyword", primary_keyword or ""),
        ("Page Type", page_type or ""),
        ("Template", template["name"] if template else ""),
        ("Word Count", str(word_count) if word_count else ""),
    ]:
        row = meta_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    doc.add_paragraph("")

    # META COPY
    if gen_meta and (generated_title or generated_description or optimised_h1):
        doc.add_heading("Meta Copy", level=2)
        if generated_title:
            p = doc.add_paragraph()
            p.add_run("Title Tag: ").bold = True
            p.add_run(generated_title)
            char_note = f"  ({len(generated_title)} chars{'  ⚠ over 80' if len(generated_title) > META_TITLE_PREFERRED_MAX else ''})"
            p.add_run(char_note)
        if generated_description:
            p = doc.add_paragraph()
            p.add_run("Meta Description: ").bold = True
            p.add_run(generated_description)
            char_note = f"  ({len(generated_description)} chars{'  ⚠ over 180' if len(generated_description) > META_DESCRIPTION_PREFERRED_MAX else ''})"
            p.add_run(char_note)
        if optimised_h1:
            p = doc.add_paragraph()
            p.add_run("Optimized H1: ").bold = True
            p.add_run(optimised_h1)
        doc.add_paragraph("")

    # PAGE COPY
    if gen_page_copy and section_results and template:
        doc.add_heading("Page Copy", level=2)
        for section in template.get("sections", []):
            sec_name = section["name"]
            text = section_results.get(sec_name, "")
            if text:
                if is_versioned_page_copy:
                    _append_page_copy_markdown(
                        doc,
                        section,
                        text,
                        skip_heading_text=(
                            canonical_document_h1
                            if str(section.get("heading_level") or "").casefold() == "h1"
                            else ""
                        ),
                    )
                else:
                    doc.add_heading(section["label"], level=3)
                    for paragraph in text.split("\n\n"):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
        doc.add_paragraph("")

    # FAQs
    if gen_faqs and faq_items:
        doc.add_heading("FAQs", level=2)
        for i, faq in enumerate(faq_items, 1):
            q = doc.add_paragraph()
            q.add_run(f"Q{i}: {faq['question']}").bold = True
            a = doc.add_paragraph()
            a.add_run(f"A: {faq['answer']}")
            doc.add_paragraph("")

        if faq_schema:
            doc.add_heading("FAQ Schema JSON-LD", level=3)
            doc.add_paragraph(f'<script type="application/ld+json">\n{faq_schema}\n</script>')

    # Competitors
    if competitor_urls:
        doc.add_heading("Competitors Referenced", level=3)
        for cu in competitor_urls:
            doc.add_paragraph(cu, style="List Bullet")

    if gen_page_copy and template:
        doc.add_heading("Page Copy Diagnostics", level=3)
        diag_table = doc.add_table(rows=1, cols=4)
        diag_table.style = "Table Grid"
        for i, label in enumerate(["Section", "Words", "Keyword Slot", "Keyword Used"]):
            diag_table.rows[0].cells[i].text = label

        for section in template.get("sections", []):
            sec_name = section["name"]
            assign = keyword_assignment.get(sec_name, {})
            keyword_used = assign.get("primary") or assign.get("supporting") or ""
            text = section_results.get(sec_name, "")
            row = diag_table.add_row()
            row.cells[0].text = section["label"]
            row.cells[1].text = str(len(text.split()) if text else 0)
            row.cells[2].text = section.get("keyword_slot", "none")
            row.cells[3].text = keyword_used

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _process_job(
    job_id: str,
    rows: list,
    settings: dict,
    gsc_credentials: dict | None,
    brand_profile: dict = None,
    user_id: str = "",
):
    sb    = get_supabase()
    total = len(rows)

    _update_job(sb, job_id, user_id, {
        "status":       "running",
        "total_rows":   total,
        "current_step": "Starting...",
    })

    # Init GSC
    gsc_client = None
    if settings.get("use_gsc"):
        if not gsc_credentials:
            _update_job(sb, job_id, user_id, {"error": _GSC_UNAVAILABLE_ERROR})
        else:
            try:
                gsc_client = get_gsc_client(gsc_credentials)
            except GscOAuthConfigError:
                _update_job(sb, job_id, user_id, {"error": _GSC_CONFIG_ERROR})
            except RefreshError:
                if gsc_credentials.get("method") == "google_oauth":
                    _update_job(sb, job_id, user_id, {"error": _GSC_RECONNECT_ERROR})
                    ciphertext = gsc_credentials.get("refresh_token_ciphertext")
                    if ciphertext:
                        try:
                            mark_gsc_reconnect_required(sb, user_id, ciphertext)
                        except Exception:
                            pass
                else:
                    _update_job(sb, job_id, user_id, {"error": _GSC_UNAVAILABLE_ERROR})
            except Exception:
                _update_job(sb, job_id, user_id, {"error": _GSC_UNAVAILABLE_ERROR})
    gsc_auth_method = _safe_gsc_auth_method(settings, gsc_credentials, gsc_client)
    if settings.get("use_gsc"):
        _update_job(sb, job_id, user_id, {"current_step": f"GSC auth method: {gsc_auth_method}"})

    branded_terms = [b.strip() for b in settings.get("brand_name", "").split() if b.strip()]
    full_brand = settings.get("full_brand_name", "").strip()
    if full_brand:
        branded_terms = list(set(branded_terms + [w.lower() for w in re.findall(r"[a-zA-Z]+", full_brand) if len(w) >= 3]))
    branded_input = settings.get("branded_terms_input", "").strip()
    if branded_input:
        branded_terms = list(set(branded_terms + [t.strip().lower() for t in branded_input.splitlines() if t.strip()]))

    used_keywords: set = set()
    results = []

    for idx, row in enumerate(rows):
        url = (row.get("url") or "").strip()
        _update_job(sb, job_id, user_id, {"current_step": f"Row {idx+1}/{total}: starting — {url}"})

        if _is_cancelled(sb, job_id, user_id):
            _update_job(sb, job_id, user_id, {
                "status":       "cancelled",
                "current_step": f"Cancelled after {idx}/{total} rows.",
                "failed_rows":  sum(1 for r in results if _result_failed(r)),
            })
            return

        try:
            result = _process_single_row(
                row=row, settings=settings, gsc_client=gsc_client,
                branded_terms=branded_terms, used_keywords=used_keywords,
                sb=sb, job_id=job_id, row_num=idx + 1, total_rows=total,
                user_id=user_id,
                brand_profile=brand_profile,
                gsc_auth_method=gsc_auth_method,
            )
        except InterruptedError:
            _update_job(sb, job_id, user_id, {
                "status":       "cancelled",
                "current_step": f"Cancelled during row {idx + 1}.",
                "failed_rows":  sum(1 for r in results if _result_failed(r)),
                "results":      results,
            })
            return
        except Exception as exc:
            log_safe_exception(
                logger,
                "aio.row.failed",
                exc,
                job_id=job_id,
                row=idx + 1,
            )
            result = {
                "url": url,
                "error": _ROW_PROCESSING_ERROR,
                "status": "error",
                "word_count": 0,
                "docx_b64": None,
                "gsc_auth_method": gsc_auth_method,
            }

        result = _apply_cross_row_uniqueness_flags(result, results)
        results.append(result)
        _update_job(sb, job_id, user_id, {"completed_rows": idx + 1, "results": results})

        if _is_cancelled(sb, job_id, user_id):
            _update_job(sb, job_id, user_id, {
                "status":       "cancelled",
                "current_step": f"Cancelled after {idx + 1}/{total} rows.",
                "failed_rows":  sum(1 for r in results if _result_failed(r)),
            })
            return

    try:
        internal_link_suggestions = _build_internal_link_suggestions(results)
        final_step = "Done."
    except Exception:
        internal_link_suggestions = []
        final_step = "Done. Internal link suggestions unavailable."
    _update_job(sb, job_id, user_id, {
        "status":        "complete",
        "current_step":  final_step,
        "completed_rows": len(results),
        "failed_rows":   sum(1 for r in results if _result_failed(r)),
        "results":       results,
        "internal_link_suggestions": internal_link_suggestions,
    })


# ── Request models ─────────────────────────────────────────────────────────────

class AIORow(BaseModel):
    url: str
    keyword: str = ""
    page_type: str = "service"
    h1: str = ""
    template_key: str = ""
    gen_page_copy: bool = True
    gen_meta: bool = True
    gen_faqs: bool = True
    num_faqs: int | None = None


class AIOSettings(BaseModel):
    niche: str = ""
    provider: str = "Claude"
    model: str = ""
    api_key: str = ""
    dfs_login: str = ""
    dfs_password: str = ""
    business_type: str = "general"
    brand_name: str = ""
    full_brand_name: str = ""
    branded_terms_input: str = ""
    include_brand: bool = True
    forbidden_phrases: str = ""
    location_code: int = 2840
    min_volume: int = 10
    page_type: str = "service"
    template_key: str = "service_page"
    custom_template_text: str = ""
    client_brief: str = ""
    brand_profile_id: str = ""
    jina_api_key: str = ""
    scrape_pages: bool = True
    scrape_provider: Literal["jina", "firecrawl"] = "jina"
    firecrawl_fallback: bool = False
    use_gsc: bool = False
    site_url: str = ""
    gen_page_copy: bool = True
    gen_meta: bool = True
    gen_faqs: bool = True
    num_faqs: int = 5
    page_copy_guidance_profile_id: str = ""


class AIOJobRequest(BaseModel):
    name: str = ""
    rows: list[AIORow]
    settings: AIOSettings


@router.post("/run")
def run_aio_job(
    request: AIOJobRequest,
    background_tasks: BackgroundTasks,
    user=Depends(get_current_user),
    sb=Depends(get_supabase),
):
    job_id = str(uuid.uuid4())
    enforce_job_start(sb, user.id, "all-in-one", len(request.rows), 50)
    enforce_rate_limit(sb, user.id, "all-in-one", "job-create", 10)
    submitted_settings, _ = _new_job_page_quality_settings(
        request.settings.model_dump(),
        user.id,
        page_copy_requested=any(
            row.gen_page_copy and row.url.strip().startswith("http")
            for row in request.rows
        ),
    )
    try:
        runtime_settings = hydrate_job_settings(sb, user.id, submitted_settings)
    except Exception:
        raise HTTPException(
            status_code=503,
            detail="Saved credentials are temporarily unavailable.",
        ) from None
    if not runtime_settings.get("api_key") or not runtime_settings.get("dfs_password"):
        raise HTTPException(status_code=400, detail="Saved provider credentials are incomplete. Update Settings and try again.")
    if (
        request.settings.scrape_pages
        and request.settings.scrape_provider == "firecrawl"
        and not runtime_settings.get("firecrawl_api_key")
    ):
        raise HTTPException(
            status_code=400,
            detail="Add a Firecrawl API key in Settings before using Firecrawl as the primary scraper.",
        )

    gsc_credentials = None
    if request.settings.use_gsc:
        gsc_credentials = runtime_settings.get("_gsc_credentials")

    # Brand profile
    brand_profile = None
    client_profile_id = None
    if request.settings.brand_profile_id:
        try:
            bp = sb.table("brand_profiles").select("data").eq("id", request.settings.brand_profile_id).eq("user_id", user.id).execute()
            if bp.data:
                brand_profile = bp.data[0].get("data") or {}
                client_profile_id = request.settings.brand_profile_id
        except Exception:
            pass

    execute_active_job_write(lambda: sb.table("jobs").insert({
        "id":             job_id,
        "user_id":        user.id,
        "client_profile_id": client_profile_id,
        "name":           request.name or f"All in One — {len(request.rows)} URLs",
        "tool":           "all-in-one",
        "status":         "pending",
        "total_rows":     len(request.rows),
        "completed_rows": 0,
        "failed_rows":    0,
        "results":        [],
        "logs":           [],
        "rows":           [r.model_dump() for r in request.rows],
        "settings":       strip_secret_fields(submitted_settings),
        "current_step":   "Queued...",
    }).execute(), "all-in-one")

    background_tasks.add_task(
        _process_job,
        job_id=job_id,
        rows=[r.model_dump() for r in request.rows],
        settings=runtime_settings,
        gsc_credentials=gsc_credentials,
        brand_profile=brand_profile,
        user_id=user.id,
    )

    return {"job_id": job_id, "status": "running"}


@router.get("/page-copy-capabilities")
def page_copy_capabilities(
    response: Response,
    user=Depends(get_current_user),
):
    response.headers["Cache-Control"] = "private, no-store"
    payload = guidance_capability_payload(
        page_quality_creation_enabled(user.id)
    )
    payload["policy_versions"] = {
        "page_quality": PAGE_QUALITY_POLICY_VERSION,
        "adaptive": ADAPTIVE_POLICY_VERSION,
        "owned_page_mapping": OWNED_PAGE_MAPPING_VERSION,
        "source_asset_manifest": SOURCE_ASSET_MANIFEST_VERSION,
        "claim_bound_renderer": CLAIM_BOUND_RENDERER_VERSION,
        "source_block_plan": SOURCE_BLOCK_PLAN_VERSION,
    }
    return payload


@router.get("/templates")
def list_templates():
    result = {}
    for pt in ["blog", "case_study", "glossary", "homepage", "service", "local", "about", "contact", "product", "collection"]:
        t = get_templates_for_page_type(pt)
        result[pt] = [{"key": k, "name": v["name"], "description": v.get("description", "")} for k, v in t.items()]
    return result
